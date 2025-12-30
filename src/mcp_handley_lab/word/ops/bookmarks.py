"""Bookmark and caption operations.

Contains functions for:
- Listing and adding bookmarks
- Inserting cross-references
- Creating and listing captions
"""

from __future__ import annotations

import re
from typing import TYPE_CHECKING

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from docx import Document

from mcp_handley_lab.word.ops.core import (
    count_occurrence,
    iter_body_blocks,
    make_block_id,
    paragraph_kind_and_level,
    resolve_target,
)
from mcp_handley_lab.word.ops.headers import insert_field
from mcp_handley_lab.word.ops.revisions import _rev_xpath

# =============================================================================
# Constants
# =============================================================================

# Reserved bookmark prefixes (used internally by Word)
_RESERVED_BOOKMARK_PREFIXES = ("_Toc", "_Ref", "_Hlt", "_GoBack")


# =============================================================================
# Helper Functions
# =============================================================================


def _validate_bookmark_name(name: str) -> None:
    """Validate bookmark name per Word restrictions.

    Word requires:
    - Starts with a letter
    - No spaces (underscores allowed)
    - Max 40 characters
    """
    if not name:
        raise ValueError("Bookmark name cannot be empty")
    if not name[0].isalpha():
        raise ValueError("Bookmark name must start with a letter")
    if " " in name:
        raise ValueError("Bookmark name cannot contain spaces")
    if len(name) > 40:
        raise ValueError("Bookmark name cannot exceed 40 characters")


def _get_next_bookmark_id(doc: Document) -> int:
    """Get next available bookmark ID."""
    max_id = 0
    for start in _rev_xpath(doc.element, "//w:bookmarkStart"):
        bm_id = start.get(qn("w:id"))
        if bm_id:
            max_id = max(max_id, int(bm_id))
    return max_id + 1


def _is_reserved_bookmark(name: str) -> bool:
    """Check if bookmark name is a reserved Word internal bookmark."""
    return any(name.startswith(prefix) for prefix in _RESERVED_BOOKMARK_PREFIXES)


# =============================================================================
# Bookmark Operations
# =============================================================================


def build_bookmarks(doc: Document) -> list[dict]:
    """List all bookmarks in document.

    Skips reserved bookmarks (_Toc*, _Ref*, etc.) used internally by Word.
    Returns list of dicts with: id, name, block_id.
    """
    results = []
    seen_names = set()

    for start in _rev_xpath(doc.element, "//w:body//w:bookmarkStart"):
        bm_id = start.get(qn("w:id"))
        bm_name = start.get(qn("w:name"))

        if not bm_id or not bm_name:
            continue

        # Skip reserved bookmarks
        if _is_reserved_bookmark(bm_name):
            continue

        # Skip duplicates
        if bm_name in seen_names:
            continue
        seen_names.add(bm_name)

        # Find containing paragraph for block_id
        block_id = ""
        parent = start.getparent()
        while parent is not None:
            if parent.tag == qn("w:p"):
                p = Paragraph(parent, doc)
                kind, level = paragraph_kind_and_level(p)
                text = p.text or ""
                occurrence = count_occurrence(doc, kind, text, parent)
                block_id = make_block_id(kind, text, occurrence)
                break
            parent = parent.getparent()

        results.append(
            {
                "id": int(bm_id),
                "name": bm_name,
                "block_id": block_id,
            }
        )

    return results


def add_bookmark(doc: Document, name: str, paragraph: Paragraph) -> int:
    """Add bookmark at paragraph. Returns bookmark ID.

    Validates:
    - Name starts with letter
    - No spaces (Word restriction)
    - Unique within document
    """
    _validate_bookmark_name(name)

    # Check uniqueness
    for bm in build_bookmarks(doc):
        if bm["name"] == name:
            raise ValueError(f"Bookmark name already exists: {name}")

    bm_id = _get_next_bookmark_id(doc)
    p_el = paragraph._element

    # Insert bookmarkStart at beginning of paragraph
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bm_id))
    bookmark_start.set(qn("w:name"), name)
    p_el.insert(0, bookmark_start)

    # Insert bookmarkEnd at end of paragraph
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bm_id))
    p_el.append(bookmark_end)

    return bm_id


# =============================================================================
# Cross-References
# =============================================================================


def insert_cross_reference(
    paragraph: Paragraph, bookmark_name: str, ref_type: str = "text"
) -> None:
    """Insert cross-reference field.

    ref_type options:
    - 'text': REF <bookmark> field (bookmark text)
    - 'number': REF <bookmark> \\r field (paragraph number)
    - 'page': PAGEREF <bookmark> field (page number)

    Bookmark names are case-sensitive in field instructions.
    """
    if ref_type == "text":
        instr = f"REF {bookmark_name} \\h"
    elif ref_type == "number":
        instr = f"REF {bookmark_name} \\r \\h"
    elif ref_type == "page":
        instr = f"PAGEREF {bookmark_name} \\h"
    else:
        raise ValueError(
            f"Unknown ref_type: {ref_type}. Use 'text', 'number', or 'page'"
        )

    insert_field(paragraph, instr, uppercase=False, placeholder="[ref]")


# =============================================================================
# Captions
# =============================================================================


def insert_caption(
    doc: Document,
    target_id: str,
    label: str,
    caption_text: str,
    position: str = "below",
) -> str:
    """Insert caption for table/image. Returns caption block ID.

    Creates paragraph with:
    1. Style "Caption"
    2. Label text (e.g., "Figure ")
    3. SEQ <label> field
    4. Separator (": ")
    5. Caption text

    Args:
        doc: Document to modify
        target_id: Block ID of element to caption
        label: Caption label (e.g., "Figure", "Table")
        caption_text: Text after the number
        position: "below" or "above" the target element
    """
    target = resolve_target(doc, target_id)

    # Create caption paragraph
    caption_p = doc.add_paragraph(style="Caption")

    # Add label text
    caption_p.add_run(f"{label} ")

    # Add SEQ field for auto-numbering
    insert_field(caption_p, f"SEQ {label}", uppercase=False)

    # Add separator and caption text
    caption_p.add_run(f": {caption_text}")

    # Move paragraph to correct position
    # For tables (even with hierarchical targets like table_X#r0c0), caption should be
    # relative to the table itself, not a cell paragraph inside it
    caption_el = caption_p._element
    if target.base_kind == "table":
        target_el = target.base_obj._tbl
    else:
        target_el = target.leaf_el

    if position == "above":
        target_el.addprevious(caption_el)
    else:  # below
        target_el.addnext(caption_el)

    # Generate block ID for the caption
    text = caption_p.text or ""
    occurrence = count_occurrence(doc, "paragraph", text, caption_el)
    return make_block_id("paragraph", text, occurrence)


def build_captions(doc: Document) -> list[dict]:
    """List all captions (paragraphs with style "Caption" containing SEQ fields).

    Returns list of dicts with: id, label, number, text, block_id, style.
    """
    results = []

    for kind, block_obj, element in iter_body_blocks(doc):
        if kind != "paragraph":
            continue

        p = block_obj
        style_name = p.style.name if p.style else ""

        if style_name != "Caption":
            continue

        # Extract text and look for SEQ field pattern
        full_text = p.text or ""

        # Parse the caption: "Label N: text" or "Label N – text"
        # The SEQ field result is embedded in the text
        match = re.match(r"^(\w+)\s+(\d+)[:–\-]\s*(.*)$", full_text)
        if match:
            label, number, caption_text = match.groups()
            number = int(number)
        else:
            # Can't parse - use full text
            label = "Unknown"
            number = 0
            caption_text = full_text

        # Generate block ID
        occurrence = count_occurrence(doc, kind, full_text, element)
        block_id = make_block_id(kind, full_text, occurrence)

        results.append(
            {
                "id": block_id,
                "label": label,
                "number": number,
                "text": caption_text,
                "block_id": block_id,
                "style": style_name,
            }
        )

    return results
