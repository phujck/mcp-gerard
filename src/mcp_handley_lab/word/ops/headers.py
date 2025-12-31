"""Header, footer, and field operations.

Contains functions for:
- Building/reading headers and footers
- Setting header/footer text
- Appending content to headers/footers
- Clearing headers/footers
- Field insertion (PAGE, NUMPAGES, etc.)
"""

from __future__ import annotations

import json
from typing import TYPE_CHECKING

from lxml import etree

from mcp_handley_lab.word.opc.constants import qn

if TYPE_CHECKING:
    pass

from mcp_handley_lab.word.models import HeaderFooterInfo
from mcp_handley_lab.word.ops.core import (
    _make_run_with,
    make_block_id,
    table_content_for_hash,
)
from mcp_handley_lab.word.ops.tables import populate_table

# EMU per inch for dimension calculations
_EMU_PER_INCH = 914400

# =============================================================================
# Field Insertion
# =============================================================================


def insert_field(
    p_el,
    field_code: str,
    display_text: str = "",
    *,
    uppercase: bool = True,
    placeholder: str = "1",
) -> None:
    """Insert a Word field into a paragraph.

    Pure OOXML: Takes w:p element (or Paragraph wrapper for backwards compat).

    Creates proper OXML field structure with separate runs for each part:
    begin, instruction, separator, result, and end markers.
    Supports any Word field code (PAGE, NUMPAGES, DATE, TIME, AUTHOR, etc.).

    Args:
        p_el: Paragraph element (w:p) or Paragraph wrapper (for backwards compat)
        uppercase: If True (default), uppercases field_code. Set False for case-sensitive
            fields like bookmark names in cross-references.
        placeholder: Default result text shown before field updates.
    """
    code = field_code.strip().upper() if uppercase else field_code

    # Support both lxml element and Paragraph wrapper (backwards compat)
    p = getattr(p_el, "_p", p_el)

    # Run 1: Field begin
    fld_char_begin = etree.Element(qn("w:fldChar"))
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    p.append(_make_run_with(fld_char_begin))

    # Run 2: Field instruction
    instr_text = etree.Element(qn("w:instrText"))
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = f" {code} "
    p.append(_make_run_with(instr_text))

    # Run 3: Field separator
    fld_char_sep = etree.Element(qn("w:fldChar"))
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    p.append(_make_run_with(fld_char_sep))

    # Run 4: Result text (placeholder shown before field updates)
    text_elem = etree.Element(qn("w:t"))
    text_elem.text = display_text or placeholder
    p.append(_make_run_with(text_elem))

    # Run 5: Field end
    fld_char_end = etree.Element(qn("w:fldChar"))
    fld_char_end.set(qn("w:fldCharType"), "end")
    p.append(_make_run_with(fld_char_end))


def insert_page_x_of_y(
    pkg_or_doc, section_index: int, location: str = "footer"
) -> None:
    """Insert 'Page X of Y' into header or footer.

    Duck-typed: works with WordPackage (pure OOXML) or python-docx Document.

    Args:
        pkg_or_doc: WordPackage or Document to modify
        section_index: 0-based section index
        location: 'header' or 'footer'
    """
    # Pure OOXML path
    if hasattr(pkg_or_doc, "has_part"):
        pkg = pkg_or_doc
        if location not in ("header", "footer"):
            raise ValueError(f"location must be 'header' or 'footer', got: {location}")
        partname, hf_root = _ensure_hf_part(pkg, section_index, location, "default")

        # Create paragraph
        p = etree.SubElement(hf_root, qn("w:p"))

        # Add "Page " text
        r1 = etree.SubElement(p, qn("w:r"))
        t1 = etree.SubElement(r1, qn("w:t"))
        t1.text = "Page "

        # Insert PAGE field
        insert_field(p, "PAGE")

        # Add " of " text
        r2 = etree.SubElement(p, qn("w:r"))
        t2 = etree.SubElement(r2, qn("w:t"))
        t2.text = " of "

        # Insert NUMPAGES field
        insert_field(p, "NUMPAGES")

        pkg.mark_xml_dirty(partname)
        return

    # python-docx path (legacy)
    doc = pkg_or_doc
    section = doc.sections[section_index]  # Let IndexError propagate
    hf = {"footer": section.footer, "header": section.header}[
        location
    ]  # KeyError if invalid
    hf.is_linked_to_previous = False
    p = hf.add_paragraph()
    p.add_run("Page ")
    insert_field(p, "PAGE")
    p.add_run(" of ")
    insert_field(p, "NUMPAGES")


# =============================================================================
# Header/Footer Reading
# =============================================================================


def _hf_text(hf) -> str | None:
    """Extract header/footer text. Returns None if linked to previous section."""
    return (
        None if hf.is_linked_to_previous else "\n".join(p.text for p in hf.paragraphs)
    )


def _extract_hf_text_ooxml(pkg, rId: str) -> str:
    """Extract text from header/footer part by rId (pure OOXML)."""
    doc_rels = pkg.get_rels("/word/document.xml")
    rel = doc_rels.get(rId)
    if not rel:
        return ""

    # Resolve relative path (e.g., "header1.xml" -> "/word/header1.xml")
    target = rel.target
    if not target.startswith("/"):
        target = f"/word/{target}"

    if not pkg.has_part(target):
        return ""

    hf_xml = pkg.get_xml(target)

    # Extract text from all w:t elements
    text_parts = []
    for p in hf_xml.iter(qn("w:p")):
        p_text_parts = []
        for t in p.iter(qn("w:t")):
            if t.text:
                p_text_parts.append(t.text)
        text_parts.append("".join(p_text_parts))

    return "\n".join(text_parts)


def _get_sectpr_hf_refs(sectPr) -> dict[str, dict[str, str]]:
    """Get header/footer references from sectPr element.

    Returns: {'header': {'default': rId, 'first': rId, 'even': rId},
              'footer': {'default': rId, 'first': rId, 'even': rId}}
    """
    refs = {"header": {}, "footer": {}}

    for ref in sectPr.findall(qn("w:headerReference")):
        hf_type = ref.get(qn("w:type")) or "default"
        rId = ref.get(qn("r:id"))
        if rId:
            refs["header"][hf_type] = rId

    for ref in sectPr.findall(qn("w:footerReference")):
        hf_type = ref.get(qn("w:type")) or "default"
        rId = ref.get(qn("r:id"))
        if rId:
            refs["footer"][hf_type] = rId

    return refs


def _has_title_pg(sectPr) -> bool:
    """Check if section has different first page header/footer (w:titlePg)."""
    title_pg = sectPr.find(qn("w:titlePg"))
    if title_pg is None:
        return False
    # Presence means True, but check w:val if present
    val = title_pg.get(qn("w:val"))
    return val != "0" if val else True


def _has_even_odd_headers(pkg) -> bool:
    """Check if document has different odd/even headers (from settings.xml)."""
    if not pkg.has_part("/word/settings.xml"):
        return False
    settings = pkg.get_xml("/word/settings.xml")
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is None:
        return False
    val = even_odd.get(qn("w:val"))
    return val != "0" if val else True


def _build_headers_footers_ooxml(pkg) -> list[HeaderFooterInfo]:
    """Build headers/footers list from pure OOXML."""
    result = []
    body = pkg.body

    # Find all section breaks (sectPr elements)
    sectPrs = []

    # Section breaks within paragraphs (w:p/w:pPr/w:sectPr)
    for p in body.findall(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is not None:
                sectPrs.append(sectPr)

    # Final section (w:body/w:sectPr)
    body_sectPr = body.find(qn("w:sectPr"))
    if body_sectPr is not None:
        sectPrs.append(body_sectPr)

    # Check document-level settings
    has_even_odd = _has_even_odd_headers(pkg)

    # Track previous section refs for "linked to previous" detection
    prev_refs = None

    for idx, sectPr in enumerate(sectPrs):
        refs = _get_sectpr_hf_refs(sectPr)
        has_first_page = _has_title_pg(sectPr)

        # Determine if linked to previous (no own reference = linked)
        header_linked = "default" not in refs["header"] and prev_refs is not None
        footer_linked = "default" not in refs["footer"] and prev_refs is not None

        # Get text for default header/footer
        header_text = None
        footer_text = None
        if not header_linked and "default" in refs["header"]:
            header_text = _extract_hf_text_ooxml(pkg, refs["header"]["default"])
        if not footer_linked and "default" in refs["footer"]:
            footer_text = _extract_hf_text_ooxml(pkg, refs["footer"]["default"])

        info = HeaderFooterInfo(
            section_index=idx,
            header_text=header_text,
            footer_text=footer_text,
            header_is_linked=header_linked,
            footer_is_linked=footer_linked,
            has_different_first_page=has_first_page,
            has_different_odd_even=has_even_odd,
        )

        # First page headers/footers
        if has_first_page:
            if "first" in refs["header"]:
                info.first_page_header_text = _extract_hf_text_ooxml(
                    pkg, refs["header"]["first"]
                )
            if "first" in refs["footer"]:
                info.first_page_footer_text = _extract_hf_text_ooxml(
                    pkg, refs["footer"]["first"]
                )

        # Even page headers/footers
        if has_even_odd:
            if "even" in refs["header"]:
                info.even_page_header_text = _extract_hf_text_ooxml(
                    pkg, refs["header"]["even"]
                )
            if "even" in refs["footer"]:
                info.even_page_footer_text = _extract_hf_text_ooxml(
                    pkg, refs["footer"]["even"]
                )

        result.append(info)
        prev_refs = refs

    return result


def build_headers_footers(pkg_or_doc) -> list[HeaderFooterInfo]:
    """Build list of HeaderFooterInfo for all sections.

    Args:
        pkg_or_doc: WordPackage or python-docx Document (duck-typed)
    """
    # Check if it's a WordPackage (pure OOXML path)
    if hasattr(pkg_or_doc, "document_xml"):
        return _build_headers_footers_ooxml(pkg_or_doc)

    # python-docx Document path (legacy)
    doc = pkg_or_doc
    result = []
    for idx, section in enumerate(doc.sections):
        hdr, ftr = section.header, section.footer
        info = HeaderFooterInfo(
            section_index=idx,
            header_text=_hf_text(hdr),
            footer_text=_hf_text(ftr),
            header_is_linked=hdr.is_linked_to_previous,
            footer_is_linked=ftr.is_linked_to_previous,
            has_different_first_page=section.different_first_page_header_footer,
            has_different_odd_even=doc.settings.odd_and_even_pages_header_footer,
        )
        if section.different_first_page_header_footer:
            fp_hdr, fp_ftr = section.first_page_header, section.first_page_footer
            info.first_page_header_text = _hf_text(fp_hdr)
            info.first_page_footer_text = _hf_text(fp_ftr)
        if doc.settings.odd_and_even_pages_header_footer:
            ev_hdr, ev_ftr = section.even_page_header, section.even_page_footer
            info.even_page_header_text = _hf_text(ev_hdr)
            info.even_page_footer_text = _hf_text(ev_ftr)
        result.append(info)
    return result


# =============================================================================
# Header/Footer Modification (Pure OOXML Helpers)
# =============================================================================

# Namespace for headers/footers
_HF_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_HF_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_HF_NSMAP = {"w": _HF_W_NS, "r": _HF_R_NS}

# Content types
_CT_HEADER = "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
_CT_FOOTER = "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"

# Relationship types
_RT_HEADER = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
)
_RT_FOOTER = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
)

# Location to (kind, hf_type) mapping
_LOCATION_MAP = {
    "header": ("header", "default"),
    "footer": ("footer", "default"),
    "first_page_header": ("header", "first"),
    "first_page_footer": ("footer", "first"),
    "even_page_header": ("header", "even"),
    "even_page_footer": ("footer", "even"),
}


def _get_all_sectprs(pkg) -> list[etree._Element]:
    """Get all sectPr elements in document order (section breaks + final)."""
    body = pkg.body
    sectPrs = []

    # Section breaks within paragraphs (w:p/w:pPr/w:sectPr)
    for p in body.findall(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is not None:
                sectPrs.append(sectPr)

    # Final section (w:body/w:sectPr)
    body_sectPr = body.find(qn("w:sectPr"))
    if body_sectPr is not None:
        sectPrs.append(body_sectPr)

    return sectPrs


def _get_next_hf_number(pkg, kind: str) -> int:
    """Find next available header/footer number (header1.xml, header2.xml, etc.)."""
    n = 1
    while pkg.has_part(f"/word/{kind}{n}.xml"):
        n += 1
    return n


def _create_minimal_hf_xml(kind: str) -> etree._Element:
    """Create minimal header/footer XML with one empty paragraph."""
    tag = "hdr" if kind == "header" else "ftr"
    root = etree.Element(qn(f"w:{tag}"), nsmap=_HF_NSMAP)
    # Add minimal empty paragraph
    p = etree.SubElement(root, qn("w:p"))
    etree.SubElement(p, qn("w:pPr"))
    return root


def _ensure_title_pg(sectPr) -> None:
    """Ensure w:titlePg exists in sectPr (enables first page header/footer)."""
    if sectPr.find(qn("w:titlePg")) is None:
        # Insert titlePg in correct position (before pgSz, pgMar, etc.)
        title_pg = etree.Element(qn("w:titlePg"))
        # Insert at beginning of sectPr
        sectPr.insert(0, title_pg)


def _ensure_even_odd_headers(pkg) -> None:
    """Ensure w:evenAndOddHeaders exists in settings.xml."""
    if not pkg.has_part("/word/settings.xml"):
        return

    settings = pkg.get_xml("/word/settings.xml")
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        even_odd = etree.Element(qn("w:evenAndOddHeaders"))
        # Insert near beginning of settings
        settings.insert(0, even_odd)
        pkg.mark_xml_dirty("/word/settings.xml")


def _ensure_hf_part(
    pkg, section_idx: int, kind: str, hf_type: str = "default"
) -> tuple[str, etree._Element]:
    """Ensure header/footer part exists for section. Returns (partname, root_el).

    Creates part if needed, adds relationship, and inserts reference in sectPr.
    Also sets required settings flags (titlePg, evenAndOddHeaders).
    """
    sectPrs = _get_all_sectprs(pkg)
    if section_idx >= len(sectPrs):
        raise IndexError(
            f"section_index {section_idx} out of range (document has {len(sectPrs)} sections)"
        )

    sectPr = sectPrs[section_idx]
    ref_tag = qn(f"w:{kind}Reference")
    r_id_qn = qn("r:id")
    type_qn = qn("w:type")

    # Check for existing reference
    for ref in sectPr.findall(ref_tag):
        ref_type = ref.get(type_qn) or "default"
        if ref_type == hf_type:
            # Found existing reference - resolve to part
            rId = ref.get(r_id_qn)
            target = pkg.resolve_rel_target("/word/document.xml", rId)
            return target, pkg.get_xml(target)

    # No existing reference - create new part
    n = _get_next_hf_number(pkg, kind)
    partname = f"/word/{kind}{n}.xml"
    target_relative = f"{kind}{n}.xml"
    ct = _CT_HEADER if kind == "header" else _CT_FOOTER
    rt = _RT_HEADER if kind == "header" else _RT_FOOTER

    # Create minimal header/footer XML
    hf_root = _create_minimal_hf_xml(kind)
    pkg.set_xml(partname, hf_root, ct)

    # Add relationship from document.xml
    rId = pkg.relate_to("/word/document.xml", target_relative, rt)

    # Insert reference in sectPr
    ref_el = etree.Element(ref_tag)
    ref_el.set(r_id_qn, rId)
    ref_el.set(type_qn, hf_type)
    # Insert at beginning of sectPr (before pgSz, etc.)
    sectPr.insert(0, ref_el)

    # Set required settings flags
    if hf_type == "first":
        _ensure_title_pg(sectPr)
    elif hf_type == "even":
        _ensure_even_odd_headers(pkg)

    pkg.mark_xml_dirty("/word/document.xml")
    return partname, hf_root


# =============================================================================
# Header/Footer Modification (Duck-Typed)
# =============================================================================


def set_header_footer_text(
    pkg_or_doc, section_index: int, text: str, location: str
) -> None:
    """Set header/footer text. Handles all types via location attribute name.

    Duck-typed: works with WordPackage (pure OOXML) or python-docx Document.
    """
    # Pure OOXML path
    if hasattr(pkg_or_doc, "has_part"):
        pkg = pkg_or_doc
        if location not in _LOCATION_MAP:
            raise ValueError(f"Unknown location: {location}")
        kind, hf_type = _LOCATION_MAP[location]
        partname, hf_root = _ensure_hf_part(pkg, section_index, kind, hf_type)

        # Clear existing content (keep root, remove all children)
        for child in list(hf_root):
            hf_root.remove(child)

        # Add paragraph with text
        p = etree.SubElement(hf_root, qn("w:p"))
        if text:
            r = etree.SubElement(p, qn("w:r"))
            t = etree.SubElement(r, qn("w:t"))
            t.text = text

        pkg.mark_xml_dirty(partname)
        return

    # python-docx path (legacy)
    doc = pkg_or_doc
    section = doc.sections[section_index]
    if location.startswith("first_page_"):
        section.different_first_page_header_footer = True
    elif location.startswith("even_page_"):
        doc.settings.odd_and_even_pages_header_footer = True
    hf = getattr(section, location)
    for p in list(hf.paragraphs):
        p._element.getparent().remove(p._element)
    hf.add_paragraph(text)


def append_to_header_footer(
    pkg_or_doc,
    section_index: int,
    content_type: str,
    content_data: str,
    location: str,
) -> str:
    """Append paragraph or table to header/footer. Returns element_id.

    Duck-typed: works with WordPackage (pure OOXML) or python-docx Document.
    """
    # Pure OOXML path
    if hasattr(pkg_or_doc, "has_part"):
        pkg = pkg_or_doc
        if location not in _LOCATION_MAP:
            raise ValueError(f"Unknown location: {location}")
        kind, hf_type = _LOCATION_MAP[location]
        partname, hf_root = _ensure_hf_part(pkg, section_index, kind, hf_type)

        if content_type == "paragraph":
            # Create paragraph with text
            p = etree.SubElement(hf_root, qn("w:p"))
            if content_data:
                r = etree.SubElement(p, qn("w:r"))
                t = etree.SubElement(r, qn("w:t"))
                t.text = content_data

            # Count occurrences for block ID
            occurrence = (
                sum(
                    1
                    for para in hf_root.findall(qn("w:p"))
                    if "".join(t.text or "" for t in para.iter(qn("w:t")))
                    == content_data
                )
                - 1
            )
            pkg.mark_xml_dirty(partname)
            return make_block_id("paragraph", content_data, occurrence)

        if content_type == "table":
            from mcp_handley_lab.word.ops.tables import _create_table_element

            table_data = json.loads(content_data)
            rows = len(table_data)
            cols = max((len(r) for r in table_data), default=1)
            # Create table element with dimensions (6 inch width / cols, in twips)
            # 1 inch = 1440 twips
            col_width_twips = int((6.0 * 1440) / cols)
            tbl = _create_table_element(rows, cols, col_width_twips)
            populate_table(tbl, table_data)
            hf_root.append(tbl)
            pkg.mark_xml_dirty(partname)
            return make_block_id("table", table_content_for_hash(tbl), 0)

        raise ValueError(f"Unknown content_type: {content_type}")

    # python-docx path (legacy)
    doc = pkg_or_doc
    hf = getattr(doc.sections[section_index], location)
    if content_type == "paragraph":
        p = hf.add_paragraph(content_data)
        occurrence = sum(1 for para in hf.paragraphs if para.text == p.text) - 1
        return make_block_id("paragraph", p.text, occurrence)
    if content_type == "table":
        from docx.shared import Inches

        table_data = json.loads(content_data)
        rows, cols = len(table_data), max((len(r) for r in table_data), default=1)
        tbl = hf.add_table(rows=rows, cols=cols, width=Inches(6))
        populate_table(tbl._tbl, table_data)  # Pass element, not wrapper
        return make_block_id("table", table_content_for_hash(tbl._tbl), 0)
    raise ValueError(f"Unknown content_type: {content_type}")


def clear_header_footer(pkg_or_doc, section_index: int, location: str) -> None:
    """Clear header/footer content. Unlinks from previous section first.

    Duck-typed: works with WordPackage (pure OOXML) or python-docx Document.
    """
    # Pure OOXML path
    if hasattr(pkg_or_doc, "has_part"):
        pkg = pkg_or_doc
        if location not in _LOCATION_MAP:
            raise ValueError(f"Unknown location: {location}")
        kind, hf_type = _LOCATION_MAP[location]
        partname, hf_root = _ensure_hf_part(pkg, section_index, kind, hf_type)

        # Clear existing content (keep root, remove all children)
        for child in list(hf_root):
            hf_root.remove(child)

        # Add empty paragraph (Word requires at least one)
        etree.SubElement(hf_root, qn("w:p"))
        pkg.mark_xml_dirty(partname)
        return

    # python-docx path (legacy)
    doc = pkg_or_doc
    hf = getattr(doc.sections[section_index], location)
    hf.is_linked_to_previous = False
    hf_el = hf._element
    for child in list(hf_el):
        hf_el.remove(child)
    hf.add_paragraph("")  # Word requires at least one paragraph
