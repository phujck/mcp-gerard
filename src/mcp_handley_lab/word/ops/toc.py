"""Table of Contents operations.

Contains functions for:
- Checking if document has a TOC
- Getting TOC metadata
- Inserting TOC fields
- Marking TOC for update
"""

from __future__ import annotations

import re
from typing import TYPE_CHECKING

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from docx import Document

from mcp_handley_lab.word.ops.core import (
    count_occurrence,
    make_block_id,
    paragraph_kind_and_level,
    resolve_target,
)

# =============================================================================
# TOC Detection and Info
# =============================================================================


def has_toc(doc: Document) -> bool:
    """Check if document has a Table of Contents.

    Searches for:
    1. w:instrText containing "TOC" (complex field)
    2. w:fldSimple[@w:instr] starting with "TOC"
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Check for complex field with TOC
    for instr in doc.element.findall(".//w:instrText", namespaces=ns):
        if instr.text and "TOC" in instr.text.upper():
            return True

    # Check for simple field with TOC
    for fld in doc.element.findall(".//w:fldSimple", namespaces=ns):
        instr = fld.get(qn("w:instr")) or ""
        if instr.strip().upper().startswith("TOC"):
            return True

    return False


def get_toc_info(doc: Document) -> dict:
    """Get TOC metadata if exists.

    Parses heading levels from field switches (e.g., \\o "1-3").
    Returns dict compatible with TOCInfo model.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    result = {
        "exists": False,
        "heading_levels": "1-3",
        "entry_count": 0,
        "block_id": None,
        "has_sdt_wrapper": False,
        "is_dirty": False,
    }

    # Find TOC field instruction
    toc_instr = None
    toc_para = None
    is_dirty = False

    # Check for complex field with TOC
    for instr_el in doc.element.findall(".//w:instrText", namespaces=ns):
        if instr_el.text and "TOC" in instr_el.text.upper():
            toc_instr = instr_el.text
            # Find containing paragraph
            parent = instr_el.getparent()
            while parent is not None:
                if parent.tag == qn("w:p"):
                    toc_para = parent
                    break
                parent = parent.getparent()

            # Check for dirty flag on fldChar begin
            if toc_para is not None:
                for fld_char in toc_para.findall(".//w:fldChar", namespaces=ns):
                    if fld_char.get(qn("w:fldCharType")) == "begin":
                        dirty = fld_char.get(qn("w:dirty"))
                        is_dirty = dirty == "true" or dirty == "1"
                        break
            break

    # Check for simple field with TOC
    if not toc_instr:
        for fld in doc.element.findall(".//w:fldSimple", namespaces=ns):
            instr = fld.get(qn("w:instr")) or ""
            if instr.strip().upper().startswith("TOC"):
                toc_instr = instr
                # Find containing paragraph
                parent = fld.getparent()
                while parent is not None:
                    if parent.tag == qn("w:p"):
                        toc_para = parent
                        break
                    parent = parent.getparent()
                # Check dirty on simple field
                dirty = fld.get(qn("w:dirty"))
                is_dirty = dirty == "true" or dirty == "1"
                break

    if not toc_instr:
        return result

    result["exists"] = True
    result["is_dirty"] = is_dirty

    # Parse heading levels from \o switch
    levels_match = re.search(r'\\o\s*"(\d+-\d+)"', toc_instr)
    if levels_match:
        result["heading_levels"] = levels_match.group(1)

    # Get block ID for the TOC paragraph
    if toc_para is not None:
        p = Paragraph(toc_para, doc)
        kind, level = paragraph_kind_and_level(p)
        text = p.text or ""
        occurrence = count_occurrence(doc, kind, text, toc_para)
        result["block_id"] = make_block_id(kind, text, occurrence)

        # Check for SDT wrapper
        parent = toc_para.getparent()
        if parent is not None and parent.tag == qn("w:sdt"):
            result["has_sdt_wrapper"] = True

    return result


# =============================================================================
# TOC Insertion
# =============================================================================


def insert_toc(
    doc: Document,
    target_id: str,
    position: str = "before",
    heading_levels: str = "1-3",
) -> str:
    """Insert TOC field at position. Returns block ID.

    Field code: TOC \\o "1-3" \\h \\z \\u
    - \\o: heading levels
    - \\h: hyperlinks
    - \\z: hide tab leaders and page numbers in Web view
    - \\u: use applied paragraph outline level

    Sets w:dirty="true" so Word updates on open.
    """
    target = resolve_target(doc, target_id)

    # Create paragraph for TOC
    toc_para = doc.add_paragraph()

    # Build field instruction
    instr = f' TOC \\\\o "{heading_levels}" \\\\h \\\\z \\\\u '

    # Insert field with 5-run structure
    # Run 1: fldChar begin (with dirty flag)
    run1 = toc_para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    run1._r.append(fld_char_begin)

    # Run 2: instrText
    run2 = toc_para.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = instr
    run2._r.append(instr_text)

    # Run 3: fldChar separate
    run3 = toc_para.add_run()
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_char_sep)

    # Run 4: result text (placeholder - Word will replace)
    toc_para.add_run("Update this field to generate Table of Contents")

    # Run 5: fldChar end
    run5 = toc_para.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_char_end)

    # Move paragraph to correct position
    toc_el = toc_para._element
    if target.base_kind == "table":
        target_el = target.base_obj._tbl
    else:
        target_el = target.leaf_el

    if position == "before":
        target_el.addprevious(toc_el)
    else:
        target_el.addnext(toc_el)

    # Generate block ID
    text = toc_para.text or ""
    occurrence = count_occurrence(doc, "paragraph", text, toc_el)
    return make_block_id("paragraph", text, occurrence)


# =============================================================================
# TOC Update
# =============================================================================


def update_toc_field(doc: Document) -> bool:
    """Set dirty flag on TOC field begin marker.

    Sets w:dirty="true" on w:fldChar[@w:fldCharType="begin"] for complex fields,
    or w:dirty="true" on w:fldSimple for simple fields.
    Word recalculates field values when document opens.

    Returns True if TOC was found and updated, False otherwise.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Find complex field with TOC
    for instr_el in doc.element.findall(".//w:instrText", namespaces=ns):
        if instr_el.text and "TOC" in instr_el.text.upper():
            # Find containing paragraph and fldChar begin
            parent = instr_el.getparent()
            while parent is not None:
                if parent.tag == qn("w:p"):
                    # Find fldChar begin in this paragraph
                    for fld_char in parent.findall(".//w:fldChar", namespaces=ns):
                        if fld_char.get(qn("w:fldCharType")) == "begin":
                            fld_char.set(qn("w:dirty"), "true")
                            return True
                    break
                parent = parent.getparent()
            break

    # Find simple field with TOC
    for fld in doc.element.findall(".//w:fldSimple", namespaces=ns):
        instr = fld.get(qn("w:instr")) or ""
        if instr.strip().upper().startswith("TOC"):
            fld.set(qn("w:dirty"), "true")
            return True

    return False
