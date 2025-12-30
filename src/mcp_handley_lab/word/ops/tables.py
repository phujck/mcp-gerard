"""Table operations for Word documents.

Contains functions for table manipulation, formatting, and conversion.
"""

from __future__ import annotations

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tc
from docx.shared import Inches
from docx.table import Table
from lxml import etree

from mcp_handley_lab.word.models import CellInfo, RowInfo, TableLayoutInfo
from mcp_handley_lab.word.ops.core import _insert_at


def table_to_markdown(
    table: Table, max_chars: int = 500, max_rows: int = 20, max_cols: int = 10
) -> tuple[str, int, int]:
    """Convert table to markdown preview with truncation."""
    rows, cols = len(table.rows), len(table.columns)
    r_lim, c_lim = min(rows, max_rows), min(cols, max_cols)

    grid = [
        [
            table.cell(r, c).text.strip().replace("|", "\\|").replace("\n", "<br>")
            for c in range(c_lim)
        ]
        for r in range(r_lim)
    ]

    header = grid[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * len(header)) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in grid[1:])

    if rows > r_lim:
        lines.append(f"... ({rows - r_lim} more rows)")
    if cols > c_lim:
        lines.append(f"... ({cols - c_lim} more cols)")

    md = "\n".join(lines)
    if len(md) > max_chars:
        md = md[:max_chars] + "\n... (truncated)"
    return md, rows, cols


def populate_table(table: Table, data: list[list]) -> None:
    """Populate table cells from 2D list."""
    for r, row in enumerate(data):
        for c, val in enumerate(row):
            table.cell(r, c).text = str(val)


# =============================================================================
# Cell Merge Helpers
# =============================================================================


def _get_vmerge_val_from_tc(tc: CT_Tc) -> str | None:
    """Get vMerge value from a tc element.

    Returns:
        'restart' for merge origin, 'continue' for continuation, None for no merge.
    """
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    v_merge = tc_pr.find(qn("w:vMerge"))
    if v_merge is None:
        return None
    # vMerge with val="restart" is origin, vMerge without val is continue
    val = v_merge.get(qn("w:val"))
    return val if val else "continue"


def _tc_at_grid_col(tr, grid_col: int):
    """Find the tc element at a given grid column, accounting for gridSpan.

    Returns the tc element or None if not found.
    """
    c = 0
    for tc in tr.findall(qn("w:tc")):
        span = tc.grid_span
        if c == grid_col:
            return tc
        c += span
        if c > grid_col:
            # Grid column falls within a span, no tc origin at this position
            return None
    return None


def _calculate_row_span_from_xml(table: Table, start_row: int, col: int) -> int:
    """Calculate vertical span by checking vMerge='continue' in subsequent rows.

    Iterates over actual XML tc elements to detect continuation cells.
    """
    span = 1
    rows = list(table.rows)
    for r in range(start_row + 1, len(rows)):
        row_el = rows[r]._tr
        tc = _tc_at_grid_col(row_el, col)
        if tc is not None:
            vmerge = _get_vmerge_val_from_tc(tc)
            if vmerge == "continue":
                span += 1
            else:
                break
        else:
            break
    return span


def _get_cell_border(tc, side: str) -> str | None:
    """Extract border info from tc element. Returns 'style:size:color' or None."""
    tcPr = tc.find(qn("w:tcPr"))
    borders = tcPr.find(qn("w:tcBorders")) if tcPr is not None else None
    border_el = borders.find(qn(f"w:{side}")) if borders is not None else None
    if border_el is None:
        return None
    style = border_el.get(qn("w:val")) or "single"
    sz = border_el.get(qn("w:sz")) or "4"
    color = border_el.get(qn("w:color")) or "auto"
    return f"{style}:{sz}:{color}"


def _get_cell_shading(tc) -> str | None:
    """Extract fill color from tc element. Returns hex color or None."""
    tcPr = tc.find(qn("w:tcPr"))
    shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    return fill.upper() if fill and fill.lower() != "auto" else None


# =============================================================================
# Table Reading Functions
# =============================================================================


def build_table_cells(table: Table, table_id: str = "") -> list[CellInfo]:
    """Build list of CellInfo with merge information.

    Detects:
    - Horizontal merges via grid_span property
    - Vertical merges via vMerge XML attribute

    Iterates over actual XML elements (not table.cell()) to correctly
    detect continuation cells in vertical merges.

    Args:
        table: The Table object
        table_id: Base ID of the table (for hierarchical IDs)

    Returns:
        List of CellInfo with merge info for all grid positions
    """
    valign_map = {
        WD_CELL_VERTICAL_ALIGNMENT.TOP: "top",
        WD_CELL_VERTICAL_ALIGNMENT.CENTER: "center",
        WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: "bottom",
    }

    result = []
    rows = list(table.rows)

    for r, row in enumerate(rows):
        row_el = row._tr
        tc_elements = row_el.findall(qn("w:tc"))
        c = 0
        for tc in tc_elements:
            vmerge = _get_vmerge_val_from_tc(tc)
            grid_span = tc.grid_span

            if vmerge == "continue":
                # Vertical continuation cell
                result.append(
                    CellInfo(
                        row=r,
                        col=c,
                        text="",
                        hierarchical_id=f"{table_id}#r{r}c{c}" if table_id else "",
                        is_merge_origin=False,
                        grid_span=1,
                        row_span=1,
                    )
                )
                # Add horizontal continuation entries for wide continuation cells
                for span_c in range(1, grid_span):
                    result.append(
                        CellInfo(
                            row=r,
                            col=c + span_c,
                            text="",
                            hierarchical_id=(
                                f"{table_id}#r{r}c{c + span_c}" if table_id else ""
                            ),
                            is_merge_origin=False,
                            grid_span=1,
                            row_span=1,
                        )
                    )
                c += grid_span
            else:
                # Origin cell (vmerge='restart' or None) or normal cell
                row_span = (
                    _calculate_row_span_from_xml(table, r, c)
                    if vmerge == "restart"
                    else 1
                )
                # Get text and properties from the cell
                cell = table.cell(r, c)
                width_inches = cell.width.inches if cell.width else None
                valign = valign_map.get(cell.vertical_alignment)
                # Extract border and shading from tc element
                border_top = _get_cell_border(tc, "top")
                border_bottom = _get_cell_border(tc, "bottom")
                border_left = _get_cell_border(tc, "left")
                border_right = _get_cell_border(tc, "right")
                fill_color = _get_cell_shading(tc)
                result.append(
                    CellInfo(
                        row=r,
                        col=c,
                        text=cell.text or "",
                        hierarchical_id=f"{table_id}#r{r}c{c}" if table_id else "",
                        is_merge_origin=True,
                        grid_span=grid_span,
                        row_span=row_span,
                        width_inches=width_inches,
                        vertical_alignment=valign,
                        border_top=border_top,
                        border_bottom=border_bottom,
                        border_left=border_left,
                        border_right=border_right,
                        fill_color=fill_color,
                    )
                )
                # Add continuation entries for horizontal span
                for span_c in range(1, grid_span):
                    result.append(
                        CellInfo(
                            row=r,
                            col=c + span_c,
                            text="",
                            hierarchical_id=(
                                f"{table_id}#r{r}c{c + span_c}" if table_id else ""
                            ),
                            is_merge_origin=False,
                            grid_span=1,
                            row_span=1,
                        )
                    )
                c += grid_span
    return result


def build_table_layout(table: Table, table_id: str) -> TableLayoutInfo:
    """Build table layout info including row heights and alignment."""
    table_align_map = {
        WD_TABLE_ALIGNMENT.LEFT: "left",
        WD_TABLE_ALIGNMENT.CENTER: "center",
        WD_TABLE_ALIGNMENT.RIGHT: "right",
    }
    row_height_rule_map = {
        WD_ROW_HEIGHT_RULE.AUTO: "auto",
        WD_ROW_HEIGHT_RULE.AT_LEAST: "at_least",
        WD_ROW_HEIGHT_RULE.EXACTLY: "exactly",
    }

    rows = []
    for i, row in enumerate(table.rows):
        # Check for header row marker (w:tblHeader in w:trPr)
        tr_el = row._tr
        trPr = tr_el.find(qn("w:trPr"))
        is_header = trPr is not None and trPr.find(qn("w:tblHeader")) is not None
        rows.append(
            RowInfo(
                index=i,
                height_inches=row.height.inches if row.height else None,
                height_rule=row_height_rule_map.get(row.height_rule),
                is_header=is_header,
            )
        )

    return TableLayoutInfo(
        table_id=table_id,
        alignment=table_align_map.get(table.alignment),
        autofit=table.autofit,
        rows=rows,
    )


# =============================================================================
# Table Modification Functions
# =============================================================================


def insert_table_relative(
    doc: Document,
    target_el,
    table_data: list[list[str]],
    position: str,
    style_name: str = "Table Grid",
) -> Table:
    """Insert table before/after target element."""
    rows, cols = len(table_data), max((len(r) for r in table_data), default=1)
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = style_name
    populate_table(tbl, table_data)
    _insert_at(target_el, tbl._tbl, position)
    return tbl


def replace_table(doc: Document, old_tbl: Table, table_data: list[list[str]]) -> Table:
    """Replace table with new data."""
    rows, cols = len(table_data), max((len(r) for r in table_data), default=1)
    new_tbl = doc.add_table(rows=rows, cols=cols)
    populate_table(new_tbl, table_data)
    old_el = old_tbl._tbl
    old_el.addprevious(new_tbl._tbl)
    old_el.getparent().remove(old_el)
    return new_tbl


def merge_cells(
    table: Table, start_row: int, start_col: int, end_row: int, end_col: int
) -> None:
    """Merge a rectangular region of cells. All indices are 0-based."""
    rows, cols = len(table.rows), len(table.columns)

    # Validate bounds
    if not (0 <= start_row < rows and 0 <= end_row < rows):
        raise ValueError(
            f"Row indices must be 0-{rows - 1}, got start_row={start_row}, end_row={end_row}"
        )
    if not (0 <= start_col < cols and 0 <= end_col < cols):
        raise ValueError(
            f"Column indices must be 0-{cols - 1}, got start_col={start_col}, end_col={end_col}"
        )

    # Validate ordering
    if start_row > end_row or start_col > end_col:
        raise ValueError(
            f"Start must be <= end: ({start_row},{start_col}) to ({end_row},{end_col})"
        )

    start_cell = table.cell(start_row, start_col)
    end_cell = table.cell(end_row, end_col)
    start_cell.merge(end_cell)


def replace_table_cell(table: Table, row: int, col: int, text: str) -> None:
    """Replace text in a table cell. Row/col are 0-based."""
    table.cell(row, col).text = text


def add_table_row(table: Table, data: list[str] | None = None) -> int:
    """Add row to table. Returns new row index (0-based)."""
    row = table.add_row()
    if data:
        for i, text in enumerate(data[: len(table.columns)]):
            row.cells[i].text = text
    return len(table.rows) - 1


def add_table_column(
    table: Table, width_inches: float = 1.0, data: list[str] | None = None
) -> int:
    """Add column to table. Width required by python-docx API. Returns new col index."""
    table.add_column(Inches(width_inches))
    col_idx = len(table.columns) - 1
    if data:
        for i, text in enumerate(data[: len(table.rows)]):
            table.cell(i, col_idx).text = text
    return col_idx


def delete_table_row(table: Table, row_index: int) -> None:
    """Delete row from table (0-based index)."""
    row = table.rows[row_index]
    row._element.getparent().remove(row._element)


def delete_table_column(table: Table, col_index: int) -> None:
    """Delete column from table (0-based index). Removes grid definition and cells."""
    # 1. Remove the grid column definition (required for valid Word XML)
    tbl_grid = table._tbl.tblGrid
    if col_index < len(tbl_grid.gridCol_lst):
        grid_col = tbl_grid.gridCol_lst[col_index]
        grid_col.getparent().remove(grid_col)

    # 2. Remove the cell from every row
    for row in table.rows:
        if col_index < len(row.cells):
            cell = row.cells[col_index]
            cell._element.getparent().remove(cell._element)


# =============================================================================
# Table Layout Functions
# =============================================================================


def set_table_alignment(table: Table, alignment: str) -> None:
    """Set table horizontal alignment. Valid: left, center, right."""
    alignment_map = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }
    alignment_lower = alignment.lower()
    if alignment_lower not in alignment_map:
        raise ValueError(
            f"Invalid alignment '{alignment}'. Valid: {list(alignment_map.keys())}"
        )
    table.alignment = alignment_map[alignment_lower]


def set_row_height(
    table: Table, row_index: int, height_inches: float, rule: str = "at_least"
) -> None:
    """Set row height. Rule: auto, at_least, exactly. Default 'at_least' prevents clipping."""
    rule_map = {
        "auto": WD_ROW_HEIGHT_RULE.AUTO,
        "at_least": WD_ROW_HEIGHT_RULE.AT_LEAST,
        "exactly": WD_ROW_HEIGHT_RULE.EXACTLY,
    }
    rule_val = rule.lower()
    if rule_val not in rule_map:
        raise ValueError(f"Invalid rule '{rule}'. Valid: {list(rule_map.keys())}")
    row = table.rows[row_index]  # Let IndexError propagate
    row.height = None if rule_val == "auto" else Inches(height_inches)
    row.height_rule = rule_map[rule_val]


def set_table_fixed_layout(table: Table, column_widths: list[float]) -> None:
    """Set table to fixed layout with explicit column widths (inches)."""
    table.autofit = False
    for i, width in enumerate(column_widths):
        if i < len(table.columns):
            table.columns[i].width = Inches(width)


def set_cell_width(table: Table, row: int, col: int, width_inches: float) -> None:
    """Set cell width."""
    table.cell(row, col).width = Inches(width_inches)


def set_cell_vertical_alignment(
    table: Table, row: int, col: int, alignment: str
) -> None:
    """Set cell vertical alignment. Valid: top, center, bottom."""
    valign_map = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }
    alignment_lower = alignment.lower()
    if alignment_lower not in valign_map:
        raise ValueError(
            f"Invalid alignment '{alignment}'. Valid: {list(valign_map.keys())}"
        )
    table.cell(row, col).vertical_alignment = valign_map[alignment_lower]


def set_cell_borders(
    table: Table,
    row: int,
    col: int,
    top: str | None = None,
    bottom: str | None = None,
    left: str | None = None,
    right: str | None = None,
) -> None:
    """Set cell borders. Format: 'style:size:color' e.g. 'single:24:000000'.

    Args:
        table: The table object
        row: 0-based row index
        col: 0-based column index
        top, bottom, left, right: Border specs in 'style:size:color' format
            - style: single, double, dotted, dashed, etc.
            - size: in eighths of a point (24 = 3pt)
            - color: hex color (e.g., '000000' for black)
    """
    cell = table.cell(row, col)
    tc = cell._tc

    # Get or create tcPr
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.Element(qn("w:tcPr"))
        tc.insert(0, tcPr)

    # Get or create tcBorders
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))

    def set_border(side: str, spec: str) -> None:
        style, sz, color = spec.split(":")
        border_el = tcBorders.find(qn(f"w:{side}"))
        if border_el is None:
            border_el = etree.SubElement(tcBorders, qn(f"w:{side}"))
        border_el.set(qn("w:val"), style)
        border_el.set(qn("w:sz"), sz)
        border_el.set(qn("w:color"), color)

    if top:
        set_border("top", top)
    if bottom:
        set_border("bottom", bottom)
    if left:
        set_border("left", left)
    if right:
        set_border("right", right)


def set_cell_shading(table: Table, row: int, col: int, fill_color: str) -> None:
    """Set cell background color.

    Args:
        table: The table object
        row: 0-based row index
        col: 0-based column index
        fill_color: Hex color (e.g., 'FF0000' for red)
    """
    cell = table.cell(row, col)
    tc = cell._tc

    # Get or create tcPr
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.Element(qn("w:tcPr"))
        tc.insert(0, tcPr)

    # Get or create shd
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = etree.SubElement(tcPr, qn("w:shd"))

    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_color.upper())


def set_header_row(table: Table, row_index: int, is_header: bool = True) -> None:
    """Mark row as header (repeats on each page in multi-page tables).

    Args:
        table: The table object
        row_index: 0-based row index
        is_header: True to mark as header, False to unmark
    """
    row = table.rows[row_index]
    tr = row._tr

    # Get or create trPr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = etree.Element(qn("w:trPr"))
        tr.insert(0, trPr)

    # Find existing tblHeader
    tblHeader = trPr.find(qn("w:tblHeader"))

    if is_header:
        if tblHeader is None:
            etree.SubElement(trPr, qn("w:tblHeader"))
    else:
        if tblHeader is not None:
            trPr.remove(tblHeader)


def get_header_rows(table: Table) -> list[int]:
    """Get indices of rows marked as headers.

    Returns:
        List of 0-based row indices that are marked as headers
    """
    result = []
    for i, row in enumerate(table.rows):
        tr = row._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None and trPr.find(qn("w:tblHeader")) is not None:
            result.append(i)
    return result
