"""Image and text box operations.

Contains functions for:
- Embedded/inline image extraction and metadata
- Floating (anchored) image insertion with positioning
- Image resolution and deletion
- Text box discovery (DrawingML + VML)
- Text box content reading and editing
"""

from __future__ import annotations

import random
from typing import TYPE_CHECKING

from lxml import etree
from lxml.etree import ElementBase as _LxmlElementBase

from mcp_handley_lab.word.opc.constants import NSMAP as OXML_NSMAP
from mcp_handley_lab.word.opc.constants import qn

if TYPE_CHECKING:
    from docx import Document

from mcp_handley_lab.word.models import ImageInfo
from mcp_handley_lab.word.ops.core import (
    _EMU_PER_INCH,
    _IMAGE_ID_RE,
    _iter_all_paragraphs,
    _iter_all_runs_in_paragraph,
    content_hash,
    make_block_id,
    mark_dirty,
    paragraph_kind_and_level,
    resolve_target,
    table_content_for_hash,
)

# =============================================================================
# Constants
# =============================================================================

# Wrap type mapping: API value <-> XML element name (without wp: prefix)
_WRAP_API_TO_XML = {
    "square": "wrapSquare",
    "tight": "wrapTight",
    "through": "wrapThrough",
    "top_and_bottom": "wrapTopAndBottom",
    "none": "wrapNone",
}
_WRAP_XML_TO_API = {v: k for k, v in _WRAP_API_TO_XML.items()}

# Namespaces for text box discovery (DrawingML + VML)
_TEXTBOX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "v": "urn:schemas-microsoft-com:vml",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
}


# =============================================================================
# XPath Helpers (duck-typed for lxml compatibility)
# =============================================================================


def _img_xpath(element, expr: str, ns: dict | None = None) -> list:
    """XPath using ElementBase.xpath for namespace compatibility.

    Uses lxml ElementBase.xpath to ensure namespaces parameter works
    correctly with both raw lxml elements and python-docx BaseOxmlElement.
    """
    return _LxmlElementBase.xpath(element, expr, namespaces=ns or OXML_NSMAP)


def _img_find(element, expr: str, ns: dict | None = None):
    """Find first matching element or None (like lxml.find but using xpath helper)."""
    results = _img_xpath(element, expr, ns)
    return results[0] if results else None


# =============================================================================
# Image Hash and Extraction Helpers
# =============================================================================


def get_embedded_image_hash(doc: Document, blip) -> str | None:
    """Get SHA1 hash for embedded image. Returns None for linked images."""
    # Support both python-docx elements (with .embed attribute) and raw lxml
    try:
        rel_id = blip.embed  # python-docx accessor
    except AttributeError:
        # Raw lxml element - use qualified namespace
        r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        rel_id = blip.get(f"{{{r_ns}}}embed")
    if not rel_id:
        return None  # Linked image - can't hash external resource
    image_part = doc.part.related_parts[rel_id]
    return image_part.image.sha1[:8]


def _extract_anchor_position(anchor) -> dict:
    """Extract positioning info from wp:anchor element."""
    result = {"position_type": "anchor", "behind_doc": False}

    # behindDoc attribute (no namespace prefix)
    behind = anchor.get("behindDoc")
    result["behind_doc"] = behind == "1" if behind else False

    # Horizontal position
    pos_h = _img_find(anchor, "./wp:positionH")
    if pos_h is not None:
        result["relative_from_h"] = pos_h.get("relativeFrom")
        offset_h = _img_find(pos_h, "./wp:posOffset")
        if offset_h is not None and offset_h.text:
            result["position_h"] = int(offset_h.text) / _EMU_PER_INCH

    # Vertical position
    pos_v = _img_find(anchor, "./wp:positionV")
    if pos_v is not None:
        result["relative_from_v"] = pos_v.get("relativeFrom")
        offset_v = _img_find(pos_v, "./wp:posOffset")
        if offset_v is not None and offset_v.text:
            result["position_v"] = int(offset_v.text) / _EMU_PER_INCH

    # Wrap type (use mapping for consistent API values)
    for xml_name, api_value in _WRAP_XML_TO_API.items():
        wrap_el = _img_find(anchor, f"./wp:{xml_name}")
        if wrap_el is not None:
            result["wrap_type"] = api_value
            break

    return result


def _extract_images_from_run(
    doc: Document,
    run,
    run_idx: int,
    block_id: str,
    image_hash_counts: dict[str, int],
    images: list[ImageInfo],
) -> None:
    """Extract images from a single run (both inline and anchored)."""
    image_idx_in_run = 0
    for drawing in _img_xpath(run._element, ".//w:drawing"):
        # Check for both inline and anchor images
        inline = _img_find(drawing, ".//wp:inline")
        anchor = _img_find(drawing, ".//wp:anchor")

        container = inline if inline is not None else anchor
        if container is None:
            continue

        # Guard access to pic element (skip charts/smartart)
        # Try python-docx attribute accessors first, fall back to XPath for raw lxml
        try:
            blip = container.graphic.graphicData.pic.blipFill.blip
        except AttributeError:
            # Raw lxml element - use XPath (lowercase 'ml' in namespace)
            blip_ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
            blip_list = _img_xpath(container, ".//a:blip", blip_ns)
            if not blip_list:
                continue
            blip = blip_list[0]

        h = get_embedded_image_hash(doc, blip)
        if h is None:
            continue  # Skip linked images

        # Track image occurrence globally
        img_occurrence = image_hash_counts.get(h, 0)
        image_hash_counts[h] = img_occurrence + 1

        # Get metadata via XML (avoid InlineShape construction issues)
        # Handle both python-docx elements (with .embed) and raw lxml elements
        try:
            rel_id = blip.embed  # python-docx accessor
        except AttributeError:
            r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            rel_id = blip.get(f"{{{r_ns}}}embed")
        image_part = doc.part.related_parts[rel_id]
        # Handle extent for both python-docx elements and raw lxml elements
        try:
            extent = container.extent
            width_emu = extent.cx if extent is not None else 0
            height_emu = extent.cy if extent is not None else 0
        except AttributeError:
            # Raw lxml element - find wp:extent element directly
            extent_el = _img_find(container, "./wp:extent")
            if extent_el is not None:
                width_emu = int(extent_el.get("cx", 0))
                height_emu = int(extent_el.get("cy", 0))
            else:
                width_emu, height_emu = 0, 0

        # Build base image info
        info_kwargs = {
            "id": f"image_{h}_{img_occurrence}",
            "width_inches": width_emu / _EMU_PER_INCH,
            "height_inches": height_emu / _EMU_PER_INCH,
            "content_type": image_part.content_type,
            "block_id": block_id,
            "run_index": run_idx,
            "image_index_in_run": image_idx_in_run,
            "filename": image_part.image.filename or "",
        }

        # Add anchor-specific positioning if floating image
        if anchor is not None:
            info_kwargs.update(_extract_anchor_position(anchor))

        images.append(ImageInfo(**info_kwargs))
        image_idx_in_run += 1


def _extract_images_from_paragraph(
    doc: Document,
    para,
    block_id: str,
    image_hash_counts: dict[str, int],
) -> list[ImageInfo]:
    """Extract images from a paragraph, updating occurrence counts.

    Uses iter_inner_content() indexing to match build_runs() indexing.
    """

    from docx.text.hyperlink import Hyperlink

    images: list[ImageInfo] = []
    run_idx = 0
    for item in para.iter_inner_content():
        if isinstance(item, Hyperlink):
            for run in item.runs:
                _extract_images_from_run(
                    doc, run, run_idx, block_id, image_hash_counts, images
                )
                run_idx += 1
        else:  # Run
            _extract_images_from_run(
                doc, item, run_idx, block_id, image_hash_counts, images
            )
            run_idx += 1
    return images


# =============================================================================
# Image Building and Resolution (Pure OOXML)
# =============================================================================


def _get_image_hash_ooxml(pkg, rel_id: str) -> str | None:
    """Get SHA1 hash for embedded image (pure OOXML). Returns None if not found."""
    import hashlib

    doc_rels = pkg.get_rels("/word/document.xml")
    rel = doc_rels.get(rel_id)
    if not rel or rel.is_external:
        return None  # Linked image - can't hash

    # Resolve relative path
    target = rel.target
    if not target.startswith("/"):
        target = f"/word/{target}"

    if not pkg.has_part(target):
        return None

    image_bytes = pkg.get_bytes(target)
    return hashlib.sha1(image_bytes).hexdigest()[:8]


def _get_image_content_type_ooxml(pkg, rel_id: str) -> str:
    """Get content type for image (pure OOXML)."""
    doc_rels = pkg.get_rels("/word/document.xml")
    rel = doc_rels.get(rel_id)
    if not rel:
        return ""

    target = rel.target
    if not target.startswith("/"):
        target = f"/word/{target}"

    try:
        return pkg._content_types[target]
    except KeyError:
        return ""


def _get_image_filename_ooxml(rel_target: str) -> str:
    """Get filename from relationship target."""
    import posixpath

    return posixpath.basename(rel_target)


def _extract_images_from_run_ooxml(
    pkg,
    run_el,
    run_idx: int,
    block_id: str,
    image_hash_counts: dict[str, int],
) -> list[ImageInfo]:
    """Extract images from a single run element (pure OOXML)."""
    images = []
    image_idx_in_run = 0

    for drawing in _img_xpath(run_el, ".//w:drawing"):
        inline = _img_find(drawing, ".//wp:inline")
        anchor = _img_find(drawing, ".//wp:anchor")

        container = inline if inline is not None else anchor
        if container is None:
            continue

        # Find blip element (skip charts/smartart without blip)
        blip_ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
        blip_list = _img_xpath(container, ".//a:blip", blip_ns)
        if not blip_list:
            continue
        blip = blip_list[0]

        # Get embed relationship ID
        r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        rel_id = blip.get(f"{{{r_ns}}}embed")
        if not rel_id:
            continue  # Linked image

        h = _get_image_hash_ooxml(pkg, rel_id)
        if h is None:
            continue

        # Track image occurrence
        img_occurrence = image_hash_counts.get(h, 0)
        image_hash_counts[h] = img_occurrence + 1

        # Get dimensions
        extent_el = _img_find(container, "./wp:extent")
        if extent_el is not None:
            width_emu = int(extent_el.get("cx", 0))
            height_emu = int(extent_el.get("cy", 0))
        else:
            width_emu, height_emu = 0, 0

        # Get content type and filename
        content_type = _get_image_content_type_ooxml(pkg, rel_id)
        doc_rels = pkg.get_rels("/word/document.xml")
        rel = doc_rels.get(rel_id)
        filename = _get_image_filename_ooxml(rel.target) if rel else ""

        info_kwargs = {
            "id": f"image_{h}_{img_occurrence}",
            "width_inches": width_emu / _EMU_PER_INCH,
            "height_inches": height_emu / _EMU_PER_INCH,
            "content_type": content_type,
            "block_id": block_id,
            "run_index": run_idx,
            "image_index_in_run": image_idx_in_run,
            "filename": filename,
        }

        if anchor is not None:
            info_kwargs.update(_extract_anchor_position(anchor))

        images.append(ImageInfo(**info_kwargs))
        image_idx_in_run += 1

    return images


def _extract_images_from_paragraph_ooxml(
    pkg,
    p_el,
    block_id: str,
    image_hash_counts: dict[str, int],
) -> list[ImageInfo]:
    """Extract images from a paragraph element (pure OOXML)."""
    images = []
    run_idx = 0

    for child in p_el:
        tag = child.tag
        if tag == qn("w:r"):
            images.extend(
                _extract_images_from_run_ooxml(
                    pkg, child, run_idx, block_id, image_hash_counts
                )
            )
            run_idx += 1
        elif tag == qn("w:hyperlink"):
            # Process runs inside hyperlink
            for run_el in child.iter(qn("w:r")):
                images.extend(
                    _extract_images_from_run_ooxml(
                        pkg, run_el, run_idx, block_id, image_hash_counts
                    )
                )
                run_idx += 1

    return images


def _extract_paragraph_text_ooxml(p_el) -> str:
    """Extract text from paragraph element."""
    text_parts = []
    for t in p_el.iter(qn("w:t")):
        if t.text:
            text_parts.append(t.text)
    return "".join(text_parts)


def _build_images_ooxml(pkg) -> list[ImageInfo]:
    """Build images list from pure OOXML."""
    images = []
    block_hash_counts: dict[str, int] = {}
    image_hash_counts: dict[str, int] = {}

    body = pkg.body

    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            block_type, _ = paragraph_kind_and_level(child)
            text = _extract_paragraph_text_ooxml(child)
            block_hash_key = f"{block_type}_{content_hash(text)}"
            block_occurrence = block_hash_counts.get(block_hash_key, 0)
            block_id = make_block_id(block_type, text, block_occurrence)
            block_hash_counts[block_hash_key] = block_occurrence + 1

            images.extend(
                _extract_images_from_paragraph_ooxml(
                    pkg, child, block_id, image_hash_counts
                )
            )

        elif tag == qn("w:tbl"):
            table_content = table_content_for_hash(child)
            block_hash_key = f"table_{content_hash(table_content)}"
            block_occurrence = block_hash_counts.get(block_hash_key, 0)
            table_block_id = make_block_id("table", table_content, block_occurrence)
            block_hash_counts[block_hash_key] = block_occurrence + 1

            # Search all cells for images
            for r_idx, tr in enumerate(child.findall(qn("w:tr"))):
                for c_idx, tc in enumerate(tr.findall(qn("w:tc"))):
                    for p_idx, p in enumerate(tc.findall(qn("w:p"))):
                        hier_block_id = f"{table_block_id}#r{r_idx}c{c_idx}/p{p_idx}"
                        images.extend(
                            _extract_images_from_paragraph_ooxml(
                                pkg, p, hier_block_id, image_hash_counts
                            )
                        )

    return images


# =============================================================================
# Image Building and Resolution
# =============================================================================


def build_images(pkg_or_doc) -> list[ImageInfo]:
    """Build list of ImageInfo from document body, including tables.

    Args:
        pkg_or_doc: WordPackage or python-docx Document (duck-typed)
    """
    # Check if it's a WordPackage (pure OOXML path)
    if hasattr(pkg_or_doc, "document_xml"):
        return _build_images_ooxml(pkg_or_doc)

    # python-docx Document path (legacy)
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph as DocxParagraph

    doc = pkg_or_doc
    images = []
    block_hash_counts: dict[str, int] = {}  # For block_id computation
    image_hash_counts: dict[str, int] = {}  # For image occurrence

    # Iterate body blocks directly using python-docx
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            para = DocxParagraph(child, doc)
            block_type, _ = paragraph_kind_and_level(child)
            text = para.text or ""
            block_hash_key = f"{block_type}_{content_hash(text)}"
            block_occurrence = block_hash_counts.get(block_hash_key, 0)
            block_id = make_block_id(block_type, text, block_occurrence)
            block_hash_counts[block_hash_key] = block_occurrence + 1

            images.extend(
                _extract_images_from_paragraph(doc, para, block_id, image_hash_counts)
            )

        elif isinstance(child, CT_Tbl):
            tbl = Table(child, doc)
            # Use element for hash computation
            table_content = table_content_for_hash(child)
            block_hash_key = f"table_{content_hash(table_content)}"
            block_occurrence = block_hash_counts.get(block_hash_key, 0)
            table_block_id = make_block_id("table", table_content, block_occurrence)
            block_hash_counts[block_hash_key] = block_occurrence + 1

            # Search all cells for images with hierarchical block_id
            rows, cols = len(tbl.rows), len(tbl.columns)
            for r in range(rows):
                for c in range(cols):
                    cell = tbl.cell(r, c)
                    for p_idx, para in enumerate(cell.paragraphs):
                        hier_block_id = f"{table_block_id}#r{r}c{c}/p{p_idx}"
                        images.extend(
                            _extract_images_from_paragraph(
                                doc, para, hier_block_id, image_hash_counts
                            )
                        )

    return images


def _find_image_in_paragraph(doc: Document, para, target_hash: str):
    """Yield each wp:inline or wp:anchor element in paragraph matching target_hash.

    Uses iter_inner_content() traversal to match build_images() indexing.
    """
    # Pass element to _iter_all_runs_in_paragraph, handle both wrapper and element
    p_el = para._element if hasattr(para, "_element") else para
    for run_el in _iter_all_runs_in_paragraph(p_el):
        for drawing in _img_xpath(run_el, ".//w:drawing"):
            inline = _img_find(drawing, ".//wp:inline")
            anchor = _img_find(drawing, ".//wp:anchor")
            container = inline if inline is not None else anchor
            if container is None:
                continue
            # Try python-docx attribute accessors first, fall back to XPath for raw lxml
            try:
                blip = container.graphic.graphicData.pic.blipFill.blip
            except AttributeError:
                # Raw lxml element - use XPath (lowercase 'ml' in namespace)
                blip_ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
                blip_list = _img_xpath(container, ".//a:blip", blip_ns)
                if not blip_list:
                    continue
                blip = blip_list[0]
            h = get_embedded_image_hash(doc, blip)
            if h == target_hash:
                yield container


def resolve_image(doc: Document, image_id: str) -> tuple:
    """Find embedded image by content-addressable ID. Returns (inline_el, para_el)."""
    target_hash, occurrence_str = _IMAGE_ID_RE.match(image_id).groups()
    target_occurrence = int(occurrence_str)

    occurrence_count = 0
    for p_el in _iter_all_paragraphs(doc):  # Now yields only elements
        for inline in _find_image_in_paragraph(doc, p_el, target_hash):
            if occurrence_count == target_occurrence:
                return inline, p_el
            occurrence_count += 1

    raise ValueError(f"Image not found: {image_id}")


def count_image_occurrence(doc: Document, target_hash: str, target_para_el) -> int:
    """Count embedded images with same hash before target paragraph."""
    occurrence = 0
    for p_el in _iter_all_paragraphs(doc):  # Now yields only elements
        for _inline in _find_image_in_paragraph(doc, p_el, target_hash):
            if p_el is target_para_el:
                return occurrence
            occurrence += 1
    raise ValueError("Target paragraph not found")


# =============================================================================
# Image Insertion and Deletion
# =============================================================================


def insert_image(
    doc: Document,
    image_path: str,
    target_id: str,
    position: str,
    width_inches: float = 0,
    height_inches: float = 0,
) -> str:
    """Insert image at target location.

    Supports hierarchical target IDs:
    - table_abc_0#r0c1/p0 -> Insert into paragraph in cell
    - table_abc_0#r0c1 -> Insert into first paragraph of cell
    - table_abc_0 -> Insert before/after table
    - paragraph_abc_0 -> Insert before/after paragraph
    """

    from docx.shared import Inches

    target = resolve_target(doc, target_id)

    # Get image hash first
    _, image = doc.part.get_or_add_image(image_path)
    h = image.sha1[:8]

    width = Inches(width_inches) if width_inches else None
    height = Inches(height_inches) if height_inches else None

    if target.leaf_kind == "paragraph":
        # Insert into this paragraph
        para = target.leaf_obj
        run = para.add_run()
        run.add_picture(image_path, width, height)
        mark_dirty(doc)
        occurrence = count_image_occurrence(doc, h, para._element)
        return f"image_{h}_{occurrence}"

    if target.leaf_kind == "cell":
        # Use first paragraph of cell (create if needed)
        cell = target.leaf_obj
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = para.add_run()
        run.add_picture(image_path, width, height)
        mark_dirty(doc)
        occurrence = count_image_occurrence(doc, h, para._element)
        return f"image_{h}_{occurrence}"

    # Target is table or paragraph at base level -> insert before/after
    new_para = doc.add_paragraph()
    run = new_para.add_run()
    run.add_picture(image_path, width, height)

    if position == "before":
        target.leaf_el.addprevious(new_para._element)
    else:
        target.leaf_el.addnext(new_para._element)

    mark_dirty(doc)
    occurrence = count_image_occurrence(doc, h, new_para._element)
    return f"image_{h}_{occurrence}"


def delete_image(doc: Document, image_id: str) -> None:
    """Delete an image. Removes containing paragraph if only whitespace remains."""
    from docx.text.paragraph import Paragraph

    inline_el, para_el = resolve_image(doc, image_id)

    # Remove the drawing element (parent of inline)
    drawing_el = inline_el.getparent()
    drawing_el.getparent().remove(drawing_el)

    # Check if paragraph has any remaining content (text, drawings, fields, etc.)
    para = Paragraph(para_el, doc)
    has_content = bool(para.text.strip())
    if not has_content:
        for run_el in _iter_all_runs_in_paragraph(
            para_el
        ):  # Pass element, returns elements
            # Check for drawings, fields, or other significant content
            if _img_xpath(run_el, ".//w:drawing") or _img_xpath(run_el, ".//w:fldChar"):
                has_content = True
                break

    if not has_content:
        para_el.getparent().remove(para_el)

    # Mark document.xml as modified for WordPackage
    mark_dirty(doc)


def insert_floating_image(
    doc: Document,
    image_path: str,
    target_id: str,
    position_h: float,
    position_v: float,
    relative_h: str = "column",
    relative_v: str = "paragraph",
    wrap_type: str = "square",
    width_inches: float = 0,
    height_inches: float = 0,
    behind_doc: bool = False,
) -> str:
    """Insert floating (anchored) image at target location.

    Args:
        doc: Document object
        image_path: Path to image file
        target_id: Block ID for anchor paragraph
        position_h: Horizontal position in inches
        position_v: Vertical position in inches
        relative_h: Horizontal reference ("column", "page", "margin", "character")
        relative_v: Vertical reference ("paragraph", "page", "margin", "line")
        wrap_type: Text wrap ("square", "tight", "through", "top_and_bottom", "none")
        width_inches: Image width (0 = auto from image)
        height_inches: Image height (0 = auto from image)
        behind_doc: True to place image behind text

    Returns:
        Image ID (image_{hash}_{occurrence})
    """
    target = resolve_target(doc, target_id)

    # Get/add image to package and get relationship ID
    rId, image = doc.part.get_or_add_image(image_path)
    h = image.sha1[:8]

    # Calculate default dimensions in EMUs from pixels and DPI
    # EMU = (pixels * 914400) / DPI
    default_cx = int(image.px_width * _EMU_PER_INCH / image.horz_dpi)
    default_cy = int(image.px_height * _EMU_PER_INCH / image.vert_dpi)

    # Calculate dimensions in EMUs
    if width_inches and height_inches:
        cx = int(width_inches * _EMU_PER_INCH)
        cy = int(height_inches * _EMU_PER_INCH)
    elif width_inches:
        cx = int(width_inches * _EMU_PER_INCH)
        cy = int(cx * default_cy / default_cx)
    elif height_inches:
        cy = int(height_inches * _EMU_PER_INCH)
        cx = int(cy * default_cx / default_cy)
    else:
        cx, cy = default_cx, default_cy

    # Convert position to EMUs
    offset_h = int(position_h * _EMU_PER_INCH)
    offset_v = int(position_v * _EMU_PER_INCH)

    # Generate unique IDs for drawing elements
    doc_pr_id = random.randint(100000, 999999)

    # Map wrap_type API value to XML element name
    wrap_xml_name = _WRAP_API_TO_XML.get(wrap_type, "wrapSquare")
    wrap_element = (
        f"<wp:{wrap_xml_name}/>"
        if wrap_type == "none"
        else f'<wp:{wrap_xml_name} wrapText="bothSides"/>'
    )

    # Build anchor XML using OOXML structure
    # Note: namespace URIs use lowercase 'ml' to match python-docx's oxml_nsmap
    anchor_xml = f"""
    <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
               distT="0" distB="0" distL="114300" distR="114300"
               simplePos="0" relativeHeight="251658240" behindDoc="{1 if behind_doc else 0}"
               locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="{relative_h}">
            <wp:posOffset>{offset_h}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="{relative_v}">
            <wp:posOffset>{offset_v}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        {wrap_element}
        <wp:docPr id="{doc_pr_id}" name="Picture {doc_pr_id}"/>
        <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic>
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic>
                    <pic:nvPicPr>
                        <pic:cNvPr id="{doc_pr_id}" name="Picture {doc_pr_id}"/>
                        <pic:cNvPicPr/>
                    </pic:nvPicPr>
                    <pic:blipFill>
                        <a:blip r:embed="{rId}"/>
                        <a:stretch>
                            <a:fillRect/>
                        </a:stretch>
                    </pic:blipFill>
                    <pic:spPr>
                        <a:xfrm>
                            <a:off x="0" y="0"/>
                            <a:ext cx="{cx}" cy="{cy}"/>
                        </a:xfrm>
                        <a:prstGeom prst="rect">
                            <a:avLst/>
                        </a:prstGeom>
                    </pic:spPr>
                </pic:pic>
            </a:graphicData>
        </a:graphic>
    </wp:anchor>
    """

    # Parse anchor XML
    anchor_el = etree.fromstring(anchor_xml.encode())

    # Create drawing wrapper
    drawing_el = etree.Element(qn("w:drawing"))
    drawing_el.append(anchor_el)

    # Find target paragraph and add drawing
    if target.leaf_kind == "paragraph":
        para = target.leaf_obj
    elif target.leaf_kind == "cell":
        cell = target.leaf_obj
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    else:
        # Create new paragraph for floating image
        para = doc.add_paragraph()
        if hasattr(target, "leaf_el"):
            target.leaf_el.addnext(para._element)

    # Add run with drawing
    run = para.add_run()
    run._element.append(drawing_el)

    # Mark document.xml as modified for WordPackage
    mark_dirty(doc)

    occurrence = count_image_occurrence(doc, h, para._element)
    return f"image_{h}_{occurrence}"


# =============================================================================
# Text Box Support
# =============================================================================


def _textbox_xpath(element, expr: str) -> list:
    """Execute XPath with text box namespaces."""
    return _LxmlElementBase.xpath(element, expr, namespaces=_TEXTBOX_NS)


def _extract_text_from_txbxcontent(txbx_content) -> tuple[str, int]:
    """Extract text and paragraph count from w:txbxContent element."""
    paragraphs = _textbox_xpath(txbx_content, ".//w:p")
    para_texts = []
    for p in paragraphs:
        # Get all text nodes within the paragraph
        texts = _textbox_xpath(p, ".//w:t/text()")
        para_texts.append("".join(texts))
    return "\n".join(para_texts), len(paragraphs)


def _get_textbox_dimensions(ancestor) -> tuple[float, float]:
    """Extract width/height in inches from DrawingML extent or VML style."""
    # DrawingML: look for wp:extent
    extents = _textbox_xpath(ancestor, ".//wp:extent")
    if extents:
        ext = extents[0]
        cx = ext.get("cx")  # EMUs
        cy = ext.get("cy")
        width = int(cx) / _EMU_PER_INCH if cx else 0.0
        height = int(cy) / _EMU_PER_INCH if cy else 0.0
        return width, height

    # VML: parse style attribute
    shapes = _textbox_xpath(ancestor, ".//*[starts-with(name(), 'v:')][@style]")
    if shapes:
        style = shapes[0].get("style", "")
        width = height = 0.0
        for part in style.split(";"):
            if ":" in part:
                key, val = part.split(":", 1)
                key = key.strip().lower()
                val = val.strip().lower()
                if key == "width" and val.endswith("pt"):
                    width = float(val[:-2]) / 72.0
                elif key == "height" and val.endswith("pt"):
                    height = float(val[:-2]) / 72.0
        return width, height

    return 0.0, 0.0


def _get_textbox_id_and_name(
    ancestor, source_type: str, index: int
) -> tuple[str, str | None]:
    """Extract stable ID and name from DrawingML docPr or VML shape."""
    if source_type == "drawingml":
        # Look for wp:docPr
        doc_prs = _textbox_xpath(ancestor, ".//wp:docPr")
        if doc_prs:
            doc_pr = doc_prs[0]
            tb_id = doc_pr.get("id", "")
            name = doc_pr.get("name")
            if tb_id:
                return f"textbox_{tb_id}", name
    else:
        # VML: use v:shape @id
        shapes = _textbox_xpath(ancestor, ".//*[starts-with(name(), 'v:')][@id]")
        if shapes:
            shape_id = shapes[0].get("id", "")
            if shape_id:
                return f"textbox_vml_{shape_id}", None

    # Fallback to index-based ID
    return f"textbox_{source_type}_{index}", None


def _get_position_type(ancestor) -> str:
    """Determine if text box is anchored (floating) or inline."""
    # DrawingML: wp:anchor = floating, wp:inline = inline
    # Check both descendants (from w:drawing) and self (from wp:anchor/wp:inline)
    if _textbox_xpath(ancestor, ".//wp:anchor") or ancestor.tag == qn("wp:anchor"):
        return "anchor"
    if _textbox_xpath(ancestor, ".//wp:inline") or ancestor.tag == qn("wp:inline"):
        return "inline"
    # VML in w:pict is typically anchored
    if ancestor.tag == qn("w:pict") or _textbox_xpath(ancestor, ".//v:shape"):
        return "anchor"
    return "unknown"


def _get_wrap_type(ancestor) -> str | None:
    """Get wrap type from DrawingML wrap elements."""
    wrap_map = {
        "wrapSquare": "square",
        "wrapTight": "tight",
        "wrapThrough": "through",
        "wrapTopAndBottom": "topAndBottom",
        "wrapNone": "none",
    }
    for wrap_name, wrap_type in wrap_map.items():
        if _textbox_xpath(ancestor, f".//wp:{wrap_name}"):
            return wrap_type
    return None


def build_text_boxes(pkg) -> list[dict]:
    """Build list of all text boxes from both DrawingML and VML.

    Discovery strategy (per OpenAI review):
    1. Search for ALL w:txbxContent elements in document
    2. Classify source by ancestor chain (w:drawing = DrawingML, w:pict = VML)
    3. Handle mc:AlternateContent (both mc:Choice and mc:Fallback)

    Args:
        pkg: WordPackage or python-docx Document (duck-typed)
    """
    text_boxes = []
    seen_ids = set()

    # Get document element (duck-typed for WordPackage vs Document)
    if hasattr(pkg, "document_xml"):
        doc_element = pkg.document_xml
    else:
        doc_element = pkg.element

    # Search for all w:txbxContent elements
    txbx_contents = _textbox_xpath(doc_element, "//w:txbxContent")

    for idx, txbx in enumerate(txbx_contents):
        # Determine source type by ancestor chain
        parent_chain = list(txbx.iterancestors())
        has_drawing = any(
            p.tag == qn("w:drawing") or "drawing" in str(p.tag).lower()
            for p in parent_chain
        )
        has_pict = any(p.tag == qn("w:pict") for p in parent_chain)

        if has_drawing:
            source_type = "drawingml"
            # Find the w:drawing ancestor for dimension extraction
            ancestor = next(
                (p for p in parent_chain if p.tag == qn("w:drawing")),
                txbx.getparent(),
            )
        elif has_pict:
            source_type = "vml"
            ancestor = next(
                (p for p in parent_chain if p.tag == qn("w:pict")),
                txbx.getparent(),
            )
        else:
            # Unknown container type
            source_type = "unknown"
            ancestor = txbx.getparent()

        # Extract text and paragraph count
        text, para_count = _extract_text_from_txbxcontent(txbx)

        # Get ID and name
        tb_id, name = _get_textbox_id_and_name(ancestor, source_type, idx)

        # Skip duplicates (mc:AlternateContent can have same content twice)
        if tb_id in seen_ids:
            continue
        seen_ids.add(tb_id)

        # Get dimensions
        width, height = _get_textbox_dimensions(ancestor)

        # Get position type
        position_type = _get_position_type(ancestor)

        # Get wrap type
        wrap_type = _get_wrap_type(ancestor)

        text_boxes.append(
            {
                "id": tb_id,
                "name": name,
                "text": text,
                "paragraph_count": para_count,
                "width_inches": round(width, 2),
                "height_inches": round(height, 2),
                "position_type": position_type,
                "source_type": source_type,
                "wrap_type": wrap_type,
            }
        )

    return text_boxes


def _find_textbox_content_by_id(pkg, textbox_id: str):
    """Find w:txbxContent element by text box ID.

    Args:
        pkg: WordPackage or python-docx Document (duck-typed)
    """
    # Get document element (duck-typed)
    if hasattr(pkg, "document_xml"):
        doc_element = pkg.document_xml
    else:
        doc_element = pkg.element

    text_boxes = build_text_boxes(pkg)
    for idx, tb in enumerate(text_boxes):
        if tb["id"] == textbox_id:
            # Re-find the actual element
            txbx_contents = _textbox_xpath(doc_element, "//w:txbxContent")
            # Need to find the matching one by index
            # Since build_text_boxes filters duplicates, we track which we've seen
            seen_ids = set()
            for tidx, txbx in enumerate(txbx_contents):
                parent_chain = list(txbx.iterancestors())
                has_drawing = any(
                    p.tag == qn("w:drawing") or "drawing" in str(p.tag).lower()
                    for p in parent_chain
                )
                has_pict = any(p.tag == qn("w:pict") for p in parent_chain)
                source_type = (
                    "drawingml" if has_drawing else "vml" if has_pict else "unknown"
                )
                # Find appropriate ancestor element
                drawing_ancestor = next(
                    (p for p in parent_chain if p.tag == qn("w:drawing")), None
                )
                pict_ancestor = next(
                    (p for p in parent_chain if p.tag == qn("w:pict")), None
                )
                if drawing_ancestor is not None:
                    ancestor = drawing_ancestor
                elif pict_ancestor is not None:
                    ancestor = pict_ancestor
                else:
                    ancestor = txbx.getparent()
                tb_id, _ = _get_textbox_id_and_name(ancestor, source_type, tidx)
                if tb_id in seen_ids:
                    continue
                seen_ids.add(tb_id)
                if tb_id == textbox_id:
                    return txbx
    return None


def read_text_box_content(pkg, textbox_id: str) -> list[dict]:
    """Read paragraphs inside a text box.

    Returns list of dicts with 'index', 'text', and basic formatting info.

    Args:
        pkg: WordPackage or python-docx Document (duck-typed)
    """
    txbx = _find_textbox_content_by_id(pkg, textbox_id)
    if txbx is None:
        raise ValueError(f"Text box not found: {textbox_id}")

    paragraphs = _textbox_xpath(txbx, "./w:p")
    result = []
    for i, p in enumerate(paragraphs):
        texts = _textbox_xpath(p, ".//w:t/text()")
        text = "".join(texts)
        result.append(
            {
                "index": i,
                "text": text,
                "id": f"{textbox_id}/p{i}",
            }
        )
    return result


def edit_text_box_text(
    doc: Document, textbox_id: str, para_index: int, new_text: str
) -> None:
    """Edit text in a text box paragraph.

    Replaces all text in the specified paragraph with new_text.
    """
    txbx = _find_textbox_content_by_id(doc, textbox_id)
    if txbx is None:
        raise ValueError(f"Text box not found: {textbox_id}")

    paragraphs = _textbox_xpath(txbx, "./w:p")
    if para_index < 0 or para_index >= len(paragraphs):
        raise ValueError(
            f"Paragraph index {para_index} out of range (0-{len(paragraphs) - 1})"
        )

    p = paragraphs[para_index]

    # Find all runs and clear their text
    runs = _textbox_xpath(p, "./w:r")
    if runs:
        # Clear all runs and set text in first one
        for run in runs:
            for t in _textbox_xpath(run, "./w:t"):
                t.text = ""
        # Set new text in first run's first w:t
        first_run = runs[0]
        t_elements = _textbox_xpath(first_run, "./w:t")
        if t_elements:
            t_elements[0].text = new_text
        else:
            # Create w:t element
            t = etree.Element(qn("w:t"))
            t.text = new_text
            first_run.append(t)
    else:
        # No runs - create one
        run = etree.Element(qn("w:r"))
        t = etree.Element(qn("w:t"))
        t.text = new_text
        run.append(t)
        p.append(run)

    # Mark document.xml as modified for WordPackage
    mark_dirty(doc)
