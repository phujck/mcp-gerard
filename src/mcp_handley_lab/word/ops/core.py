"""Core utilities for Word document operations.

Contains shared helpers used across all ops modules:
- Content-addressed ID generation and resolution
- Block iteration and traversal
- Target resolution for hierarchical addressing
- Common XML element operations
"""

from __future__ import annotations

import hashlib
import json
import re
from collections.abc import Iterator
from dataclasses import dataclass
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap as oxml_nsmap
from docx.oxml.table import CT_Tbl, CT_Tc
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    from docx.text.run import Run

from mcp_handley_lab.word.models import Block

# Regexes for parsing IDs and styles
_HEADING_RE = re.compile(r"^Heading ([1-9])$")
_ID_RE = re.compile(r"^(paragraph|heading[1-9]|table)_([0-9a-f]{8})_(\d+)$")
_IMAGE_ID_RE = re.compile(r"^image_([0-9a-f]{8})_(\d+)$")

# Hierarchical path segment patterns (0-based indices)
_CELL_RE = re.compile(r"^r(\d+)c(\d+)$")
_PARA_RE = re.compile(r"^p(\d+)$")
_TABLE_RE = re.compile(r"^tbl(\d+)$")

# Unit conversions
_EMU_PER_INCH = 914400

# PathSegment: (kind, indices) where kind='cell'|'para'|'tbl'
PathSegment = tuple[str, tuple[int, ...]]


@dataclass
class ResolvedTarget:
    """Result of resolving a hierarchical target ID."""

    base_id: str
    base_kind: str  # 'table' | 'paragraph' | 'heading1' etc
    base_obj: Paragraph | Table
    base_occurrence: int
    leaf_kind: str  # 'table' | 'cell' | 'paragraph'
    leaf_obj: Paragraph | Table | _Cell
    leaf_el: CT_P | CT_Tbl | CT_Tc  # XML element (Paragraph, Table, or Cell)


# =============================================================================
# Content Hashing and ID Generation
# =============================================================================


def content_hash(text: str) -> str:
    """8-char SHA256 of normalized text for content-addressable IDs."""
    normalized = re.sub(r"\s+", " ", text.strip().lower())
    return hashlib.sha256(normalized.encode()).hexdigest()[:8]


def make_block_id(block_type: str, text: str, occurrence: int) -> str:
    """Generate content-addressable block ID. Single source of truth."""
    return f"{block_type}_{content_hash(text)}_{occurrence}"


def table_content_for_hash(table: Table) -> str:
    """Get full table content as canonical string for hashing."""
    return json.dumps(
        [[c.text for c in r.cells] for r in table.rows],
        ensure_ascii=False,
        separators=(",", ":"),
    )


def paragraph_kind_and_level(p: Paragraph) -> tuple[str, int]:
    """Detect if paragraph is a heading and return (kind, level).

    For content-hash IDs, kind includes level: 'heading1', 'heading2', etc.
    """
    style_name = p.style.name if p.style else "Normal"
    m = _HEADING_RE.match(style_name)
    if m:
        level = int(m.group(1))
        return (f"heading{level}", level)
    return ("paragraph", 0)


# =============================================================================
# Block Iteration
# =============================================================================


def iter_body_blocks(
    doc: Document,
) -> Iterator[tuple[str, Paragraph | Table, CT_P | CT_Tbl]]:
    """Yield (block_kind, block_obj, element) in true document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, doc), child)
        elif isinstance(child, CT_Tbl):
            yield ("table", Table(child, doc), child)


def _iter_all_paragraphs(doc: Document) -> Iterator[tuple[Paragraph, CT_P]]:
    """Iterate over all paragraphs in document body and tables."""
    for kind, obj, el in iter_body_blocks(doc):
        if kind == "paragraph":
            yield obj, el
        elif kind == "table":
            for row in obj.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para, para._element


def _iter_all_runs_in_paragraph(para: Paragraph) -> Iterator[Run]:
    """Iterate all runs in paragraph including those inside hyperlinks.

    Uses iter_inner_content() to match build_runs() indexing.
    """
    for item in para.iter_inner_content():
        if isinstance(item, Hyperlink):
            yield from item.runs
        else:  # Run
            yield item


# =============================================================================
# Target Resolution
# =============================================================================


def parse_target_id(target_id: str) -> tuple[str, list[PathSegment]]:
    """Parse target_id into (base_id, path_segments).

    Supports hierarchical IDs: table_abc12345_0#r0c1/p0
    Returns (base_id, []) for simple IDs without path.
    """
    base_id, *rest = target_id.split("#", 1)
    if not rest:
        return target_id, []

    segments: list[PathSegment] = []
    for part in rest[0].split("/"):
        if m := _CELL_RE.match(part):
            segments.append(("cell", (int(m[1]), int(m[2]))))
        elif m := _PARA_RE.match(part):
            segments.append(("para", (int(m[1]),)))
        elif m := _TABLE_RE.match(part):
            segments.append(("tbl", (int(m[1]),)))
        else:
            raise ValueError(f"Invalid path segment: {part}")
    return base_id, segments


def resolve_path(
    container: Table | _Cell, segments: list[PathSegment]
) -> Paragraph | Table | _Cell:
    """Resolve path segments from container. Returns leaf object.

    Transition rules:
    - From Table -> cell (r,c) only
    - From Cell -> para (p) or tbl (tbl)
    - From nested Table -> cell (r,c)
    """
    current: Table | _Cell | Paragraph = container

    for kind, indices in segments:
        if kind == "cell":
            row, col = indices
            current = current.cell(
                row, col
            )  # AttributeError if not Table, IndexError if OOB
        elif kind == "para":
            current = current.paragraphs[indices[0]]  # AttributeError if not Cell
        elif kind == "tbl":
            current = current.tables[indices[0]]  # AttributeError if not Cell

    return current


def _resolve_base_block(
    doc: Document, base_id: str
) -> tuple[str, Paragraph | Table, CT_P | CT_Tbl, int]:
    """Resolve base block ID to (block_type, obj, element, occurrence).

    Uses content-hash IDs: {type}_{hash}_{occurrence}
    Searches all blocks for matching type+hash, then skips to Nth occurrence.
    """
    target_type, target_hash, occurrence_str = _ID_RE.match(base_id).groups()
    target_occurrence = int(occurrence_str)

    occurrence_count = 0
    for kind, obj, el in iter_body_blocks(doc):
        if kind == "table":
            if target_type == "table":
                table_content = table_content_for_hash(obj)
                if content_hash(table_content) == target_hash:
                    if occurrence_count == target_occurrence:
                        return "table", obj, el, occurrence_count
                    occurrence_count += 1
        else:
            block_type, _ = paragraph_kind_and_level(obj)
            if target_type == block_type:
                text = obj.text or ""
                if content_hash(text) == target_hash:
                    if occurrence_count == target_occurrence:
                        return block_type, obj, el, occurrence_count
                    occurrence_count += 1
    raise ValueError(f"Block not found: {base_id}")


def resolve_target(doc: Document, target_id: str) -> ResolvedTarget:
    """Resolve target_id to ResolvedTarget with base and leaf info.

    Supports hierarchical IDs: table_abc12345_0#r0c1/p0
    """
    base_id, path_segments = parse_target_id(target_id)

    # Resolve base block
    base_kind, base_obj, base_el, base_occurrence = _resolve_base_block(doc, base_id)

    if not path_segments:
        return ResolvedTarget(
            base_id=base_id,
            base_kind=base_kind,
            base_obj=base_obj,
            base_occurrence=base_occurrence,
            leaf_kind=base_kind,
            leaf_obj=base_obj,
            leaf_el=base_el,
        )

    # Resolve path within container (AttributeError if base_obj not a Table)
    leaf_obj = resolve_path(base_obj, path_segments)

    # Determine leaf kind and element based on type
    if isinstance(leaf_obj, Table):
        return ResolvedTarget(
            base_id=base_id,
            base_kind=base_kind,
            base_obj=base_obj,
            base_occurrence=base_occurrence,
            leaf_kind="table",
            leaf_obj=leaf_obj,
            leaf_el=leaf_obj._tbl,
        )
    if isinstance(leaf_obj, _Cell):
        return ResolvedTarget(
            base_id=base_id,
            base_kind=base_kind,
            base_obj=base_obj,
            base_occurrence=base_occurrence,
            leaf_kind="cell",
            leaf_obj=leaf_obj,
            leaf_el=leaf_obj._tc,
        )
    # Paragraph
    return ResolvedTarget(
        base_id=base_id,
        base_kind=base_kind,
        base_obj=base_obj,
        base_occurrence=base_occurrence,
        leaf_kind="paragraph",
        leaf_obj=leaf_obj,
        leaf_el=leaf_obj._element,
    )


def find_paragraph_by_id(doc: Document, target_id: str) -> Paragraph | None:
    """Find a paragraph by its block ID.

    Returns the Paragraph object if found and is a paragraph, None otherwise.
    """
    try:
        target = resolve_target(doc, target_id)
        if target.leaf_kind == "paragraph" and isinstance(target.leaf_obj, Paragraph):
            return target.leaf_obj
        return None
    except ValueError:
        return None


def count_occurrence(
    doc: Document, block_type: str, text: str, target_el: CT_P | CT_Tbl
) -> int:
    """Count how many blocks with same type+hash appear before target_el.

    Used after edits to compute the correct occurrence index for returned ID.
    For tables, 'text' should be the full table content (from table_content_for_hash).
    """
    target_hash = content_hash(text)
    occurrence = 0
    for kind, obj, el in iter_body_blocks(doc):
        if kind == "table":
            if block_type == "table":
                table_content = table_content_for_hash(obj)
                if content_hash(table_content) == target_hash:
                    if el is target_el:
                        return occurrence
                    occurrence += 1
        else:
            actual_type, _ = paragraph_kind_and_level(obj)
            if block_type == actual_type:
                block_text = obj.text or ""
                if content_hash(block_text) == target_hash:
                    if el is target_el:
                        return occurrence
                    occurrence += 1
    raise ValueError("target_el not found in document")


# =============================================================================
# Block Building
# =============================================================================


def build_blocks(
    doc: Document,
    offset: int = 0,
    limit: int = 50,
    heading_only: bool = False,
    search_query: str = "",
) -> tuple[list[Block], int]:
    """Build list of Block objects from document body.

    Uses content-hash IDs: {type}_{hash}_{occurrence}
    """
    # Import here to avoid circular dependency (table_to_markdown is in tables.py)
    from mcp_handley_lab.word.ops.tables import table_to_markdown

    blocks = []
    matched = 0
    query_lower = search_query.lower() if search_query else ""
    hash_counts: dict[str, int] = {}  # {type_hash: count} for occurrence tracking

    for kind, obj, _el in iter_body_blocks(doc):
        if kind == "paragraph":
            block_type, level = paragraph_kind_and_level(obj)
            text = obj.text or ""
            style = obj.style.name if obj.style else "Normal"

            # Track occurrence for this type+hash combination
            hash_key = f"{block_type}_{content_hash(text)}"
            occurrence = hash_counts.get(hash_key, 0)
            hash_counts[hash_key] = occurrence + 1

            if heading_only and not block_type.startswith("heading"):
                continue
            if query_lower and query_lower not in text.lower():
                continue

            if matched >= offset and len(blocks) < limit:
                blocks.append(
                    Block(
                        id=make_block_id(block_type, text, occurrence),
                        type=block_type,
                        text=text,
                        style=style,
                        level=level,
                    )
                )
            matched += 1

        elif kind == "table":
            md, rows, cols = table_to_markdown(obj)
            style = obj.style.name if obj.style else ""
            # Use full content for hash (not truncated markdown)
            table_content = table_content_for_hash(obj)

            # Track occurrence for this type+hash combination
            hash_key = f"table_{content_hash(table_content)}"
            occurrence = hash_counts.get(hash_key, 0)
            hash_counts[hash_key] = occurrence + 1

            if heading_only:
                continue
            # Search full table content, not truncated markdown preview
            if query_lower and query_lower not in table_content.lower():
                continue

            if matched >= offset and len(blocks) < limit:
                blocks.append(
                    Block(
                        id=make_block_id("table", table_content, occurrence),
                        type="table",
                        text=md,  # Keep markdown for display
                        style=style,
                        rows=rows,
                        cols=cols,
                    )
                )
            matched += 1

    return blocks, matched


# =============================================================================
# Element Operations
# =============================================================================


def _make_run_with(*elements) -> OxmlElement:
    """Create a w:r element containing the given child elements."""
    r = OxmlElement("w:r")
    for elem in elements:
        r.append(elem)
    return r


def _insert_at(target_el, new_el, position: str) -> None:
    """Insert new_el before or after target_el."""
    (target_el.addprevious if position == "before" else target_el.addnext)(new_el)


# =============================================================================
# Re-exports for convenience
# =============================================================================

# Make oxml_nsmap available for other ops modules
nsmap = oxml_nsmap
