"""Word document MCP tool - read and edit operations."""

from __future__ import annotations

import json
from typing import TYPE_CHECKING

from mcp.server.fastmcp import FastMCP
from pydantic import Field

from mcp_handley_lab.word import document as word_ops

if TYPE_CHECKING:
    pass
from mcp_handley_lab.word.models import (
    BookmarkInfo,
    CaptionInfo,
    CommentInfo,
    ContentControlInfo,
    DocumentReadResult,
    EditResult,
    EquationInfo,
    FootnoteInfo,
    ListInfo,
    RevisionInfo,
    TextBoxInfo,
    TOCInfo,
)
from mcp_handley_lab.word.opc import WordPackage

mcp = FastMCP("Word Document Tool")


# Helper dispatch dicts for header/footer operations
_HF_SET_OPS = {
    "set_header": "header",
    "set_footer": "footer",
    "set_first_page_header": "first_page_header",
    "set_first_page_footer": "first_page_footer",
    "set_even_page_header": "even_page_header",
    "set_even_page_footer": "even_page_footer",
}
_HF_APPEND_OPS = {"append_header": "header", "append_footer": "footer"}
_HF_CLEAR_OPS = {"clear_header": "header", "clear_footer": "footer"}


def _recalc_table_id(doc, t) -> str:
    """Recalculate table element ID after modification. Requires base_kind == 'table'."""
    if t.base_kind != "table":
        raise ValueError("Expected base_kind=table for table ID recalculation")
    tbl_el = t.base_el  # Use element, not wrapper
    content = word_ops.table_content_for_hash(tbl_el)
    occurrence = word_ops.count_occurrence(doc, "table", content, tbl_el)
    return word_ops.make_block_id("table", content, occurrence)


def _recalc_block_id(doc, t) -> str:
    """Recalculate block element ID after modification. Uses t.leaf_kind directly."""
    text = t.leaf_obj.text or ""
    occurrence = word_ops.count_occurrence(doc, t.leaf_kind, text, t.leaf_el)
    return word_ops.make_block_id(t.leaf_kind, text, occurrence)


def _get_target_paragraph(target):
    """Get paragraph from target. For cells, uses first paragraph (must exist)."""
    return (
        target.leaf_obj.paragraphs[0] if target.leaf_kind == "cell" else target.leaf_obj
    )


@mcp.tool(
    description="Read Word document content. Scopes: 'meta' (doc info), 'outline' (headings only), 'blocks' (all content), 'search' (find text), 'table_cells' (cells of a table), 'table_layout' (table alignment/autofit/row heights), 'runs' (text runs in a paragraph), 'comments' (all comments), 'headers_footers' (headers/footers per section), 'page_setup' (margins, orientation per section), 'images' (embedded inline images), 'hyperlinks' (all hyperlinks with URLs), 'styles' (all document styles), 'style' (detailed formatting for a specific style by name in target_id), 'revisions' (tracked changes/revisions), 'list' (list properties for a paragraph by target_id), 'text_boxes' (all text boxes/floating content), 'text_box_content' (paragraphs inside a text box by target_id), 'bookmarks' (all bookmarks), 'captions' (all captions), 'toc' (table of contents info), 'footnotes' (all footnotes and endnotes), 'content_controls' (all content controls/SDTs), 'equations' (math equations with simplified text). Block IDs are content-addressed (type_hash_occurrence) and CHANGE when content changes or after inserts/deletes shift occurrence index - use element_id from edit response for chaining."
)
def read(
    file_path: str = Field(..., description="Path to .docx file"),
    scope: str = Field(
        "outline",
        description="What to read: 'meta', 'outline', 'blocks', 'search', 'table_cells', 'table_layout', 'runs', 'comments', 'headers_footers', 'page_setup', 'images', 'hyperlinks', 'styles', 'style', 'revisions', 'list', 'text_boxes', 'text_box_content', 'bookmarks', 'captions', 'toc', 'footnotes', 'content_controls', 'equations'",
    ),
    target_id: str = Field(
        "",
        description="Block ID for table_cells/runs/table_layout/list scopes, style name for 'style' scope, or text box ID for 'text_box_content' scope",
    ),
    search_query: str = Field("", description="Text to search for (scope='search')"),
    limit: int = Field(50, description="Max blocks to return"),
    offset: int = Field(0, description="Pagination offset"),
) -> DocumentReadResult:
    """Read Word document content."""
    # Scopes that need WordPackage (pure OOXML)
    _PKG_SCOPES = {
        "meta",
        "outline",
        "blocks",
        "search",
        "table_cells",
        "table_layout",
        "revisions",
        "list",
        "bookmarks",
        "captions",
        "toc",
        "content_controls",
        "equations",
        "footnotes",
        "hyperlinks",
        "text_boxes",
        "text_box_content",
        "page_setup",
        "styles",
        "style",
        "runs",
        "comments",
        "headers_footers",
        "images",
    }
    # Scopes that need python-docx Document (for now)
    _DOC_SCOPES: set[str] = set()

    # Load appropriate object(s)
    pkg = (
        WordPackage.open(file_path) if scope in _PKG_SCOPES or scope == "meta" else None
    )

    if scope == "meta":
        from docx import Document

        doc = Document(file_path)
        meta = word_ops.get_document_meta(doc)
        _, block_count = word_ops.build_blocks(pkg, offset=0, limit=0)
        return DocumentReadResult(block_count=block_count, meta=meta)
    if scope == "outline":
        blocks, block_count = word_ops.build_blocks(
            pkg, offset=offset, limit=limit, heading_only=True
        )
        return DocumentReadResult(block_count=block_count, blocks=blocks)
    if scope == "blocks":
        blocks, block_count = word_ops.build_blocks(pkg, offset=offset, limit=limit)
        return DocumentReadResult(block_count=block_count, blocks=blocks)
    if scope == "search":
        blocks, block_count = word_ops.build_blocks(
            pkg, offset=offset, limit=limit, search_query=search_query
        )
        return DocumentReadResult(block_count=block_count, blocks=blocks)
    if scope == "table_cells":
        from mcp_handley_lab.word.opc.constants import qn

        t = word_ops.resolve_target(pkg, target_id)
        tbl_el = t.base_el if t.base_kind == "table" else t.leaf_el
        cells = word_ops.build_table_cells(tbl_el, t.base_id)
        # Count rows/cols from element
        rows = tbl_el.findall(qn("w:tr"))
        max_cols = max((len(tr.findall(qn("w:tc"))) for tr in rows), default=0)
        return DocumentReadResult(
            block_count=len(cells),
            cells=cells,
            table_rows=len(rows),
            table_cols=max_cols,
        )
    if scope == "table_layout":
        t = word_ops.resolve_target(pkg, target_id)
        tbl_el = t.base_el if t.base_kind == "table" else t.leaf_el
        table_layout = word_ops.build_table_layout(tbl_el, t.base_id)
        return DocumentReadResult(block_count=1, table_layout=table_layout)
    if scope == "runs":
        # Pure OOXML path - find paragraph element, get doc rels for hyperlink URLs
        t = word_ops.resolve_target(pkg, target_id)
        if t.leaf_kind != "paragraph":
            raise ValueError(
                f"Target is not a paragraph: {target_id} (resolved to {t.leaf_kind})"
            )
        p_el = t.leaf_el
        doc_rels = pkg.get_rels("/word/document.xml")
        runs = word_ops.build_runs(p_el, doc_rels)
        paragraph_format = word_ops.build_paragraph_format(p_el)
        return DocumentReadResult(
            block_count=len(runs), runs=runs, paragraph_format=paragraph_format
        )
    if scope == "comments":
        comments_data = word_ops.build_comments_with_threading(pkg)
        comments = [CommentInfo(**c) for c in comments_data]
        return DocumentReadResult(block_count=len(comments), comments=comments)
    if scope == "headers_footers":
        headers_footers = word_ops.build_headers_footers(pkg)
        return DocumentReadResult(
            block_count=len(headers_footers), headers_footers=headers_footers
        )
    if scope == "page_setup":
        page_setup = word_ops.build_page_setup(pkg)
        return DocumentReadResult(block_count=len(page_setup), page_setup=page_setup)
    if scope == "images":
        images = word_ops.build_images(pkg)
        return DocumentReadResult(block_count=len(images), images=images)
    if scope == "hyperlinks":
        hyperlinks = word_ops.build_hyperlinks(pkg)
        return DocumentReadResult(block_count=len(hyperlinks), hyperlinks=hyperlinks)
    if scope == "styles":
        styles = word_ops.build_styles(pkg)
        return DocumentReadResult(block_count=len(styles), styles=styles)
    if scope == "style":
        style_format = word_ops.get_style_format(pkg, target_id)
        return DocumentReadResult(block_count=1, style_format=style_format)
    if scope == "revisions":
        has_changes = word_ops.has_tracked_changes(pkg)
        revisions = (
            [RevisionInfo(**r) for r in word_ops.read_tracked_changes(pkg)]
            if has_changes
            else []
        )
        return DocumentReadResult(
            block_count=len(revisions),
            has_tracked_changes=has_changes,
            revisions=revisions,
        )
    if scope == "list":
        if not target_id:
            raise ValueError("target_id required for list scope")
        p_el = word_ops.find_paragraph_by_id(pkg, target_id)
        if p_el is None:
            raise ValueError(f"Paragraph not found: {target_id}")
        list_info_dict = word_ops.get_list_info(pkg, p_el)
        list_info = ListInfo(**list_info_dict) if list_info_dict else None
        return DocumentReadResult(
            block_count=1 if list_info else 0, list_info=list_info
        )
    if scope == "text_boxes":
        text_boxes_data = word_ops.build_text_boxes(pkg)
        text_boxes = [TextBoxInfo(**tb) for tb in text_boxes_data]
        return DocumentReadResult(block_count=len(text_boxes), text_boxes=text_boxes)
    if scope == "text_box_content":
        if not target_id:
            raise ValueError("target_id required for text_box_content scope")
        paragraphs = word_ops.read_text_box_content(pkg, target_id)
        # Return paragraphs as blocks for consistency
        from mcp_handley_lab.word.models import Block

        blocks = [
            Block(
                id=p["id"],
                type="paragraph",
                text=p["text"],
                style="Normal",
            )
            for p in paragraphs
        ]
        return DocumentReadResult(block_count=len(blocks), blocks=blocks)
    if scope == "bookmarks":
        bookmarks_data = word_ops.build_bookmarks(pkg)
        bookmarks = [BookmarkInfo(**bm) for bm in bookmarks_data]
        return DocumentReadResult(block_count=len(bookmarks), bookmarks=bookmarks)
    if scope == "captions":
        captions_data = word_ops.build_captions(pkg)
        captions = [CaptionInfo(**c) for c in captions_data]
        return DocumentReadResult(block_count=len(captions), captions=captions)
    if scope == "toc":
        toc_data = word_ops.get_toc_info(pkg)
        toc_info = TOCInfo(**toc_data)
        return DocumentReadResult(
            block_count=1 if toc_info.exists else 0, toc_info=toc_info
        )
    if scope == "footnotes":
        footnotes_data = word_ops.build_footnotes(pkg)
        footnotes = [FootnoteInfo(**fn) for fn in footnotes_data]
        return DocumentReadResult(block_count=len(footnotes), footnotes=footnotes)
    if scope == "content_controls":
        cc_data = word_ops.build_content_controls(pkg)
        controls = [ContentControlInfo(**cc) for cc in cc_data]
        return DocumentReadResult(block_count=len(controls), content_controls=controls)
    if scope == "equations":
        eq_data = word_ops.build_equations(pkg)
        equations = [EquationInfo(**eq) for eq in eq_data]
        return DocumentReadResult(block_count=len(equations), equations=equations)
    raise ValueError(f"Unknown scope: {scope}")


@mcp.tool(
    description="Edit Word document. Operations: 'create', 'insert_before', 'insert_after', 'append', 'delete', 'replace', 'style', 'edit_cell', 'edit_run', 'edit_style', 'add_comment', 'reply_comment', 'resolve_comment', 'unresolve_comment', 'set_header', 'set_footer', 'set_first_page_header', 'set_first_page_footer', 'set_even_page_header', 'set_even_page_footer', 'append_header', 'append_footer', 'clear_header', 'clear_footer', 'set_margins', 'set_orientation', 'set_columns', 'set_line_numbering', 'set_custom_property', 'delete_custom_property', 'create_style', 'delete_style', 'insert_image', 'insert_floating_image', 'delete_image', 'add_row', 'add_column', 'delete_row', 'delete_column', 'add_page_break', 'add_break', 'set_meta', 'add_section', 'merge_cells', 'set_table_alignment', 'set_table_autofit', 'set_table_fixed_layout', 'set_row_height', 'set_cell_width', 'set_cell_vertical_alignment', 'set_cell_borders', 'set_cell_shading', 'set_header_row', 'add_tab_stop', 'clear_tab_stops', 'insert_field', 'insert_page_x_of_y', 'accept_change', 'reject_change', 'accept_all_changes', 'reject_all_changes', 'set_list_level', 'promote_list', 'demote_list', 'restart_numbering', 'remove_list', 'edit_text_box', 'add_bookmark', 'add_hyperlink', 'insert_cross_ref', 'insert_caption', 'insert_toc', 'update_toc', 'add_footnote', 'delete_footnote', 'set_content_control'. Block IDs are content-addressed and CHANGE when content changes or after inserts/deletes shift occurrence index. Always use element_id from response for chaining operations on modified content. Note: 'create' makes a doc with an initial empty paragraph (python-docx behavior). 'update_toc' sets dirty flag; Word updates content on open."
)
def edit(
    file_path: str = Field(..., description="Path to .docx file"),
    operation: str = Field(
        ...,
        description="Operation: 'create', 'insert_before', 'insert_after', 'append', 'delete', 'replace', 'style', 'edit_cell', 'edit_run', 'edit_style', 'create_style', 'delete_style', 'add_comment', 'reply_comment', 'resolve_comment', 'unresolve_comment', 'set_header', 'set_footer', 'set_first_page_header', 'set_first_page_footer', 'set_even_page_header', 'set_even_page_footer', 'append_header', 'append_footer', 'clear_header', 'clear_footer', 'set_margins', 'set_orientation', 'set_columns', 'set_line_numbering', 'set_custom_property', 'delete_custom_property', 'insert_image', 'delete_image', 'add_row', 'add_column', 'delete_row', 'delete_column', 'add_page_break', 'add_break', 'set_meta', 'add_section', 'merge_cells', 'set_table_alignment', 'set_table_autofit', 'set_table_fixed_layout', 'set_row_height', 'set_cell_width', 'set_cell_vertical_alignment', 'set_cell_borders', 'set_cell_shading', 'set_header_row', 'add_tab_stop', 'clear_tab_stops', 'insert_field', 'insert_page_x_of_y', 'accept_change', 'reject_change', 'accept_all_changes', 'reject_all_changes', 'set_list_level', 'promote_list', 'demote_list', 'restart_numbering', 'remove_list', 'edit_text_box', 'add_bookmark', 'add_hyperlink', 'insert_cross_ref', 'insert_caption', 'insert_toc', 'update_toc', 'add_footnote', 'delete_footnote', 'set_content_control'",
    ),
    target_id: str = Field(
        "",
        description="Block ID from read() - required for insert/delete/replace/style/edit_cell/edit_run/add_comment/insert_image/list operations and table operations. For edit_style: style name. For delete_style: style name to delete. For delete_image: the image ID. For accept_change/reject_change: the revision ID from read(scope='revisions'). For edit_text_box: the text box ID from read(scope='text_boxes'). For add_bookmark/add_hyperlink/insert_cross_ref/insert_caption/insert_toc: the paragraph or block ID (or table cell ID). For reply_comment/resolve_comment/unresolve_comment: the comment ID. For add_footnote: the paragraph ID. For delete_footnote: the footnote ID (from read(scope='footnotes')). For set_content_control: the content control ID (from read(scope='content_controls')).",
    ),
    content_type: str = Field(
        "paragraph",
        description="Type: 'paragraph', 'heading', 'table'",
    ),
    content_data: str = Field(
        "",
        description="Content: text or JSON. For set_table_alignment: 'left'/'center'/'right'. For set_table_autofit: 'true'/'false'. For set_table_fixed_layout: JSON array of widths. For set_row_height: JSON {height, rule}. For set_cell_width: width in inches. For set_cell_vertical_alignment: 'top'/'center'/'bottom'. For set_cell_borders: JSON {top?, bottom?, left?, right?} with values as 'style:size:color' (e.g., 'single:24:000000'). For set_cell_shading: hex color (e.g., 'FF0000'). For set_header_row: 'true'/'false' (mark row as header). For set_columns: JSON {num_columns, spacing_inches?, separator?}. For set_line_numbering: JSON {enabled?, restart?, start?, count_by?, distance_inches?}. For set_custom_property: JSON {name, value, type?} where type is 'string'/'int'/'bool'/'datetime'/'float' (default 'string'). For delete_custom_property: property name. For create_style: JSON {name, style_type?, base_style?} where style_type is 'paragraph'/'character'/'table' (default 'paragraph'), base_style is style to inherit from (default 'Normal'). For add_tab_stop: JSON {position, alignment, leader}. For insert_field: field code (PAGE, NUMPAGES, DATE, TIME). For insert_page_x_of_y: 'header' or 'footer'. For set_list_level: level 0-8. For restart_numbering: start value (default 1). For insert_toc: JSON {position: 'before'/'after', heading_levels: '1-3'}. For add_footnote: JSON {text, note_type?, position?} where note_type is 'footnote'/'endnote' (default 'footnote'), position is 'after'/'before' (default 'after'). For delete_footnote: JSON {note_type?} where note_type is 'footnote'/'endnote' (default 'footnote'). For set_content_control: the new value - for dropdown must match one of the options, for checkbox use 'true'/'false', for date use ISO format. For add_hyperlink: JSON {text, address?, fragment?} where text is visible link text, address is URL (external link), fragment is bookmark name (internal link) or URL anchor.",
    ),
    style_name: str = Field(
        "", description="Apply Word style: 'Heading 1', 'Normal', etc."
    ),
    formatting: str = Field(
        "",
        description='Direct formatting JSON. Text/Run: {"bold": true, "color": "FF0000", "font_size": 14, "highlight_color": "yellow", "strike": true, "subscript": true, "superscript": true, "style": "Strong"}. Character styles: "Strong", "Emphasis", "Hyperlink", etc. Paragraph: {"left_indent": 0.5, "right_indent": 0.5, "first_line_indent": 0.5, "space_before": 12, "space_after": 12, "line_spacing": 1.5, "keep_with_next": true, "page_break_before": true} (indents in inches, spacing in points, line_spacing < 5 is multiplier). Margins: {"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25} in inches. Images: {"width": 4} or {"height": 3} in inches.',
    ),
    heading_level: int = Field(
        1, description="Heading level 1-9 (only for content_type='heading')"
    ),
    row: int = Field(0, description="Row number (0-based, for edit_cell operation)"),
    col: int = Field(0, description="Column number (0-based, for edit_cell operation)"),
    run_index: int = Field(
        -1,
        description="Run index (0-based, for edit_run operation). Use read() with scope='runs' to find indices.",
    ),
    author: str = Field("", description="Comment author name (for add_comment)"),
    initials: str = Field("", description="Comment author initials (for add_comment)"),
    section_index: int = Field(
        0,
        description="Section index (0-based, for set_header/set_footer/set_margins/set_orientation). Use read() with scope='page_setup' to see sections.",
    ),
) -> EditResult:
    """Edit Word document."""
    from docx import Document
    from docx.text.paragraph import Paragraph

    # Create operation is special - creates new document
    if operation == "create":
        doc = Document()
        obj = word_ops.create_element(
            doc, content_type, content_data, style_name, heading_level
        )
        level = heading_level if content_type == "heading" else 0
        element_id = word_ops.get_element_id(doc, obj, level)
        doc.save(file_path)
        return EditResult(
            success=True, element_id=element_id, message=f"Created: {file_path}"
        )

    # All other operations work on existing document
    doc = Document(file_path)
    element_id = target_id
    comment_id = None
    message = f"Completed {operation}"

    if operation in ("insert_before", "insert_after"):
        t = word_ops.resolve_target(doc, target_id)
        position = "before" if operation == "insert_before" else "after"
        obj = word_ops.insert_content_relative(
            doc,
            t.leaf_el,
            position,
            content_type,
            content_data,
            style_name,
            heading_level,
        )
        level = heading_level if content_type == "heading" else 0
        element_id = word_ops.get_element_id(doc, obj, level)
        message = f"Inserted {content_type} {position} {target_id}"

    elif operation == "append":
        obj = word_ops.create_element(
            doc, content_type, content_data, style_name, heading_level
        )
        level = heading_level if content_type == "heading" else 0
        element_id = word_ops.get_element_id(doc, obj, level)
        message = f"Appended {content_type} to document"

    elif operation == "delete":
        t = word_ops.resolve_target(doc, target_id)
        word_ops.delete_block(t.base_obj)
        message = f"Deleted block {target_id}"

    elif operation == "replace":
        t = word_ops.resolve_target(doc, target_id)
        if t.leaf_kind.startswith("heading") or t.leaf_kind == "paragraph":
            t.leaf_obj.text = content_data
            element_id = _recalc_block_id(doc, t)
        elif t.leaf_kind == "table":
            table_data = json.loads(content_data)
            new_tbl_el = word_ops.replace_table(t.leaf_el, table_data)  # Takes element
            table_content = word_ops.table_content_for_hash(new_tbl_el)
            occurrence = word_ops.count_occurrence(
                doc,
                "table",
                table_content,
                new_tbl_el,  # Element, not wrapper
            )
            element_id = word_ops.make_block_id("table", table_content, occurrence)
        else:
            raise ValueError(f"Unsupported leaf_kind: {t.leaf_kind}")
        message = f"Replaced content of {target_id}"

    elif operation == "style":
        t = word_ops.resolve_target(doc, target_id)
        if style_name:
            t.leaf_obj.style = style_name
        if t.base_kind == "table":
            element_id = _recalc_table_id(doc, t)
        else:
            if formatting:
                fmt = json.loads(formatting)
                word_ops.apply_paragraph_formatting(t.leaf_obj, fmt)
            element_id = _recalc_block_id(doc, t)
        message = f"Applied style to {target_id}"

    elif operation == "edit_cell":
        t = word_ops.resolve_target(doc, target_id)
        word_ops.replace_table_cell(t.base_el, row, col, content_data)  # Takes element
        element_id = _recalc_table_id(doc, t)
        message = f"Updated cell r{row}c{col}"

    elif operation == "edit_run":
        t = word_ops.resolve_target(doc, target_id)
        if content_data:
            word_ops.edit_run_text(t.leaf_obj, run_index, content_data)
        if formatting:
            fmt = json.loads(formatting)
            word_ops.edit_run_formatting(t.leaf_obj, run_index, fmt)
        element_id = _recalc_block_id(doc, t)
        message = f"Updated run {run_index}"

    elif operation == "edit_style":
        fmt = json.loads(formatting)
        word_ops.edit_style(doc, target_id, fmt)
        message = f"Modified style: {target_id}"

    elif operation == "add_comment":
        t = word_ops.resolve_target(doc, target_id)
        comment_id = word_ops.add_comment_to_block(
            doc, t.leaf_obj, content_data, author, initials
        )
        message = f"Added comment {comment_id} to {target_id}"

    elif operation == "reply_comment":
        parent_id = int(target_id)  # target_id is parent comment ID
        comment_id = word_ops.reply_to_comment(
            doc, parent_id, content_data, author, initials
        )
        message = f"Added reply {comment_id} to comment {parent_id}"

    elif operation == "resolve_comment":
        comment_id = int(target_id)
        word_ops.resolve_comment(doc, comment_id)
        message = f"Resolved comment {comment_id}"

    elif operation == "unresolve_comment":
        comment_id = int(target_id)
        word_ops.unresolve_comment(doc, comment_id)
        message = f"Unresolved comment {comment_id}"

    elif operation in _HF_SET_OPS:
        location = _HF_SET_OPS[operation]
        word_ops.set_header_footer_text(doc, section_index, content_data, location)
        message = f"Set {location.replace('_', ' ')} for section {section_index}"

    elif operation in _HF_APPEND_OPS:
        location = _HF_APPEND_OPS[operation]
        element_id = word_ops.append_to_header_footer(
            doc, section_index, content_type, content_data, location
        )
        message = f"Appended {content_type} to {location} of section {section_index}"

    elif operation in _HF_CLEAR_OPS:
        location = _HF_CLEAR_OPS[operation]
        word_ops.clear_header_footer(doc, section_index, location)
        message = f"Cleared {location} for section {section_index}"

    elif operation == "set_margins":
        margins = json.loads(formatting)
        word_ops.set_page_margins(
            doc,
            section_index,
            top=margins["top"],
            bottom=margins["bottom"],
            left=margins["left"],
            right=margins["right"],
        )
        message = f"Set margins for section {section_index}"

    elif operation == "set_orientation":
        word_ops.set_page_orientation(doc, section_index, content_data)
        message = f"Set orientation to {content_data} for section {section_index}"

    elif operation == "set_columns":
        col_data = json.loads(content_data)
        word_ops.set_section_columns(
            doc,
            section_index,
            int(col_data["num_columns"]),
            float(col_data.get("spacing_inches", 0.5)),
            col_data.get("separator", False),
        )
        message = f"Set {col_data['num_columns']} columns for section {section_index}"

    elif operation == "set_line_numbering":
        ln_data = json.loads(content_data)
        word_ops.set_line_numbering(
            doc,
            section_index,
            enabled=ln_data.get("enabled", True),
            restart=ln_data.get("restart", "newPage"),
            start=int(ln_data.get("start", 1)),
            count_by=int(ln_data.get("count_by", 1)),
            distance_inches=float(ln_data.get("distance_inches", 0.5)),
        )
        enabled = ln_data.get("enabled", True)
        action = "Enabled" if enabled else "Disabled"
        message = f"{action} line numbering for section {section_index}"

    elif operation == "set_custom_property":
        prop_data = json.loads(content_data)
        word_ops.set_custom_property(
            doc,
            name=prop_data["name"],
            value=prop_data["value"],
            prop_type=prop_data.get("type", "string"),
        )
        message = f"Set custom property '{prop_data['name']}'"

    elif operation == "delete_custom_property":
        prop_name = content_data
        deleted = word_ops.delete_custom_property(doc, prop_name)
        if deleted:
            message = f"Deleted custom property '{prop_name}'"
        else:
            message = f"Custom property '{prop_name}' not found"

    elif operation == "create_style":
        style_data = json.loads(content_data)
        formatting_dict = json.loads(formatting) if formatting else None
        style_id = word_ops.create_style(
            doc,
            name=style_data["name"],
            style_type=style_data.get("style_type", "paragraph"),
            base_style=style_data.get("base_style", "Normal"),
            formatting=formatting_dict,
        )
        element_id = style_id
        message = f"Created style '{style_data['name']}'"

    elif operation == "delete_style":
        deleted = word_ops.delete_style(doc, target_id)
        if deleted:
            message = f"Deleted style '{target_id}'"
        else:
            message = f"Style '{target_id}' not found or is builtin"

    elif operation == "insert_image":
        fmt = json.loads(formatting) if formatting else {}
        element_id = word_ops.insert_image(
            doc,
            content_data,
            target_id,
            "after",
            width_inches=float(fmt.get("width", 0)),
            height_inches=float(fmt.get("height", 0)),
        )
        message = "Inserted image"

    elif operation == "delete_image":
        word_ops.delete_image(doc, target_id)
        message = f"Deleted {target_id}"

    elif operation == "insert_floating_image":
        # content_data = image path
        # formatting JSON: position_h, position_v, relative_h, relative_v, wrap_type,
        #                  width, height, behind_doc
        fmt = json.loads(formatting) if formatting else {}
        element_id = word_ops.insert_floating_image(
            doc,
            content_data,
            target_id,
            position_h=float(fmt.get("position_h", 0)),
            position_v=float(fmt.get("position_v", 0)),
            relative_h=fmt.get("relative_h", "column"),
            relative_v=fmt.get("relative_v", "paragraph"),
            wrap_type=fmt.get("wrap_type", "square"),
            width_inches=float(fmt.get("width", 0)),
            height_inches=float(fmt.get("height", 0)),
            behind_doc=bool(fmt.get("behind_doc", False)),
        )
        message = "Inserted floating image"

    elif operation == "add_row":
        t = word_ops.resolve_target(doc, target_id)
        data = json.loads(content_data) if content_data else None
        row_idx = word_ops.add_table_row(t.base_el, data)  # Takes element
        element_id = _recalc_table_id(doc, t)
        message = f"Added row {row_idx}"

    elif operation == "add_column":
        t = word_ops.resolve_target(doc, target_id)
        fmt = json.loads(formatting) if formatting else {}
        width = float(fmt.get("width", 1.0))
        data = json.loads(content_data) if content_data else None
        col_idx = word_ops.add_table_column(t.base_el, width, data)  # Takes element
        element_id = _recalc_table_id(doc, t)
        message = f"Added column {col_idx}"

    elif operation == "delete_row":
        t = word_ops.resolve_target(doc, target_id)
        word_ops.delete_table_row(t.base_el, row)  # Takes element
        element_id = _recalc_table_id(doc, t)
        message = f"Deleted row {row}"

    elif operation == "delete_column":
        t = word_ops.resolve_target(doc, target_id)
        word_ops.delete_table_column(t.base_el, col)  # Takes element
        element_id = _recalc_table_id(doc, t)
        message = f"Deleted column {col}"

    elif operation == "add_page_break":
        p = word_ops.add_page_break(doc)
        text = p.text or ""
        occurrence = word_ops.count_occurrence(doc, "paragraph", text, p._element)
        element_id = word_ops.make_block_id("paragraph", text, occurrence)
        message = "Added page break"

    elif operation == "add_break":
        t = word_ops.resolve_target(doc, target_id)
        break_type = content_data or "page"
        p = word_ops.add_break_after(doc, t.leaf_el, break_type)
        text = p.text or ""
        occurrence = word_ops.count_occurrence(doc, "paragraph", text, p._element)
        element_id = word_ops.make_block_id("paragraph", text, occurrence)
        message = f"Added {break_type} break after {target_id}"

    elif operation == "set_meta":
        meta_data = json.loads(content_data) if content_data else {}
        word_ops.set_document_meta(
            doc,
            title=meta_data.get("title"),
            author=meta_data.get("author"),
            subject=meta_data.get("subject"),
            keywords=meta_data.get("keywords"),
            category=meta_data.get("category"),
        )
        message = "Updated document metadata"

    elif operation == "add_section":
        start_type = content_data or "new_page"
        section_idx = word_ops.add_section(doc, start_type)
        message = f"Added section {section_idx} ({start_type})"

    elif operation == "merge_cells":
        target = word_ops.resolve_target(doc, target_id)

        # Accept JSON with all coords, or use row/col params as fallback for start
        if not content_data:
            raise ValueError(
                "merge_cells requires content_data JSON with end_row, end_col"
            )
        try:
            merge_data = json.loads(content_data)
        except json.JSONDecodeError:
            raise ValueError(
                "merge_cells content_data must be valid JSON with end_row, end_col"
            )

        # Accept both naming conventions for start coordinates
        start_row = merge_data.get("start_row", merge_data.get("row", row))
        start_col = merge_data.get("start_col", merge_data.get("col", col))

        # end_row and end_col are required
        end_row = merge_data.get("end_row")
        end_col = merge_data.get("end_col")
        if end_row is None or end_col is None:
            raise ValueError("merge_cells requires end_row and end_col in content_data")

        word_ops.merge_cells(
            target.base_el, start_row, start_col, end_row, end_col
        )  # Takes element
        message = (
            f"Merged cells from ({start_row},{start_col}) to ({end_row},{end_col})"
        )

    elif operation == "set_table_alignment":
        target = word_ops.resolve_target(doc, target_id)
        word_ops.set_table_alignment(target.base_el, content_data)  # Takes element
        message = f"Set table alignment to {content_data}"

    elif operation == "set_table_autofit":
        target = word_ops.resolve_target(doc, target_id)
        target.base_obj.autofit = json.loads(content_data.lower())  # Needs wrapper
        message = f"Set table autofit to {target.base_obj.autofit}"

    elif operation == "set_table_fixed_layout":
        target = word_ops.resolve_target(doc, target_id)
        widths = json.loads(content_data)
        word_ops.set_table_fixed_layout(target.base_el, widths)  # Takes element
        message = f"Set table fixed layout with {len(widths)} columns"

    elif operation == "set_row_height":
        target = word_ops.resolve_target(doc, target_id)
        height_data = json.loads(content_data)
        word_ops.set_row_height(
            target.base_el,  # Takes element
            row,
            height_data["height"],
            height_data.get("rule", "at_least"),
        )
        message = f"Set row {row} height to {height_data['height']} inches"

    elif operation == "set_cell_width":
        target = word_ops.resolve_target(doc, target_id)
        word_ops.set_cell_width(
            target.base_el, row, col, float(content_data)
        )  # Takes element
        message = f"Set cell ({row},{col}) width to {content_data} inches"

    elif operation == "set_cell_vertical_alignment":
        target = word_ops.resolve_target(doc, target_id)
        word_ops.set_cell_vertical_alignment(
            target.base_el, row, col, content_data
        )  # Takes element
        message = f"Set cell ({row},{col}) vertical alignment to {content_data}"

    elif operation == "set_cell_borders":
        target = word_ops.resolve_target(doc, target_id)
        border_data = json.loads(content_data)
        word_ops.set_cell_borders(
            target.base_el,  # Takes element
            row,
            col,
            top=border_data.get("top"),
            bottom=border_data.get("bottom"),
            left=border_data.get("left"),
            right=border_data.get("right"),
        )
        message = f"Set borders on cell ({row},{col})"

    elif operation == "set_cell_shading":
        target = word_ops.resolve_target(doc, target_id)
        word_ops.set_cell_shading(
            target.base_el, row, col, content_data
        )  # Takes element
        message = f"Set shading on cell ({row},{col}) to {content_data}"

    elif operation == "set_header_row":
        target = word_ops.resolve_target(doc, target_id)
        is_header = content_data.lower() in ("true", "1", "yes")
        word_ops.set_header_row(target.base_el, row, is_header)  # Takes element
        action = "marked as" if is_header else "unmarked as"
        message = f"Row {row} {action} header row"

    elif operation == "add_tab_stop":
        target = word_ops.resolve_target(doc, target_id)
        para = _get_target_paragraph(target)
        tab_data = json.loads(content_data)
        word_ops.add_tab_stop(
            para,
            float(tab_data["position"]),
            tab_data.get("alignment", "left"),
            tab_data.get("leader", "spaces"),
        )
        message = f"Added tab stop at {tab_data['position']} inches ({tab_data.get('alignment', 'left')}, {tab_data.get('leader', 'spaces')})"

    elif operation == "clear_tab_stops":
        target = word_ops.resolve_target(doc, target_id)
        para = _get_target_paragraph(target)
        para.paragraph_format.tab_stops.clear_all()
        message = "Cleared all tab stops"

    elif operation == "insert_field":
        target = word_ops.resolve_target(doc, target_id)
        para = _get_target_paragraph(target)
        field_code = content_data.strip().upper()
        word_ops.insert_field(para, field_code)
        message = f"Inserted {field_code} field"

    elif operation == "insert_page_x_of_y":
        location = content_data.strip().lower() or "footer"
        word_ops.insert_page_x_of_y(doc, section_index, location)
        message = f"Inserted 'Page X of Y' in {location} of section {section_index}"

    elif operation == "accept_change":
        word_ops.accept_change(doc, target_id)
        message = f"Accepted change {target_id}"

    elif operation == "reject_change":
        word_ops.reject_change(doc, target_id)
        message = f"Rejected change {target_id}"

    elif operation == "accept_all_changes":
        count = word_ops.accept_all_changes(doc)
        message = f"Accepted {count} changes"

    elif operation == "reject_all_changes":
        count = word_ops.reject_all_changes(doc)
        message = f"Rejected {count} changes"

    elif operation == "set_list_level":
        paragraph = word_ops.find_paragraph_by_id(doc, target_id)
        if paragraph is None:
            raise ValueError(f"Paragraph not found: {target_id}")
        level = int(content_data) if content_data else 0
        if level < 0 or level > 8:
            raise ValueError("List level must be 0-8")
        word_ops.set_list_level(doc, paragraph, level)
        message = f"Set list level to {level}"

    elif operation == "promote_list":
        paragraph = word_ops.find_paragraph_by_id(doc, target_id)
        if paragraph is None:
            raise ValueError(f"Paragraph not found: {target_id}")
        word_ops.promote_list_item(doc, paragraph)
        message = "Promoted list item"

    elif operation == "demote_list":
        paragraph = word_ops.find_paragraph_by_id(doc, target_id)
        if paragraph is None:
            raise ValueError(f"Paragraph not found: {target_id}")
        word_ops.demote_list_item(doc, paragraph)
        message = "Demoted list item"

    elif operation == "restart_numbering":
        paragraph = word_ops.find_paragraph_by_id(doc, target_id)
        if paragraph is None:
            raise ValueError(f"Paragraph not found: {target_id}")
        start_value = int(content_data) if content_data else 1
        word_ops.restart_numbering(doc, paragraph, start_value)
        message = f"Restarted numbering at {start_value}"

    elif operation == "remove_list":
        paragraph = word_ops.find_paragraph_by_id(doc, target_id)
        if paragraph is None:
            raise ValueError(f"Paragraph not found: {target_id}")
        word_ops.remove_list_formatting(doc, paragraph)
        message = "Removed list formatting"

    elif operation == "edit_text_box":
        # target_id is the text box ID, row is the paragraph index, content_data is new text
        if not target_id:
            raise ValueError("target_id (text box ID) required for edit_text_box")
        word_ops.edit_text_box_text(doc, target_id, row, content_data)
        message = f"Edited text box {target_id} paragraph {row}"

    elif operation == "add_bookmark":
        # target_id is the paragraph/heading ID, content_data is the bookmark name
        if not target_id:
            raise ValueError("target_id (paragraph ID) required for add_bookmark")
        if not content_data:
            raise ValueError("content_data (bookmark name) required for add_bookmark")
        target = word_ops.resolve_target(doc, target_id)
        if not isinstance(target.leaf_obj, Paragraph):
            raise ValueError(f"Target must be a paragraph or heading: {target_id}")
        bm_id = word_ops.add_bookmark(doc, content_data, target.leaf_obj)
        message = f"Added bookmark '{content_data}' with ID {bm_id}"

    elif operation == "add_hyperlink":
        # target_id is paragraph/cell ID, content_data is JSON: {text, address?, fragment?}
        if not target_id:
            raise ValueError("target_id required for add_hyperlink")
        if not content_data:
            raise ValueError("content_data (JSON with text, address/fragment) required")
        try:
            link_data = json.loads(content_data)
        except json.JSONDecodeError:
            raise ValueError(
                "content_data must be valid JSON: {text, address?, fragment?}"
            )
        text = link_data.get("text", "")
        address = link_data.get("address", "")
        fragment = link_data.get("fragment", "")
        target = word_ops.resolve_target(doc, target_id)
        para = _get_target_paragraph(target)
        word_ops.add_hyperlink(para, text, address, fragment)
        # Recalc element_id - for cells, return table ID; for paragraphs, recalc block ID
        if target.leaf_kind == "cell":
            element_id = _recalc_table_id(doc, target)
        else:
            element_id = _recalc_block_id(doc, target)
        message = f"Added hyperlink '{text}' to {address or '#' + fragment}"

    elif operation == "insert_cross_ref":
        # target_id is the paragraph/heading ID, content_data is bookmark name, style_name is ref_type
        if not target_id:
            raise ValueError("target_id (paragraph ID) required for insert_cross_ref")
        if not content_data:
            raise ValueError(
                "content_data (bookmark name) required for insert_cross_ref"
            )
        target = word_ops.resolve_target(doc, target_id)
        if not isinstance(target.leaf_obj, Paragraph):
            raise ValueError(f"Target must be a paragraph or heading: {target_id}")
        ref_type = style_name if style_name else "text"
        word_ops.insert_cross_reference(target.leaf_obj, content_data, ref_type)
        message = f"Inserted cross-reference to '{content_data}' ({ref_type})"

    elif operation == "insert_caption":
        # target_id is the block to caption
        # content_data can be JSON {label, text, position} or plain string (caption text)
        if not target_id:
            raise ValueError("target_id (block ID) required for insert_caption")

        # Accept plain string OR JSON dict
        try:
            caption_data = json.loads(content_data) if content_data else {}
            if not isinstance(caption_data, dict):
                # JSON parsed to non-dict (e.g., string, list, number)
                raise TypeError("Expected dict")
            label = caption_data.get("label", "Figure")
            caption_text = caption_data.get("text", "")
            position = caption_data.get("position", "below")
        except (json.JSONDecodeError, TypeError):
            # Plain string: use as caption text with default label
            label = "Figure"
            caption_text = content_data if content_data else ""
            position = "below"

        # Validate position
        if position not in ("above", "below"):
            raise ValueError(
                f"Invalid position '{position}'. Valid: ['above', 'below']"
            )

        element_id = word_ops.insert_caption(
            doc, target_id, label, caption_text, position
        )
        message = f"Inserted {label} caption {position} {target_id}"

    elif operation == "insert_toc":
        # target_id is the block to insert before/after, content_data is JSON {position, heading_levels}
        if not target_id:
            raise ValueError("target_id (block ID) required for insert_toc")
        toc_data = json.loads(content_data) if content_data else {}
        position = toc_data.get("position", "before")
        heading_levels = toc_data.get("heading_levels", "1-3")
        element_id = word_ops.insert_toc(doc, target_id, position, heading_levels)
        message = f"Inserted TOC {position} {target_id}"

    elif operation == "update_toc":
        word_ops.update_toc_field(doc)
        message = "Set TOC dirty flag for update on open"

    elif operation == "add_footnote":
        # target_id is the paragraph/block ID, content_data is JSON {text, note_type?, position?}
        if not target_id:
            raise ValueError("target_id (paragraph ID) required for add_footnote")
        if not content_data:
            raise ValueError("content_data (JSON with text) required for add_footnote")
        fn_data = json.loads(content_data)
        fn_text = fn_data.get("text", "")
        note_type = fn_data.get("note_type", "footnote")
        position = fn_data.get("position", "after")
        # Footnotes require direct ZIP manipulation, so we save first and operate on file
        doc.save(file_path)
        fn_id = word_ops.add_footnote(
            file_path, target_id, fn_text, note_type, position
        )
        element_id = str(fn_id)
        message = f"Added {note_type} {fn_id} to {target_id}"
        # Return early - don't save again
        return EditResult(success=True, element_id=element_id, message=message)

    elif operation == "delete_footnote":
        # target_id is the footnote/endnote ID, content_data is optional JSON {note_type?}
        if not target_id:
            raise ValueError("target_id (footnote ID) required for delete_footnote")
        fn_data = json.loads(content_data) if content_data else {}
        note_type = fn_data.get("note_type", "footnote")
        note_id = int(target_id)
        # Footnotes require direct ZIP manipulation, so we save first and operate on file
        doc.save(file_path)
        word_ops.delete_footnote(file_path, note_id, note_type)
        message = f"Deleted {note_type} {note_id}"
        # Return early - don't save again
        return EditResult(success=True, element_id=target_id, message=message)

    elif operation == "set_content_control":
        # target_id is the content control ID, content_data is the new value
        if not target_id:
            raise ValueError(
                "target_id (content control ID) required for set_content_control"
            )
        if not content_data:
            raise ValueError(
                "content_data (new value) required for set_content_control"
            )
        sdt_id = int(target_id)
        word_ops.set_content_control_value(doc, sdt_id, content_data)
        message = f"Set content control {sdt_id} to '{content_data}'"

    else:
        raise ValueError(f"Unknown operation: {operation}")

    doc.save(file_path)
    return EditResult(
        success=True, element_id=element_id, comment_id=comment_id, message=message
    )


if __name__ == "__main__":
    mcp.run()
