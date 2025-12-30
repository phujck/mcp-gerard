"""Python-docx wrapper functions for Word document manipulation."""

import hashlib
import json
import re
from collections.abc import Iterator
from dataclasses import dataclass

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_COLOR_INDEX,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap as oxml_nsmap
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl, CT_Tc
from docx.oxml.text.paragraph import CT_P
from docx.shared import Emu, Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from lxml import etree

from mcp_handley_lab.word.models import (
    Block,
    CellInfo,
    CommentInfo,
    CustomPropertyInfo,
    DocumentMeta,
    HeaderFooterInfo,
    HyperlinkInfo,
    ImageInfo,
    LineNumberingInfo,
    PageSetupInfo,
    ParagraphFormatInfo,
    RowInfo,
    RunInfo,
    StyleFormatInfo,
    StyleInfo,
    TableLayoutInfo,
    TabStopInfo,
)

_HEADING_RE = re.compile(r"^Heading ([1-9])$")
_ID_RE = re.compile(r"^(paragraph|heading[1-9]|table)_([0-9a-f]{8})_(\d+)$")
_IMAGE_ID_RE = re.compile(r"^image_([0-9a-f]{8})_(\d+)$")
_EMU_PER_INCH = 914400

# Highlight color mapping (WD_COLOR_INDEX enum)
_HIGHLIGHT_MAP = {
    "yellow": WD_COLOR_INDEX.YELLOW,
    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "cyan": WD_COLOR_INDEX.TURQUOISE,
    "pink": WD_COLOR_INDEX.PINK,
    "blue": WD_COLOR_INDEX.BLUE,
    "red": WD_COLOR_INDEX.RED,
    "dark_blue": WD_COLOR_INDEX.DARK_BLUE,
    "dark_red": WD_COLOR_INDEX.DARK_RED,
    "dark_yellow": WD_COLOR_INDEX.DARK_YELLOW,
    "gray": WD_COLOR_INDEX.GRAY_25,
    "dark_gray": WD_COLOR_INDEX.GRAY_50,
    "black": WD_COLOR_INDEX.BLACK,
    "white": WD_COLOR_INDEX.WHITE,
}
_HIGHLIGHT_REVERSE = {v: k for k, v in _HIGHLIGHT_MAP.items()}

# Section start type mapping
_SECTION_START_MAP = {
    "new_page": WD_SECTION.NEW_PAGE,
    "continuous": WD_SECTION.CONTINUOUS,
    "even_page": WD_SECTION.EVEN_PAGE,
    "odd_page": WD_SECTION.ODD_PAGE,
    "new_column": WD_SECTION.NEW_COLUMN,
}

# Wrap type mapping: API value <-> XML element name (without wp: prefix)
_WRAP_API_TO_XML = {
    "square": "wrapSquare",
    "tight": "wrapTight",
    "through": "wrapThrough",
    "top_and_bottom": "wrapTopAndBottom",
    "none": "wrapNone",
}
_WRAP_XML_TO_API = {v: k for k, v in _WRAP_API_TO_XML.items()}

# Extended namespace map for SDT content controls (Word 2010/2012 extensions)
_SDT_NSMAP = {
    **oxml_nsmap,
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
}

# Hierarchical path segment patterns (0-based indices)
_CELL_RE = re.compile(r"^r(\d+)c(\d+)$")
_PARA_RE = re.compile(r"^p(\d+)$")
_TABLE_RE = re.compile(r"^tbl(\d+)$")

# Run formatting attribute mappings (attr_path, transform)
_RUN_ATTRS = {
    "bold": "bold",
    "italic": "italic",
    "underline": "underline",
    "strike": "font.strike",
    "double_strike": "font.double_strike",
    "subscript": "font.subscript",
    "superscript": "font.superscript",
    "all_caps": "font.all_caps",
    "small_caps": "font.small_caps",
    "hidden": "font.hidden",
    "emboss": "font.emboss",
    "imprint": "font.imprint",
    "outline": "font.outline",
    "shadow": "font.shadow",
    "font_name": "font.name",
}


def populate_table(table: Table, data: list[list]) -> None:
    """Populate table cells from 2D list."""
    for r, row in enumerate(data):
        for c, val in enumerate(row):
            table.cell(r, c).text = str(val)


# PathSegment: (kind, indices) where kind='cell'|'para'|'tbl'
PathSegment = tuple[str, tuple[int, ...]]


@dataclass
class ResolvedTarget:
    """Result of resolving a hierarchical target ID."""

    base_id: str
    base_kind: str  # 'table' | 'paragraph' | 'heading1' etc
    base_obj: Paragraph | Table
    base_occurrence: int
    leaf_kind: str  # 'table' | 'cell' | 'paragraph'
    leaf_obj: Paragraph | Table | _Cell
    leaf_el: CT_P | CT_Tbl | CT_Tc  # XML element (Paragraph, Table, or Cell)


def content_hash(text: str) -> str:
    """8-char SHA256 of normalized text for content-addressable IDs."""
    normalized = re.sub(r"\s+", " ", text.strip().lower())
    return hashlib.sha256(normalized.encode()).hexdigest()[:8]


def make_block_id(block_type: str, text: str, occurrence: int) -> str:
    """Generate content-addressable block ID. Single source of truth."""
    return f"{block_type}_{content_hash(text)}_{occurrence}"


def iter_body_blocks(
    doc: Document,
) -> Iterator[tuple[str, Paragraph | Table, CT_P | CT_Tbl]]:
    """Yield (block_kind, block_obj, element) in true document order."""
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield ("paragraph", Paragraph(child, doc), child)
        elif isinstance(child, CT_Tbl):
            yield ("table", Table(child, doc), child)


def paragraph_kind_and_level(p: Paragraph) -> tuple[str, int]:
    """Detect if paragraph is a heading and return (kind, level).

    For content-hash IDs, kind includes level: 'heading1', 'heading2', etc.
    """
    style_name = p.style.name if p.style else "Normal"
    m = _HEADING_RE.match(style_name)
    if m:
        level = int(m.group(1))
        return (f"heading{level}", level)
    return ("paragraph", 0)


def table_content_for_hash(table: Table) -> str:
    """Get full table content as canonical string for hashing."""
    return json.dumps(
        [[c.text for c in r.cells] for r in table.rows],
        ensure_ascii=False,
        separators=(",", ":"),
    )


def create_element(
    doc: Document,
    content_type: str,
    content_data: str,
    style_name: str = "",
    heading_level: int = 0,
) -> Paragraph | Table:
    """Create a content element (paragraph, heading, or table)."""
    if content_type == "paragraph":
        return doc.add_paragraph(content_data, style_name or None)
    if content_type == "heading":
        return doc.add_heading(content_data, level=heading_level)
    if content_type == "table":
        table_data = json.loads(content_data)
        rows, cols = len(table_data), max((len(r) for r in table_data), default=1)
        tbl = doc.add_table(rows=rows, cols=cols)
        tbl.style = style_name or "Table Grid"
        populate_table(tbl, table_data)
        return tbl
    raise ValueError(f"Unknown content_type: {content_type}")


def get_element_id(
    doc: Document, obj: Paragraph | Table, heading_level: int = 0
) -> str:
    """Calculate content-addressed ID for an element."""
    if isinstance(obj, Table):
        block_type = "table"
        content = table_content_for_hash(obj)
        el = obj._tbl
    else:
        block_type, _ = paragraph_kind_and_level(obj)
        # Override block_type if heading_level specified (for newly created headings)
        if heading_level:
            block_type = f"heading{heading_level}"
        content = obj.text or ""
        el = obj._element
    occurrence = count_occurrence(doc, block_type, content, el)
    return make_block_id(block_type, content, occurrence)


def insert_content_relative(
    doc: Document,
    target_el: CT_P | CT_Tbl,
    position: str,
    content_type: str,
    content_data: str,
    style_name: str = "",
    heading_level: int = 0,
) -> Paragraph | Table:
    """Insert content relative to a target element."""
    obj = create_element(doc, content_type, content_data, style_name, heading_level)
    el = obj._tbl if isinstance(obj, Table) else obj._element
    _insert_at(target_el, el, position)
    return obj


def table_to_markdown(
    table: Table, max_chars: int = 500, max_rows: int = 20, max_cols: int = 10
) -> tuple[str, int, int]:
    """Convert table to markdown preview with truncation."""
    rows, cols = len(table.rows), len(table.columns)
    r_lim, c_lim = min(rows, max_rows), min(cols, max_cols)

    grid = [
        [
            table.cell(r, c).text.strip().replace("|", "\\|").replace("\n", "<br>")
            for c in range(c_lim)
        ]
        for r in range(r_lim)
    ]

    header = grid[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * len(header)) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in grid[1:])

    if rows > r_lim:
        lines.append(f"... ({rows - r_lim} more rows)")
    if cols > c_lim:
        lines.append(f"... ({cols - c_lim} more cols)")

    md = "\n".join(lines)
    if len(md) > max_chars:
        md = md[:max_chars] + "\n... (truncated)"
    return md, rows, cols


def get_document_meta(doc: Document) -> DocumentMeta:
    """Extract document metadata from core properties and custom properties."""
    cp = doc.core_properties
    custom_props = get_custom_properties(doc)
    return DocumentMeta(
        title=cp.title or "",
        author=cp.author or "",
        created=cp.created.isoformat() if cp.created else "",
        modified=cp.modified.isoformat() if cp.modified else "",
        revision=cp.revision or 0,
        sections=len(doc.sections),
        custom_properties=custom_props,
    )


def set_document_meta(doc: Document, **kwargs) -> None:
    """Update document core properties. Only updates non-None values."""
    cp = doc.core_properties
    for key, value in kwargs.items():
        if value is not None:
            setattr(cp, key, value)


# Custom properties functions


def get_custom_properties(doc: Document) -> list[CustomPropertyInfo]:
    """Get all custom document properties from docProps/custom.xml."""
    props = []
    # Custom properties are in docProps/custom.xml
    for part in doc.part.package.iter_parts():
        if part.partname == "/docProps/custom.xml":
            root = etree.fromstring(part.blob)
            ns_custom = "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"

            for prop in root.findall(f"{{{ns_custom}}}property"):
                name = prop.get("name", "")
                # Find value element - could be vt:lpwstr, vt:i4, vt:bool, vt:filetime, etc.
                value = ""
                prop_type = "string"
                for child in prop:
                    local_tag = (
                        child.tag.split("}")[-1] if "}" in child.tag else child.tag
                    )
                    if local_tag == "lpwstr":
                        value = child.text or ""
                        prop_type = "string"
                    elif local_tag == "i4":
                        value = child.text or "0"
                        prop_type = "int"
                    elif local_tag == "bool":
                        value = child.text or "false"
                        prop_type = "bool"
                    elif local_tag == "filetime":
                        value = child.text or ""
                        prop_type = "datetime"
                    elif local_tag == "r8":
                        value = child.text or "0.0"
                        prop_type = "float"
                    break
                props.append(CustomPropertyInfo(name=name, value=value, type=prop_type))
            break
    return props


def set_custom_property(
    doc: Document, name: str, value: str, prop_type: str = "string"
) -> None:
    """Set or update a custom document property.

    Args:
        doc: The Document object
        name: Property name (must be unique)
        value: Property value as string
        prop_type: One of "string", "int", "bool", "datetime", "float"
    """
    ns_custom = (
        "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
    )
    ns_vt = "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"

    # Type to element tag mapping
    type_map = {
        "string": f"{{{ns_vt}}}lpwstr",
        "int": f"{{{ns_vt}}}i4",
        "bool": f"{{{ns_vt}}}bool",
        "datetime": f"{{{ns_vt}}}filetime",
        "float": f"{{{ns_vt}}}r8",
    }

    value_tag = type_map.get(prop_type, type_map["string"])

    # Find or create custom.xml part
    custom_part = None
    for part in doc.part.package.iter_parts():
        if part.partname == "/docProps/custom.xml":
            custom_part = part
            break

    if custom_part is None:
        # Create new custom.xml using python-docx's internal mechanisms
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part

        root = etree.Element(
            f"{{{ns_custom}}}Properties",
            nsmap={None: ns_custom, "vt": ns_vt},
        )
        # Add property element
        prop = etree.SubElement(root, f"{{{ns_custom}}}property")
        prop.set("fmtid", "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}")
        prop.set("pid", "2")
        prop.set("name", name)
        value_el = etree.SubElement(prop, value_tag)
        value_el.text = value

        # Save - need to add the part to the package
        xml_bytes = etree.tostring(root, xml_declaration=True, encoding="UTF-8")
        # Create custom properties part using the package's load method
        content_type = (
            "application/vnd.openxmlformats-officedocument.custom-properties+xml"
        )
        part_uri = PackURI("/docProps/custom.xml")

        # Create the part and add to package properly
        custom_part = Part.load(part_uri, content_type, xml_bytes, doc.part.package)

        # Add relationship from package to this part
        doc.part.package.relate_to(
            custom_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties",
        )
    else:
        # Update existing custom.xml
        root = etree.fromstring(custom_part.blob)

        # Find property with this name
        existing = None
        for prop in root.findall(f"{{{ns_custom}}}property"):
            if prop.get("name") == name:
                existing = prop
                break

        if existing is not None:
            # Update existing property
            for child in list(existing):
                existing.remove(child)
            value_el = etree.SubElement(existing, value_tag)
            value_el.text = value
        else:
            # Add new property - find next pid
            max_pid = 1
            for prop in root.findall(f"{{{ns_custom}}}property"):
                pid = int(prop.get("pid", "1"))
                if pid > max_pid:
                    max_pid = pid
            new_pid = max_pid + 1

            prop = etree.SubElement(root, f"{{{ns_custom}}}property")
            prop.set("fmtid", "{D5CDD505-2E9C-101B-9397-08002B2CF9AE}")
            prop.set("pid", str(new_pid))
            prop.set("name", name)
            value_el = etree.SubElement(prop, value_tag)
            value_el.text = value

        # Save updated XML back to part
        custom_part._blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8")


def delete_custom_property(doc: Document, name: str) -> bool:
    """Delete a custom document property.

    Args:
        doc: The Document object
        name: Property name to delete

    Returns:
        True if property was found and deleted, False if not found
    """
    ns_custom = (
        "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
    )

    for part in doc.part.package.iter_parts():
        if part.partname == "/docProps/custom.xml":
            root = etree.fromstring(part.blob)

            # Find property with this name
            for prop in root.findall(f"{{{ns_custom}}}property"):
                if prop.get("name") == name:
                    root.remove(prop)
                    part._blob = etree.tostring(
                        root, xml_declaration=True, encoding="UTF-8"
                    )
                    return True
            break
    return False


def build_blocks(
    doc: Document,
    offset: int = 0,
    limit: int = 50,
    heading_only: bool = False,
    search_query: str = "",
) -> tuple[list[Block], int]:
    """Build list of Block objects from document body.

    Uses content-hash IDs: {type}_{hash}_{occurrence}
    """
    blocks = []
    matched = 0
    query_lower = search_query.lower() if search_query else ""
    hash_counts: dict[str, int] = {}  # {type_hash: count} for occurrence tracking

    for kind, obj, _el in iter_body_blocks(doc):
        if kind == "paragraph":
            block_type, level = paragraph_kind_and_level(obj)
            text = obj.text or ""
            style = obj.style.name if obj.style else "Normal"

            # Track occurrence for this type+hash combination
            hash_key = f"{block_type}_{content_hash(text)}"
            occurrence = hash_counts.get(hash_key, 0)
            hash_counts[hash_key] = occurrence + 1

            if heading_only and not block_type.startswith("heading"):
                continue
            if query_lower and query_lower not in text.lower():
                continue

            if matched >= offset and len(blocks) < limit:
                blocks.append(
                    Block(
                        id=make_block_id(block_type, text, occurrence),
                        type=block_type,
                        text=text,
                        style=style,
                        level=level,
                    )
                )
            matched += 1

        elif kind == "table":
            md, rows, cols = table_to_markdown(obj)
            style = obj.style.name if obj.style else ""
            # Use full content for hash (not truncated markdown)
            table_content = table_content_for_hash(obj)

            # Track occurrence for this type+hash combination
            hash_key = f"table_{content_hash(table_content)}"
            occurrence = hash_counts.get(hash_key, 0)
            hash_counts[hash_key] = occurrence + 1

            if heading_only:
                continue
            # Search full table content, not truncated markdown preview
            if query_lower and query_lower not in table_content.lower():
                continue

            if matched >= offset and len(blocks) < limit:
                blocks.append(
                    Block(
                        id=make_block_id("table", table_content, occurrence),
                        type="table",
                        text=md,  # Keep markdown for display
                        style=style,
                        rows=rows,
                        cols=cols,
                    )
                )
            matched += 1

    return blocks, matched


def parse_target_id(target_id: str) -> tuple[str, list[PathSegment]]:
    """Parse target_id into (base_id, path_segments).

    Supports hierarchical IDs: table_abc12345_0#r0c1/p0
    Returns (base_id, []) for simple IDs without path.
    """
    base_id, *rest = target_id.split("#", 1)
    if not rest:
        return target_id, []

    segments: list[PathSegment] = []
    for part in rest[0].split("/"):
        if m := _CELL_RE.match(part):
            segments.append(("cell", (int(m[1]), int(m[2]))))
        elif m := _PARA_RE.match(part):
            segments.append(("para", (int(m[1]),)))
        elif m := _TABLE_RE.match(part):
            segments.append(("tbl", (int(m[1]),)))
        else:
            raise ValueError(f"Invalid path segment: {part}")
    return base_id, segments


def resolve_path(
    container: Table | _Cell, segments: list[PathSegment]
) -> Paragraph | Table | _Cell:
    """Resolve path segments from container. Returns leaf object.

    Transition rules:
    - From Table -> cell (r,c) only
    - From Cell -> para (p) or tbl (tbl)
    - From nested Table -> cell (r,c)
    """
    current: Table | _Cell | Paragraph = container

    for kind, indices in segments:
        if kind == "cell":
            row, col = indices
            current = current.cell(
                row, col
            )  # AttributeError if not Table, IndexError if OOB
        elif kind == "para":
            current = current.paragraphs[indices[0]]  # AttributeError if not Cell
        elif kind == "tbl":
            current = current.tables[indices[0]]  # AttributeError if not Cell

    return current


def _resolve_base_block(
    doc: Document, base_id: str
) -> tuple[str, Paragraph | Table, CT_P | CT_Tbl, int]:
    """Resolve base block ID to (block_type, obj, element, occurrence).

    Uses content-hash IDs: {type}_{hash}_{occurrence}
    Searches all blocks for matching type+hash, then skips to Nth occurrence.
    """
    target_type, target_hash, occurrence_str = _ID_RE.match(base_id).groups()
    target_occurrence = int(occurrence_str)

    occurrence_count = 0
    for kind, obj, el in iter_body_blocks(doc):
        if kind == "table":
            if target_type == "table":
                table_content = table_content_for_hash(obj)
                if content_hash(table_content) == target_hash:
                    if occurrence_count == target_occurrence:
                        return "table", obj, el, occurrence_count
                    occurrence_count += 1
        else:
            block_type, _ = paragraph_kind_and_level(obj)
            if target_type == block_type:
                text = obj.text or ""
                if content_hash(text) == target_hash:
                    if occurrence_count == target_occurrence:
                        return block_type, obj, el, occurrence_count
                    occurrence_count += 1
    raise ValueError(f"Block not found: {base_id}")


def resolve_target(doc: Document, target_id: str) -> ResolvedTarget:
    """Resolve target_id to ResolvedTarget with base and leaf info.

    Supports hierarchical IDs: table_abc12345_0#r0c1/p0
    """
    base_id, path_segments = parse_target_id(target_id)

    # Resolve base block
    base_kind, base_obj, base_el, base_occurrence = _resolve_base_block(doc, base_id)

    if not path_segments:
        return ResolvedTarget(
            base_id=base_id,
            base_kind=base_kind,
            base_obj=base_obj,
            base_occurrence=base_occurrence,
            leaf_kind=base_kind,
            leaf_obj=base_obj,
            leaf_el=base_el,
        )

    # Resolve path within container (AttributeError if base_obj not a Table)
    leaf_obj = resolve_path(base_obj, path_segments)

    # Determine leaf kind and element based on type
    if isinstance(leaf_obj, Table):
        return ResolvedTarget(
            base_id=base_id,
            base_kind=base_kind,
            base_obj=base_obj,
            base_occurrence=base_occurrence,
            leaf_kind="table",
            leaf_obj=leaf_obj,
            leaf_el=leaf_obj._tbl,
        )
    if isinstance(leaf_obj, _Cell):
        return ResolvedTarget(
            base_id=base_id,
            base_kind=base_kind,
            base_obj=base_obj,
            base_occurrence=base_occurrence,
            leaf_kind="cell",
            leaf_obj=leaf_obj,
            leaf_el=leaf_obj._tc,
        )
    # Paragraph
    return ResolvedTarget(
        base_id=base_id,
        base_kind=base_kind,
        base_obj=base_obj,
        base_occurrence=base_occurrence,
        leaf_kind="paragraph",
        leaf_obj=leaf_obj,
        leaf_el=leaf_obj._element,
    )


def find_paragraph_by_id(doc: Document, target_id: str) -> Paragraph | None:
    """Find a paragraph by its block ID.

    Returns the Paragraph object if found and is a paragraph, None otherwise.
    """
    try:
        target = resolve_target(doc, target_id)
        if target.leaf_kind == "paragraph" and isinstance(target.leaf_obj, Paragraph):
            return target.leaf_obj
        return None
    except ValueError:
        return None


def count_occurrence(
    doc: Document, block_type: str, text: str, target_el: CT_P | CT_Tbl
) -> int:
    """Count how many blocks with same type+hash appear before target_el.

    Used after edits to compute the correct occurrence index for returned ID.
    For tables, 'text' should be the full table content (from table_content_for_hash).
    """
    target_hash = content_hash(text)
    occurrence = 0
    for kind, obj, el in iter_body_blocks(doc):
        if kind == "table":
            if block_type == "table":
                table_content = table_content_for_hash(obj)
                if content_hash(table_content) == target_hash:
                    if el is target_el:
                        return occurrence
                    occurrence += 1
        else:
            actual_type, _ = paragraph_kind_and_level(obj)
            if block_type == actual_type:
                block_text = obj.text or ""
                if content_hash(block_text) == target_hash:
                    if el is target_el:
                        return occurrence
                    occurrence += 1
    raise ValueError("target_el not found in document")


def _insert_at(target_el, new_el, position: str) -> None:
    """Insert new_el before or after target_el."""
    (target_el.addprevious if position == "before" else target_el.addnext)(new_el)


def insert_table_relative(
    doc: Document,
    target_el: CT_P | CT_Tbl,
    table_data: list[list[str]],
    position: str,
    style_name: str = "Table Grid",
) -> Table:
    """Insert table before/after target element."""
    rows, cols = len(table_data), max((len(r) for r in table_data), default=1)
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = style_name
    populate_table(tbl, table_data)
    _insert_at(target_el, tbl._tbl, position)
    return tbl


def add_page_break(doc: Document) -> Paragraph:
    """Append a page break paragraph to the document. Returns the paragraph."""

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def add_break_after(
    doc: Document, target_el: CT_P | CT_Tbl, break_type: str
) -> Paragraph:
    """Insert break after target element. break_type: 'page', 'column', 'line'."""

    break_map = {
        "page": WD_BREAK.PAGE,
        "column": WD_BREAK.COLUMN,
        "line": WD_BREAK.LINE,
    }
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(break_map[break_type])  # KeyError if invalid
    target_el.addnext(p._element)
    return p


def delete_block(obj: Paragraph | Table) -> None:
    """Delete a block from the document."""
    el = obj._element if isinstance(obj, Paragraph) else obj._tbl
    el.getparent().remove(el)


def replace_table(doc: Document, old_tbl: Table, table_data: list[list[str]]) -> Table:
    """Replace table with new data."""
    rows, cols = len(table_data), max((len(r) for r in table_data), default=1)
    new_tbl = doc.add_table(rows=rows, cols=cols)
    populate_table(new_tbl, table_data)
    old_el = old_tbl._tbl
    old_el.addprevious(new_tbl._tbl)
    old_el.getparent().remove(old_el)
    return new_tbl


_PARA_INCH_ATTRS = {"left_indent", "right_indent", "first_line_indent"}
_PARA_PT_ATTRS = {"space_before", "space_after"}
_PARA_DIRECT_ATTRS = {"keep_with_next", "page_break_before"}
_RUN_FORMAT_KEYS = set(_RUN_ATTRS) | {"style", "font_size", "color", "highlight_color"}


def apply_paragraph_formatting(p: Paragraph, fmt: dict) -> None:
    """Apply direct formatting to paragraph (affects all runs)."""
    if "alignment" in fmt:
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, fmt["alignment"].upper())

    pf = p.paragraph_format
    for key, value in fmt.items():
        if key in _PARA_INCH_ATTRS:
            setattr(pf, key, Inches(value))
        elif key in _PARA_PT_ATTRS:
            setattr(pf, key, Pt(value))
        elif key == "line_spacing":
            pf.line_spacing = value if value < 5 else Pt(value)
        elif key in _PARA_DIRECT_ATTRS:
            setattr(pf, key, value)

    for run in p.runs:
        for key, value in fmt.items():
            if key in _RUN_FORMAT_KEYS:
                _set_run_attr(run, key, value)


def _get_vmerge_val_from_tc(tc: CT_Tc) -> str | None:
    """Get vMerge value from a tc element.

    Returns:
        'restart' for merge origin, 'continue' for continuation, None for no merge.
    """

    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    v_merge = tc_pr.find(qn("w:vMerge"))
    if v_merge is None:
        return None
    # vMerge with val="restart" is origin, vMerge without val is continue
    val = v_merge.get(qn("w:val"))
    return val if val else "continue"


def _tc_at_grid_col(tr, grid_col: int):
    """Find the tc element at a given grid column, accounting for gridSpan.

    Returns the tc element or None if not found.
    """

    c = 0
    for tc in tr.findall(qn("w:tc")):
        span = tc.grid_span
        if c == grid_col:
            return tc
        c += span
        if c > grid_col:
            # Grid column falls within a span, no tc origin at this position
            return None
    return None


def _calculate_row_span_from_xml(table: Table, start_row: int, col: int) -> int:
    """Calculate vertical span by checking vMerge='continue' in subsequent rows.

    Iterates over actual XML tc elements to detect continuation cells.
    """
    span = 1
    rows = list(table.rows)
    for r in range(start_row + 1, len(rows)):
        row_el = rows[r]._tr
        tc = _tc_at_grid_col(row_el, col)
        if tc is not None:
            vmerge = _get_vmerge_val_from_tc(tc)
            if vmerge == "continue":
                span += 1
            else:
                break
        else:
            break
    return span


def _get_cell_border(tc, side: str) -> str | None:
    """Extract border info from tc element. Returns 'style:size:color' or None."""
    tcPr = tc.find(qn("w:tcPr"))
    borders = tcPr.find(qn("w:tcBorders")) if tcPr is not None else None
    border_el = borders.find(qn(f"w:{side}")) if borders is not None else None
    if border_el is None:
        return None
    style = border_el.get(qn("w:val")) or "single"
    sz = border_el.get(qn("w:sz")) or "4"
    color = border_el.get(qn("w:color")) or "auto"
    return f"{style}:{sz}:{color}"


def _get_cell_shading(tc) -> str | None:
    """Extract fill color from tc element. Returns hex color or None."""
    tcPr = tc.find(qn("w:tcPr"))
    shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
    if shd is None:
        return None
    fill = shd.get(qn("w:fill"))
    return fill.upper() if fill and fill.lower() != "auto" else None


def build_table_cells(table: Table, table_id: str = "") -> list[CellInfo]:
    """Build list of CellInfo with merge information.

    Detects:
    - Horizontal merges via grid_span property
    - Vertical merges via vMerge XML attribute

    Iterates over actual XML elements (not table.cell()) to correctly
    detect continuation cells in vertical merges.

    Args:
        table: The Table object
        table_id: Base ID of the table (for hierarchical IDs)

    Returns:
        List of CellInfo with merge info for all grid positions
    """

    valign_map = {
        WD_CELL_VERTICAL_ALIGNMENT.TOP: "top",
        WD_CELL_VERTICAL_ALIGNMENT.CENTER: "center",
        WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: "bottom",
    }

    result = []
    rows = list(table.rows)

    for r, row in enumerate(rows):
        row_el = row._tr
        tc_elements = row_el.findall(qn("w:tc"))
        c = 0
        for tc in tc_elements:
            vmerge = _get_vmerge_val_from_tc(tc)
            grid_span = tc.grid_span

            if vmerge == "continue":
                # Vertical continuation cell
                result.append(
                    CellInfo(
                        row=r,
                        col=c,
                        text="",
                        hierarchical_id=f"{table_id}#r{r}c{c}" if table_id else "",
                        is_merge_origin=False,
                        grid_span=1,
                        row_span=1,
                    )
                )
                # Add horizontal continuation entries for wide continuation cells
                for span_c in range(1, grid_span):
                    result.append(
                        CellInfo(
                            row=r,
                            col=c + span_c,
                            text="",
                            hierarchical_id=(
                                f"{table_id}#r{r}c{c + span_c}" if table_id else ""
                            ),
                            is_merge_origin=False,
                            grid_span=1,
                            row_span=1,
                        )
                    )
                c += grid_span
            else:
                # Origin cell (vmerge='restart' or None) or normal cell
                row_span = (
                    _calculate_row_span_from_xml(table, r, c)
                    if vmerge == "restart"
                    else 1
                )
                # Get text and properties from the cell
                cell = table.cell(r, c)
                width_inches = cell.width.inches if cell.width else None
                valign = valign_map.get(cell.vertical_alignment)
                # Extract border and shading from tc element
                border_top = _get_cell_border(tc, "top")
                border_bottom = _get_cell_border(tc, "bottom")
                border_left = _get_cell_border(tc, "left")
                border_right = _get_cell_border(tc, "right")
                fill_color = _get_cell_shading(tc)
                result.append(
                    CellInfo(
                        row=r,
                        col=c,
                        text=cell.text or "",
                        hierarchical_id=f"{table_id}#r{r}c{c}" if table_id else "",
                        is_merge_origin=True,
                        grid_span=grid_span,
                        row_span=row_span,
                        width_inches=width_inches,
                        vertical_alignment=valign,
                        border_top=border_top,
                        border_bottom=border_bottom,
                        border_left=border_left,
                        border_right=border_right,
                        fill_color=fill_color,
                    )
                )
                # Add continuation entries for horizontal span
                for span_c in range(1, grid_span):
                    result.append(
                        CellInfo(
                            row=r,
                            col=c + span_c,
                            text="",
                            hierarchical_id=(
                                f"{table_id}#r{r}c{c + span_c}" if table_id else ""
                            ),
                            is_merge_origin=False,
                            grid_span=1,
                            row_span=1,
                        )
                    )
                c += grid_span
    return result


def merge_cells(
    table: Table, start_row: int, start_col: int, end_row: int, end_col: int
) -> None:
    """Merge a rectangular region of cells."""
    start_cell = table.cell(start_row, start_col)
    end_cell = table.cell(end_row, end_col)
    start_cell.merge(end_cell)


def replace_table_cell(table: Table, row: int, col: int, text: str) -> None:
    """Replace text in a table cell. Row/col are 0-based."""
    table.cell(row, col).text = text


def add_table_row(table: Table, data: list[str] | None = None) -> int:
    """Add row to table. Returns new row index (0-based)."""
    row = table.add_row()
    if data:
        for i, text in enumerate(data[: len(table.columns)]):
            row.cells[i].text = text
    return len(table.rows) - 1


def add_table_column(
    table: Table, width_inches: float = 1.0, data: list[str] | None = None
) -> int:
    """Add column to table. Width required by python-docx API. Returns new col index."""

    table.add_column(Inches(width_inches))
    col_idx = len(table.columns) - 1
    if data:
        for i, text in enumerate(data[: len(table.rows)]):
            table.cell(i, col_idx).text = text
    return col_idx


def delete_table_row(table: Table, row_index: int) -> None:
    """Delete row from table (0-based index)."""
    row = table.rows[row_index]
    row._element.getparent().remove(row._element)


def delete_table_column(table: Table, col_index: int) -> None:
    """Delete column from table (0-based index). Removes grid definition and cells."""
    # 1. Remove the grid column definition (required for valid Word XML)
    tbl_grid = table._tbl.tblGrid
    if col_index < len(tbl_grid.gridCol_lst):
        grid_col = tbl_grid.gridCol_lst[col_index]
        grid_col.getparent().remove(grid_col)

    # 2. Remove the cell from every row
    for row in table.rows:
        if col_index < len(row.cells):
            cell = row.cells[col_index]
            cell._element.getparent().remove(cell._element)


# --- Table Layout Functions ---


def build_table_layout(table: Table, table_id: str) -> TableLayoutInfo:
    """Build table layout info including row heights and alignment."""

    table_align_map = {
        WD_TABLE_ALIGNMENT.LEFT: "left",
        WD_TABLE_ALIGNMENT.CENTER: "center",
        WD_TABLE_ALIGNMENT.RIGHT: "right",
    }
    row_height_rule_map = {
        WD_ROW_HEIGHT_RULE.AUTO: "auto",
        WD_ROW_HEIGHT_RULE.AT_LEAST: "at_least",
        WD_ROW_HEIGHT_RULE.EXACTLY: "exactly",
    }

    rows = []
    for i, row in enumerate(table.rows):
        # Check for header row marker (w:tblHeader in w:trPr)
        tr_el = row._tr
        trPr = tr_el.find(qn("w:trPr"))
        is_header = trPr is not None and trPr.find(qn("w:tblHeader")) is not None
        rows.append(
            RowInfo(
                index=i,
                height_inches=row.height.inches if row.height else None,
                height_rule=row_height_rule_map.get(row.height_rule),
                is_header=is_header,
            )
        )

    return TableLayoutInfo(
        table_id=table_id,
        alignment=table_align_map.get(table.alignment),
        autofit=table.autofit,
        rows=rows,
    )


def set_table_alignment(table: Table, alignment: str) -> None:
    """Set table horizontal alignment."""

    alignment_map = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }
    table.alignment = alignment_map[alignment.lower()]  # KeyError if invalid


def set_row_height(
    table: Table, row_index: int, height_inches: float, rule: str = "at_least"
) -> None:
    """Set row height. Default rule is 'at_least' to prevent text clipping."""

    rule_map = {
        "auto": WD_ROW_HEIGHT_RULE.AUTO,
        "at_least": WD_ROW_HEIGHT_RULE.AT_LEAST,
        "exactly": WD_ROW_HEIGHT_RULE.EXACTLY,
    }
    rule_val = rule.lower()
    row = table.rows[row_index]  # Let IndexError propagate
    row.height = None if rule_val == "auto" else Inches(height_inches)
    row.height_rule = rule_map[rule_val]  # KeyError if invalid


def set_table_fixed_layout(table: Table, column_widths: list[float]) -> None:
    """Set table to fixed layout with explicit column widths (inches)."""

    table.autofit = False
    for i, width in enumerate(column_widths):
        if i < len(table.columns):
            table.columns[i].width = Inches(width)


def set_cell_width(table: Table, row: int, col: int, width_inches: float) -> None:
    """Set cell width."""

    table.cell(row, col).width = Inches(width_inches)


def set_cell_vertical_alignment(
    table: Table, row: int, col: int, alignment: str
) -> None:
    """Set cell vertical alignment."""

    valign_map = {
        "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
        "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
        "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
    }
    table.cell(row, col).vertical_alignment = valign_map[alignment.lower()]


def set_cell_borders(
    table: Table,
    row: int,
    col: int,
    top: str | None = None,
    bottom: str | None = None,
    left: str | None = None,
    right: str | None = None,
) -> None:
    """Set cell borders. Format: 'style:size:color' e.g. 'single:24:000000'.

    Args:
        table: The table object
        row: 0-based row index
        col: 0-based column index
        top, bottom, left, right: Border specs in 'style:size:color' format
            - style: single, double, dotted, dashed, etc.
            - size: in eighths of a point (24 = 3pt)
            - color: hex color (e.g., '000000' for black)
    """
    cell = table.cell(row, col)
    tc = cell._tc

    # Get or create tcPr
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.Element(qn("w:tcPr"))
        tc.insert(0, tcPr)

    # Get or create tcBorders
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))

    def set_border(side: str, spec: str) -> None:
        style, sz, color = spec.split(":")
        border_el = tcBorders.find(qn(f"w:{side}"))
        if border_el is None:
            border_el = etree.SubElement(tcBorders, qn(f"w:{side}"))
        border_el.set(qn("w:val"), style)
        border_el.set(qn("w:sz"), sz)
        border_el.set(qn("w:color"), color)

    if top:
        set_border("top", top)
    if bottom:
        set_border("bottom", bottom)
    if left:
        set_border("left", left)
    if right:
        set_border("right", right)


def set_cell_shading(table: Table, row: int, col: int, fill_color: str) -> None:
    """Set cell background color.

    Args:
        table: The table object
        row: 0-based row index
        col: 0-based column index
        fill_color: Hex color (e.g., 'FF0000' for red)
    """
    cell = table.cell(row, col)
    tc = cell._tc

    # Get or create tcPr
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.Element(qn("w:tcPr"))
        tc.insert(0, tcPr)

    # Get or create shd
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = etree.SubElement(tcPr, qn("w:shd"))

    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_color.upper())


def set_header_row(table: Table, row_index: int, is_header: bool = True) -> None:
    """Mark row as header (repeats on each page in multi-page tables).

    Args:
        table: The table object
        row_index: 0-based row index
        is_header: True to mark as header, False to unmark
    """
    row = table.rows[row_index]
    tr = row._tr

    # Get or create trPr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = etree.Element(qn("w:trPr"))
        tr.insert(0, trPr)

    # Find existing tblHeader
    tblHeader = trPr.find(qn("w:tblHeader"))

    if is_header:
        if tblHeader is None:
            etree.SubElement(trPr, qn("w:tblHeader"))
    else:
        if tblHeader is not None:
            trPr.remove(tblHeader)


def get_header_rows(table: Table) -> list[int]:
    """Get indices of rows marked as headers.

    Returns:
        List of 0-based row indices that are marked as headers
    """
    result = []
    for i, row in enumerate(table.rows):
        tr = row._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None and trPr.find(qn("w:tblHeader")) is not None:
            result.append(i)
    return result


def _build_run_info(
    run, index: int, is_hyperlink: bool = False, hyperlink_url: str | None = None
) -> RunInfo:
    """Build RunInfo from a Run object."""
    return RunInfo(
        index=index,
        text=run.text or "",
        bold=run.bold,
        italic=run.italic,
        underline=run.underline,
        font_name=run.font.name,
        font_size=run.font.size.pt if run.font.size else None,
        color=str(run.font.color.rgb)
        if run.font.color and run.font.color.rgb
        else None,
        highlight_color=_HIGHLIGHT_REVERSE.get(run.font.highlight_color),
        strike=run.font.strike,
        double_strike=run.font.double_strike,
        subscript=run.font.subscript,
        superscript=run.font.superscript,
        style=run.style.name if run.style else None,
        is_hyperlink=is_hyperlink,
        hyperlink_url=hyperlink_url,
        # Additional font properties
        all_caps=run.font.all_caps,
        small_caps=run.font.small_caps,
        hidden=run.font.hidden,
        emboss=run.font.emboss,
        imprint=run.font.imprint,
        outline=run.font.outline,
        shadow=run.font.shadow,
    )


def build_runs(paragraph: Paragraph) -> list[RunInfo]:
    """Build list of RunInfo for all runs in a paragraph, including hyperlink runs."""

    result = []
    idx = 0
    for item in paragraph.iter_inner_content():
        if isinstance(item, Hyperlink):
            url = item.url
            for run in item.runs:
                result.append(
                    _build_run_info(run, idx, is_hyperlink=True, hyperlink_url=url)
                )
                idx += 1
        else:  # Run
            result.append(_build_run_info(item, idx))
            idx += 1
    return result


def build_hyperlinks(doc: Document) -> list[HyperlinkInfo]:
    """Build list of all hyperlinks in the document."""

    result = []
    idx = 0
    for para, _el in _iter_all_paragraphs(doc):
        for item in para.iter_inner_content():
            if isinstance(item, Hyperlink):
                result.append(
                    HyperlinkInfo(
                        index=idx,
                        text=item.text,
                        url=item.url,
                        address=item.address,
                        fragment=item.fragment,
                        is_external=bool(item._hyperlink.rId),
                    )
                )
                idx += 1
    return result


def build_styles(doc: Document) -> list[StyleInfo]:
    """Build list of all styles in the document."""

    style_type_map = {
        WD_STYLE_TYPE.PARAGRAPH: "paragraph",
        WD_STYLE_TYPE.CHARACTER: "character",
        WD_STYLE_TYPE.TABLE: "table",
        WD_STYLE_TYPE.LIST: "list",
    }
    result = []
    for style in doc.styles:
        # Some style types (e.g., _NumberingStyle) don't have all attributes
        base = getattr(style, "base_style", None)
        next_style = getattr(style, "next_paragraph_style", None)
        hidden = getattr(style, "hidden", False)
        quick_style = getattr(style, "quick_style", False)
        result.append(
            StyleInfo(
                name=style.name,
                style_id=style.style_id,
                type=style_type_map.get(style.type, "unknown"),
                builtin=style.builtin,
                base_style=base.name if base else None,
                next_style=(
                    next_style.name if next_style and next_style != style else None
                ),
                hidden=hidden,
                quick_style=quick_style,
            )
        )
    return result


def get_style_format(doc: Document, style_name: str) -> StyleFormatInfo:
    """Get detailed formatting for a specific style."""

    # Alignment normalization map (enum -> API value)
    alignment_to_api = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }

    style = doc.styles[style_name]
    font = style.font

    # Determine style type
    style_type_map = {
        WD_STYLE_TYPE.PARAGRAPH: "paragraph",
        WD_STYLE_TYPE.CHARACTER: "character",
        WD_STYLE_TYPE.TABLE: "table",
        WD_STYLE_TYPE.LIST: "list",
    }
    style_type = style_type_map.get(style.type, "unknown")

    # Extract paragraph format if available
    pf = getattr(style, "paragraph_format", None)
    alignment = None
    left_indent = None
    space_before = None
    space_after = None
    line_spacing = None

    if pf:
        alignment = alignment_to_api.get(pf.alignment) if pf.alignment else None
        left_indent = pf.left_indent.inches if pf.left_indent else None
        space_before = pf.space_before.pt if pf.space_before else None
        space_after = pf.space_after.pt if pf.space_after else None
        # Handle both float multipliers and Length objects (same as build_paragraph_format)
        line_spacing = (
            pf.line_spacing
            if isinstance(pf.line_spacing, float)
            else (pf.line_spacing.pt if pf.line_spacing else None)
        )

    return StyleFormatInfo(
        name=style.name,
        style_id=style.style_id,
        type=style_type,
        font_name=font.name,
        font_size=font.size.pt if font.size else None,
        bold=font.bold,
        italic=font.italic,
        color=str(font.color.rgb) if font.color and font.color.rgb else None,
        alignment=alignment,
        left_indent=left_indent,
        space_before=space_before,
        space_after=space_after,
        line_spacing=line_spacing,
    )


def edit_style(doc: Document, style_name: str, fmt: dict) -> None:
    """Modify a style definition."""

    style = doc.styles[style_name]
    font = style.font

    # Font properties
    if "font_name" in fmt:
        font.name = fmt["font_name"]
    if "font_size" in fmt:
        font.size = Pt(fmt["font_size"])
    if "bold" in fmt:
        font.bold = fmt["bold"]
    if "italic" in fmt:
        font.italic = fmt["italic"]
    if "color" in fmt:
        font.color.rgb = RGBColor.from_string(fmt["color"].lstrip("#"))

    # Paragraph properties (paragraph/table styles only)
    pf = getattr(style, "paragraph_format", None)
    if pf:
        if "alignment" in fmt:
            pf.alignment = getattr(WD_ALIGN_PARAGRAPH, fmt["alignment"].upper())
        if "left_indent" in fmt:
            pf.left_indent = Inches(fmt["left_indent"])
        if "space_before" in fmt:
            pf.space_before = Pt(fmt["space_before"])
        if "space_after" in fmt:
            pf.space_after = Pt(fmt["space_after"])
        if "line_spacing" in fmt:
            val = fmt["line_spacing"]
            pf.line_spacing = val if val < 5 else Pt(val)


def create_style(
    doc: Document,
    name: str,
    style_type: str = "paragraph",
    base_style: str = "Normal",
    formatting: dict | None = None,
) -> str:
    """Create a new custom style.

    Args:
        doc: The Document object
        name: Style name (must be unique)
        style_type: 'paragraph', 'character', or 'table'
        base_style: Name of style to inherit from (default: 'Normal')
        formatting: Optional formatting dict (same as edit_style)

    Returns:
        The style ID of the created style
    """
    from docx.enum.style import WD_STYLE_TYPE

    type_map = {
        "paragraph": WD_STYLE_TYPE.PARAGRAPH,
        "character": WD_STYLE_TYPE.CHARACTER,
        "table": WD_STYLE_TYPE.TABLE,
    }

    wd_style_type = type_map.get(style_type.lower(), WD_STYLE_TYPE.PARAGRAPH)

    # Create the style
    style = doc.styles.add_style(name, wd_style_type)

    # Set base style if specified
    if base_style and base_style in doc.styles:
        style.base_style = doc.styles[base_style]

    # Apply formatting if provided
    if formatting:
        edit_style(doc, name, formatting)

    return style.style_id


def delete_style(doc: Document, style_name: str) -> bool:
    """Delete a custom style.

    Args:
        doc: The Document object
        style_name: Name of the style to delete

    Returns:
        True if deleted, False if style is builtin or not found

    Note:
        Built-in styles cannot be deleted.
    """
    try:
        style = doc.styles[style_name]
    except KeyError:
        return False

    # Cannot delete builtin styles
    if style.builtin:
        return False

    # Delete by removing from styles element
    style._element.getparent().remove(style._element)
    return True


def build_tab_stops(paragraph: Paragraph) -> list[TabStopInfo]:
    """Build list of tab stops for a paragraph."""

    tab_align_map = {
        WD_TAB_ALIGNMENT.LEFT: "left",
        WD_TAB_ALIGNMENT.CENTER: "center",
        WD_TAB_ALIGNMENT.RIGHT: "right",
        WD_TAB_ALIGNMENT.DECIMAL: "decimal",
        WD_TAB_ALIGNMENT.BAR: "bar",
    }
    tab_leader_map = {
        WD_TAB_LEADER.SPACES: "spaces",
        WD_TAB_LEADER.DOTS: "dots",
        WD_TAB_LEADER.HEAVY: "heavy",
        WD_TAB_LEADER.MIDDLE_DOT: "middle_dot",
        None: "spaces",  # None means default (spaces)
    }

    result = []
    for tab in paragraph.paragraph_format.tab_stops:
        alignment = tab_align_map.get(tab.alignment, "unknown")
        leader = tab_leader_map.get(tab.leader, "unknown")
        result.append(
            TabStopInfo(
                position_inches=round(tab.position.inches, 4),
                alignment=alignment,
                leader=leader,
            )
        )
    return result


def add_tab_stop(
    paragraph: Paragraph,
    position_inches: float,
    alignment: str = "left",
    leader: str = "spaces",
) -> None:
    """Add a tab stop to a paragraph."""

    align_map = {
        "left": WD_TAB_ALIGNMENT.LEFT,
        "center": WD_TAB_ALIGNMENT.CENTER,
        "right": WD_TAB_ALIGNMENT.RIGHT,
        "decimal": WD_TAB_ALIGNMENT.DECIMAL,
    }
    leader_map = {
        "spaces": WD_TAB_LEADER.SPACES,
        "dots": WD_TAB_LEADER.DOTS,
        "heavy": WD_TAB_LEADER.HEAVY,
        "middle_dot": WD_TAB_LEADER.MIDDLE_DOT,
    }
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(position_inches),
        align_map[alignment.lower()],  # KeyError if invalid
        leader_map[leader.lower()],  # KeyError if invalid
    )


# --- Document Fields ---


def _make_run_with(*elements):
    """Create a w:r element containing the given child elements."""
    r = OxmlElement("w:r")
    for elem in elements:
        r.append(elem)
    return r


def insert_field(
    paragraph: Paragraph,
    field_code: str,
    display_text: str = "",
    *,
    uppercase: bool = True,
    placeholder: str = "1",
) -> None:
    """Insert a Word field into a paragraph.

    Creates proper OXML field structure with separate runs for each part:
    begin, instruction, separator, result, and end markers.
    Supports any Word field code (PAGE, NUMPAGES, DATE, TIME, AUTHOR, etc.).

    Args:
        uppercase: If True (default), uppercases field_code. Set False for case-sensitive
            fields like bookmark names in cross-references.
        placeholder: Default result text shown before field updates.
    """
    code = field_code.strip().upper() if uppercase else field_code
    p = paragraph._p

    # Run 1: Field begin
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    p.append(_make_run_with(fld_char_begin))

    # Run 2: Field instruction
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {code} "
    p.append(_make_run_with(instr_text))

    # Run 3: Field separator
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    p.append(_make_run_with(fld_char_sep))

    # Run 4: Result text (placeholder shown before field updates)
    text_elem = OxmlElement("w:t")
    text_elem.text = display_text or placeholder
    p.append(_make_run_with(text_elem))

    # Run 5: Field end
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    p.append(_make_run_with(fld_char_end))


def insert_page_x_of_y(
    doc: Document, section_index: int, location: str = "footer"
) -> None:
    """Insert 'Page X of Y' into header or footer.

    Args:
        doc: Document to modify
        section_index: 0-based section index
        location: 'header' or 'footer'
    """
    section = doc.sections[section_index]  # Let IndexError propagate
    hf = {"footer": section.footer, "header": section.header}[
        location
    ]  # KeyError if invalid
    hf.is_linked_to_previous = False
    p = hf.add_paragraph()
    p.add_run("Page ")
    insert_field(p, "PAGE")
    p.add_run(" of ")
    insert_field(p, "NUMPAGES")


def build_paragraph_format(paragraph: Paragraph) -> ParagraphFormatInfo:
    """Extract paragraph formatting properties."""

    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    pf = paragraph.paragraph_format
    alignment = alignment_map.get(paragraph.alignment)
    return ParagraphFormatInfo(
        alignment=alignment,
        left_indent=pf.left_indent.inches if pf.left_indent else None,
        right_indent=pf.right_indent.inches if pf.right_indent else None,
        first_line_indent=pf.first_line_indent.inches if pf.first_line_indent else None,
        space_before=pf.space_before.pt if pf.space_before else None,
        space_after=pf.space_after.pt if pf.space_after else None,
        line_spacing=(
            pf.line_spacing
            if isinstance(pf.line_spacing, float)
            else (pf.line_spacing.pt if pf.line_spacing else None)
        ),
        keep_with_next=pf.keep_with_next,
        page_break_before=pf.page_break_before,
        tab_stops=build_tab_stops(paragraph),
    )


def _resolve_run_by_inner_index(paragraph: Paragraph, run_index: int):
    """Resolve run by iter_inner_content index (matching build_runs indexing).

    This ensures edit_run operations use the same indexing as read(scope='runs'),
    which includes runs inside hyperlinks.
    """

    idx = 0
    for item in paragraph.iter_inner_content():
        if isinstance(item, Hyperlink):
            for run in item.runs:
                if idx == run_index:
                    return run
                idx += 1
        else:  # Run
            if idx == run_index:
                return item
            idx += 1
    raise IndexError(f"Run index {run_index} out of range (paragraph has {idx} runs)")


def edit_run_text(paragraph: Paragraph, run_index: int, text: str) -> None:
    """Edit run text. Uses iter_inner_content indexing (includes hyperlink runs)."""
    run = _resolve_run_by_inner_index(paragraph, run_index)
    run.text = text


def _set_run_attr(run, key: str, value) -> None:
    """Set a run attribute by key. Handles nested paths like 'font.bold'."""
    if key == "style":
        run.style = value
    elif key == "font_size":
        run.font.size = Pt(float(value))
    elif key == "color":
        run.font.color.rgb = RGBColor.from_string(value.lstrip("#"))
    elif key == "highlight_color":
        run.font.highlight_color = _HIGHLIGHT_MAP[value.lower()]
    elif key in _RUN_ATTRS:
        path = _RUN_ATTRS[key]
        obj, attr = (run, path) if "." not in path else (run.font, path.split(".")[1])
        setattr(obj, attr, value)


def edit_run_formatting(paragraph: Paragraph, run_index: int, fmt: dict) -> None:
    """Apply formatting to a specific run. Uses iter_inner_content indexing."""
    run = _resolve_run_by_inner_index(paragraph, run_index)
    for key, value in fmt.items():
        _set_run_attr(run, key, value)


def build_comments(doc: Document) -> list[CommentInfo]:
    """Build list of CommentInfo from document comments."""
    return [
        CommentInfo(
            id=c.comment_id,
            author=c.author,
            initials=c.initials,
            timestamp=c.timestamp.isoformat() if c.timestamp else None,
            text=c.text,
        )
        for c in doc.comments
    ]


def add_comment_to_block(
    doc: Document,
    paragraph: Paragraph,
    text: str,
    author: str = "",
    initials: str = "",
) -> int:
    """Add a comment anchored to all runs in a paragraph. Returns comment_id."""
    return doc.add_comment(
        runs=paragraph.runs, text=text, author=author, initials=initials
    ).comment_id


def _hf_text(hf) -> str | None:
    """Extract header/footer text. Returns None if linked to previous section."""
    return (
        None if hf.is_linked_to_previous else "\n".join(p.text for p in hf.paragraphs)
    )


def build_headers_footers(doc: Document) -> list[HeaderFooterInfo]:
    """Build list of HeaderFooterInfo for all sections."""
    result = []
    for idx, section in enumerate(doc.sections):
        hdr, ftr = section.header, section.footer
        info = HeaderFooterInfo(
            section_index=idx,
            header_text=_hf_text(hdr),
            footer_text=_hf_text(ftr),
            header_is_linked=hdr.is_linked_to_previous,
            footer_is_linked=ftr.is_linked_to_previous,
            has_different_first_page=section.different_first_page_header_footer,
            has_different_odd_even=doc.settings.odd_and_even_pages_header_footer,
        )
        if section.different_first_page_header_footer:
            fp_hdr, fp_ftr = section.first_page_header, section.first_page_footer
            info.first_page_header_text = _hf_text(fp_hdr)
            info.first_page_footer_text = _hf_text(fp_ftr)
        if doc.settings.odd_and_even_pages_header_footer:
            ev_hdr, ev_ftr = section.even_page_header, section.even_page_footer
            info.even_page_header_text = _hf_text(ev_hdr)
            info.even_page_footer_text = _hf_text(ev_ftr)
        result.append(info)
    return result


def set_header_footer_text(
    doc: Document, section_index: int, text: str, location: str
) -> None:
    """Set header/footer text. Handles all types via location attribute name."""
    section = doc.sections[section_index]
    if location.startswith("first_page_"):
        section.different_first_page_header_footer = True
    elif location.startswith("even_page_"):
        doc.settings.odd_and_even_pages_header_footer = True
    hf = getattr(section, location)
    for p in list(hf.paragraphs):
        p._element.getparent().remove(p._element)
    hf.add_paragraph(text)


def append_to_header_footer(
    doc: Document,
    section_index: int,
    content_type: str,
    content_data: str,
    location: str,
) -> str:
    """Append paragraph or table to header/footer. Returns element_id."""
    hf = getattr(doc.sections[section_index], location)
    if content_type == "paragraph":
        p = hf.add_paragraph(content_data)
        occurrence = sum(1 for para in hf.paragraphs if para.text == p.text) - 1
        return make_block_id("paragraph", p.text, occurrence)
    if content_type == "table":
        table_data = json.loads(content_data)
        rows, cols = len(table_data), max((len(r) for r in table_data), default=1)
        tbl = hf.add_table(rows=rows, cols=cols, width=Inches(6))
        populate_table(tbl, table_data)
        return make_block_id("table", table_content_for_hash(tbl), 0)
    raise ValueError(f"Unknown content_type: {content_type}")


def clear_header_footer(doc: Document, section_index: int, location: str) -> None:
    """Clear header/footer content. Unlinks from previous section first."""
    hf = getattr(doc.sections[section_index], location)
    hf.is_linked_to_previous = False
    hf_el = hf._element
    for child in list(hf_el):
        hf_el.remove(child)
    hf.add_paragraph("")  # Word requires at least one paragraph


def build_page_setup(doc: Document) -> list[PageSetupInfo]:
    """Build list of PageSetupInfo for all sections."""

    result = []
    for idx, section in enumerate(doc.sections):
        s = section
        sectPr = s._sectPr

        # Extract column settings from w:cols
        columns = 1
        column_spacing = 0.5
        column_separator = False
        cols_el = sectPr.find(qn("w:cols"))
        if cols_el is not None:
            num = cols_el.get(qn("w:num"))
            columns = int(num) if num else 1
            space = cols_el.get(qn("w:space"))
            if space:
                column_spacing = round(int(space) / 1440, 2)  # twips to inches
            sep = cols_el.get(qn("w:sep"))
            column_separator = sep == "1" or sep == "true"

        # Extract line numbering from w:lnNumType
        line_numbering = None
        ln_el = sectPr.find(qn("w:lnNumType"))
        if ln_el is not None:
            restart_map = {
                "newPage": "newPage",
                "newSection": "newSection",
                "continuous": "continuous",
            }
            restart_val = ln_el.get(qn("w:restart")) or "newPage"
            start_val = ln_el.get(qn("w:start")) or "1"
            count_by_val = ln_el.get(qn("w:countBy")) or "1"
            distance_val = ln_el.get(qn("w:distance")) or "720"
            line_numbering = LineNumberingInfo(
                enabled=True,
                restart=restart_map.get(restart_val, "newPage"),
                start=int(start_val),
                count_by=int(count_by_val),
                distance_inches=round(int(distance_val) / 1440, 2),  # twips to inches
            )

        result.append(
            PageSetupInfo(
                section_index=idx,
                orientation="landscape"
                if s.orientation == WD_ORIENT.LANDSCAPE
                else "portrait",
                page_width=round(s.page_width / 914400, 2) if s.page_width else 0.0,
                page_height=round(s.page_height / 914400, 2) if s.page_height else 0.0,
                top_margin=round(s.top_margin / 914400, 2) if s.top_margin else 0.0,
                bottom_margin=round(s.bottom_margin / 914400, 2)
                if s.bottom_margin
                else 0.0,
                left_margin=round(s.left_margin / 914400, 2) if s.left_margin else 0.0,
                right_margin=round(s.right_margin / 914400, 2)
                if s.right_margin
                else 0.0,
                columns=columns,
                column_spacing=column_spacing,
                column_separator=column_separator,
                line_numbering=line_numbering,
            )
        )
    return result


def set_page_margins(
    doc: Document,
    section_index: int,
    top: float,
    bottom: float,
    left: float,
    right: float,
) -> None:
    """Set page margins for a section. Values in inches."""

    section = doc.sections[section_index]
    section.top_margin = Emu(int(top * 914400))
    section.bottom_margin = Emu(int(bottom * 914400))
    section.left_margin = Emu(int(left * 914400))
    section.right_margin = Emu(int(right * 914400))


def set_page_orientation(doc: Document, section_index: int, orientation: str) -> None:
    """Set page orientation for a section. 'portrait' or 'landscape'."""

    section = doc.sections[section_index]
    w, h = section.page_width, section.page_height

    orient_lower = orientation.lower()
    section.orientation = (
        WD_ORIENT.LANDSCAPE if orient_lower == "landscape" else WD_ORIENT.PORTRAIT
    )
    if orient_lower == "landscape" and h > w or orient_lower == "portrait" and w > h:
        section.page_width, section.page_height = Emu(h), Emu(w)


def add_section(doc: Document, start_type: str = "new_page") -> int:
    """Add new section. Returns new section index (0-based).

    start_type: 'new_page', 'continuous', 'even_page', 'odd_page', 'new_column'.
    """
    section_start = _SECTION_START_MAP.get(start_type.lower(), WD_SECTION.NEW_PAGE)
    doc.add_section(section_start)
    return len(doc.sections) - 1


# OOXML schema order for w:sectPr children (partial list for insertion)
_SECTPR_ORDER = [
    "footnotePr",
    "endnotePr",
    "type",
    "pgSz",
    "pgMar",
    "paperSrc",
    "pgBorders",
    "lnNumType",
    "pgNumType",
    "cols",
    "formProt",
    "vAlign",
    "noEndnote",
    "titlePg",
    "textDirection",
    "bidi",
    "rtlGutter",
    "docGrid",
]


def _insert_sectpr_element(sectPr, element, local_name: str) -> None:
    """Insert element into sectPr at schema-correct position."""
    try:
        target_idx = _SECTPR_ORDER.index(local_name)
    except ValueError:
        # Unknown element, append at end
        sectPr.append(element)
        return

    # Find first child that should come after this element
    for i, child in enumerate(sectPr):
        child_local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        try:
            child_idx = _SECTPR_ORDER.index(child_local)
            if child_idx > target_idx:
                sectPr.insert(i, element)
                return
        except ValueError:
            pass

    # No later element found, append
    sectPr.append(element)


def set_section_columns(
    doc: Document,
    section_index: int,
    num_columns: int,
    spacing_inches: float = 0.5,
    separator: bool = False,
) -> None:
    """Set multi-column layout for a section.

    Args:
        doc: The Document object
        section_index: 0-based section index
        num_columns: Number of columns (1-16)
        spacing_inches: Space between columns in inches
        separator: True to show line between columns
    """
    section = doc.sections[section_index]
    sectPr = section._sectPr

    # Get or create w:cols element
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = etree.Element(qn("w:cols"))
        _insert_sectpr_element(sectPr, cols, "cols")

    # Set attributes
    cols.set(qn("w:num"), str(num_columns))
    cols.set(qn("w:space"), str(int(spacing_inches * 1440)))  # inches to twips
    cols.set(qn("w:sep"), "1" if separator else "0")


def set_line_numbering(
    doc: Document,
    section_index: int,
    enabled: bool = True,
    restart: str = "newPage",
    start: int = 1,
    count_by: int = 1,
    distance_inches: float = 0.5,
) -> None:
    """Enable/configure line numbering for a section.

    Args:
        doc: The Document object
        section_index: 0-based section index
        enabled: True to enable, False to disable line numbering
        restart: When to restart: 'newPage', 'newSection', or 'continuous'
        start: Starting number
        count_by: Show number every N lines
        distance_inches: Distance from margin in inches
    """
    section = doc.sections[section_index]
    sectPr = section._sectPr

    # Find existing w:lnNumType
    lnNumType = sectPr.find(qn("w:lnNumType"))

    if not enabled:
        # Remove if exists
        if lnNumType is not None:
            sectPr.remove(lnNumType)
        return

    # Create if not exists
    if lnNumType is None:
        lnNumType = etree.Element(qn("w:lnNumType"))
        _insert_sectpr_element(sectPr, lnNumType, "lnNumType")

    # Set attributes
    lnNumType.set(qn("w:restart"), restart)
    lnNumType.set(qn("w:start"), str(start))
    lnNumType.set(qn("w:countBy"), str(count_by))
    lnNumType.set(qn("w:distance"), str(int(distance_inches * 1440)))  # inches to twips


# Image support functions


def get_embedded_image_hash(doc: Document, blip) -> str | None:
    """Get SHA1 hash for embedded image. Returns None for linked images."""
    # Support both python-docx elements (with .embed attribute) and raw lxml
    try:
        rel_id = blip.embed  # python-docx accessor
    except AttributeError:
        # Raw lxml element - use qualified namespace
        r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        rel_id = blip.get(f"{{{r_ns}}}embed")
    if not rel_id:
        return None  # Linked image - can't hash external resource
    image_part = doc.part.related_parts[rel_id]
    return image_part.image.sha1[:8]


def _extract_anchor_position(anchor) -> dict:
    """Extract positioning info from wp:anchor element."""
    result = {"position_type": "anchor", "behind_doc": False}

    # behindDoc attribute (no namespace prefix)
    behind = anchor.get("behindDoc")
    result["behind_doc"] = behind == "1" if behind else False

    # Horizontal position
    pos_h = anchor.find("wp:positionH", namespaces=oxml_nsmap)
    if pos_h is not None:
        result["relative_from_h"] = pos_h.get("relativeFrom")
        offset_h = pos_h.find("wp:posOffset", namespaces=oxml_nsmap)
        if offset_h is not None and offset_h.text:
            result["position_h"] = int(offset_h.text) / _EMU_PER_INCH

    # Vertical position
    pos_v = anchor.find("wp:positionV", namespaces=oxml_nsmap)
    if pos_v is not None:
        result["relative_from_v"] = pos_v.get("relativeFrom")
        offset_v = pos_v.find("wp:posOffset", namespaces=oxml_nsmap)
        if offset_v is not None and offset_v.text:
            result["position_v"] = int(offset_v.text) / _EMU_PER_INCH

    # Wrap type (use mapping for consistent API values)
    for xml_name, api_value in _WRAP_XML_TO_API.items():
        wrap_el = anchor.find(f"wp:{xml_name}", namespaces=oxml_nsmap)
        if wrap_el is not None:
            result["wrap_type"] = api_value
            break

    return result


def _extract_images_from_run(
    doc: Document,
    run,
    run_idx: int,
    block_id: str,
    image_hash_counts: dict[str, int],
    images: list[ImageInfo],
) -> None:
    """Extract images from a single run (both inline and anchored)."""
    image_idx_in_run = 0
    for drawing in run._element.findall(".//w:drawing", namespaces=oxml_nsmap):
        # Check for both inline and anchor images
        inline = drawing.find(".//wp:inline", namespaces=oxml_nsmap)
        anchor = drawing.find(".//wp:anchor", namespaces=oxml_nsmap)

        container = inline if inline is not None else anchor
        if container is None:
            continue

        # Guard access to pic element (skip charts/smartart)
        # Try python-docx attribute accessors first, fall back to XPath for raw lxml
        try:
            blip = container.graphic.graphicData.pic.blipFill.blip
        except AttributeError:
            # Raw lxml element - use XPath (lowercase 'ml' in namespace)
            blip_ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
            blip_list = container.findall(".//a:blip", namespaces=blip_ns)
            if not blip_list:
                continue
            blip = blip_list[0]

        h = get_embedded_image_hash(doc, blip)
        if h is None:
            continue  # Skip linked images

        # Track image occurrence globally
        img_occurrence = image_hash_counts.get(h, 0)
        image_hash_counts[h] = img_occurrence + 1

        # Get metadata via XML (avoid InlineShape construction issues)
        # Handle both python-docx elements (with .embed) and raw lxml elements
        try:
            rel_id = blip.embed  # python-docx accessor
        except AttributeError:
            r_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            rel_id = blip.get(f"{{{r_ns}}}embed")
        image_part = doc.part.related_parts[rel_id]
        # Handle extent for both python-docx elements and raw lxml elements
        try:
            extent = container.extent
            width_emu = extent.cx if extent is not None else 0
            height_emu = extent.cy if extent is not None else 0
        except AttributeError:
            # Raw lxml element - find wp:extent element directly
            extent_el = container.find("wp:extent", namespaces=oxml_nsmap)
            if extent_el is not None:
                width_emu = int(extent_el.get("cx", 0))
                height_emu = int(extent_el.get("cy", 0))
            else:
                width_emu, height_emu = 0, 0

        # Build base image info
        info_kwargs = {
            "id": f"image_{h}_{img_occurrence}",
            "width_inches": width_emu / _EMU_PER_INCH,
            "height_inches": height_emu / _EMU_PER_INCH,
            "content_type": image_part.content_type,
            "block_id": block_id,
            "run_index": run_idx,
            "image_index_in_run": image_idx_in_run,
            "filename": image_part.image.filename or "",
        }

        # Add anchor-specific positioning if floating image
        if anchor is not None:
            info_kwargs.update(_extract_anchor_position(anchor))

        images.append(ImageInfo(**info_kwargs))
        image_idx_in_run += 1


def _extract_images_from_paragraph(
    doc: Document,
    para,
    block_id: str,
    image_hash_counts: dict[str, int],
) -> list[ImageInfo]:
    """Extract images from a paragraph, updating occurrence counts.

    Uses iter_inner_content() indexing to match build_runs() indexing.
    """

    images: list[ImageInfo] = []
    run_idx = 0
    for item in para.iter_inner_content():
        if isinstance(item, Hyperlink):
            for run in item.runs:
                _extract_images_from_run(
                    doc, run, run_idx, block_id, image_hash_counts, images
                )
                run_idx += 1
        else:  # Run
            _extract_images_from_run(
                doc, item, run_idx, block_id, image_hash_counts, images
            )
            run_idx += 1
    return images


def build_images(doc: Document) -> list[ImageInfo]:
    """Build list of ImageInfo from document body, including tables."""
    images = []
    block_hash_counts: dict[str, int] = {}  # For block_id computation
    image_hash_counts: dict[str, int] = {}  # For image occurrence

    for kind, obj, _el in iter_body_blocks(doc):
        if kind == "paragraph":
            # Use SAME logic as build_blocks for block_id
            block_type, _ = paragraph_kind_and_level(obj)
            text = obj.text or ""
            block_hash_key = f"{block_type}_{content_hash(text)}"
            block_occurrence = block_hash_counts.get(block_hash_key, 0)
            block_id = make_block_id(block_type, text, block_occurrence)
            block_hash_counts[block_hash_key] = block_occurrence + 1

            images.extend(
                _extract_images_from_paragraph(doc, obj, block_id, image_hash_counts)
            )

        elif kind == "table":
            # Compute table's block_id
            table_content = table_content_for_hash(obj)
            block_hash_key = f"table_{content_hash(table_content)}"
            block_occurrence = block_hash_counts.get(block_hash_key, 0)
            table_block_id = make_block_id("table", table_content, block_occurrence)
            block_hash_counts[block_hash_key] = block_occurrence + 1

            # Search all cells for images with hierarchical block_id
            # Use true grid coordinates to handle merged cells correctly
            rows, cols = len(obj.rows), len(obj.columns)
            for r in range(rows):
                for c in range(cols):
                    cell = obj.cell(r, c)
                    for p_idx, para in enumerate(cell.paragraphs):
                        # Hierarchical block_id: table_abc_0#r0c0/p0
                        hier_block_id = f"{table_block_id}#r{r}c{c}/p{p_idx}"
                        images.extend(
                            _extract_images_from_paragraph(
                                doc, para, hier_block_id, image_hash_counts
                            )
                        )

    return images


def _iter_all_paragraphs(doc: Document):
    """Iterate over all paragraphs in document body and tables."""
    for kind, obj, el in iter_body_blocks(doc):
        if kind == "paragraph":
            yield obj, el
        elif kind == "table":
            for row in obj.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para, para._element


def _iter_all_runs_in_paragraph(para):
    """Iterate all runs in paragraph including those inside hyperlinks.

    Uses iter_inner_content() to match build_runs() indexing.
    """

    for item in para.iter_inner_content():
        if isinstance(item, Hyperlink):
            yield from item.runs
        else:  # Run
            yield item


def _find_image_in_paragraph(doc: Document, para, target_hash: str):
    """Yield each wp:inline or wp:anchor element in paragraph matching target_hash.

    Uses iter_inner_content() traversal to match build_images() indexing.
    """
    for run in _iter_all_runs_in_paragraph(para):
        for drawing in run._element.findall(".//w:drawing", namespaces=oxml_nsmap):
            inline = drawing.find(".//wp:inline", namespaces=oxml_nsmap)
            anchor = drawing.find(".//wp:anchor", namespaces=oxml_nsmap)
            container = inline if inline is not None else anchor
            if container is None:
                continue
            # Try python-docx attribute accessors first, fall back to XPath for raw lxml
            try:
                blip = container.graphic.graphicData.pic.blipFill.blip
            except AttributeError:
                # Raw lxml element - use XPath (lowercase 'ml' in namespace)
                blip_ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
                blip_list = container.findall(".//a:blip", namespaces=blip_ns)
                if not blip_list:
                    continue
                blip = blip_list[0]
            h = get_embedded_image_hash(doc, blip)
            if h == target_hash:
                yield container


def resolve_image(doc: Document, image_id: str) -> tuple:
    """Find embedded image by content-addressable ID. Returns (inline_el, para_el)."""
    target_hash, occurrence_str = _IMAGE_ID_RE.match(image_id).groups()
    target_occurrence = int(occurrence_str)

    occurrence_count = 0
    for para, para_el in _iter_all_paragraphs(doc):
        for inline in _find_image_in_paragraph(doc, para, target_hash):
            if occurrence_count == target_occurrence:
                return inline, para_el
            occurrence_count += 1

    raise ValueError(f"Image not found: {image_id}")


def insert_image(
    doc: Document,
    image_path: str,
    target_id: str,
    position: str,
    width_inches: float = 0,
    height_inches: float = 0,
) -> str:
    """Insert image at target location.

    Supports hierarchical target IDs:
    - table_abc_0#r0c1/p0 -> Insert into paragraph in cell
    - table_abc_0#r0c1 -> Insert into first paragraph of cell
    - table_abc_0 -> Insert before/after table
    - paragraph_abc_0 -> Insert before/after paragraph
    """

    target = resolve_target(doc, target_id)

    # Get image hash first
    _, image = doc.part.get_or_add_image(image_path)
    h = image.sha1[:8]

    width = Inches(width_inches) if width_inches else None
    height = Inches(height_inches) if height_inches else None

    if target.leaf_kind == "paragraph":
        # Insert into this paragraph
        para = target.leaf_obj
        run = para.add_run()
        run.add_picture(image_path, width, height)
        occurrence = count_image_occurrence(doc, h, para._element)
        return f"image_{h}_{occurrence}"

    if target.leaf_kind == "cell":
        # Use first paragraph of cell (create if needed)
        cell = target.leaf_obj
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        run = para.add_run()
        run.add_picture(image_path, width, height)
        occurrence = count_image_occurrence(doc, h, para._element)
        return f"image_{h}_{occurrence}"

    # Target is table or paragraph at base level -> insert before/after
    new_para = doc.add_paragraph()
    run = new_para.add_run()
    run.add_picture(image_path, width, height)

    if position == "before":
        target.leaf_el.addprevious(new_para._element)
    else:
        target.leaf_el.addnext(new_para._element)

    occurrence = count_image_occurrence(doc, h, new_para._element)
    return f"image_{h}_{occurrence}"


def count_image_occurrence(doc: Document, target_hash: str, target_para_el) -> int:
    """Count embedded images with same hash before target paragraph."""
    occurrence = 0
    for para, para_el in _iter_all_paragraphs(doc):
        for _inline in _find_image_in_paragraph(doc, para, target_hash):
            if para_el is target_para_el:
                return occurrence
            occurrence += 1
    raise ValueError("Target paragraph not found")


def delete_image(doc: Document, image_id: str) -> None:
    """Delete an image. Removes containing paragraph if only whitespace remains."""
    inline_el, para_el = resolve_image(doc, image_id)

    # Remove the drawing element (parent of inline)
    drawing_el = inline_el.getparent()
    drawing_el.getparent().remove(drawing_el)

    # Check if paragraph has any remaining content (text, drawings, fields, etc.)
    para = Paragraph(para_el, doc)
    has_content = bool(para.text.strip())
    if not has_content:
        for run in _iter_all_runs_in_paragraph(para):
            # Check for drawings, fields, or other significant content
            run_el = run._element
            if run_el.findall(".//w:drawing", namespaces=oxml_nsmap) or run_el.findall(
                ".//w:fldChar", namespaces=oxml_nsmap
            ):
                has_content = True
                break

    if not has_content:
        para_el.getparent().remove(para_el)


def insert_floating_image(
    doc: Document,
    image_path: str,
    target_id: str,
    position_h: float,
    position_v: float,
    relative_h: str = "column",
    relative_v: str = "paragraph",
    wrap_type: str = "square",
    width_inches: float = 0,
    height_inches: float = 0,
    behind_doc: bool = False,
) -> str:
    """Insert floating (anchored) image at target location.

    Args:
        doc: Document object
        image_path: Path to image file
        target_id: Block ID for anchor paragraph
        position_h: Horizontal position in inches
        position_v: Vertical position in inches
        relative_h: Horizontal reference ("column", "page", "margin", "character")
        relative_v: Vertical reference ("paragraph", "page", "margin", "line")
        wrap_type: Text wrap ("square", "tight", "through", "top_and_bottom", "none")
        width_inches: Image width (0 = auto from image)
        height_inches: Image height (0 = auto from image)
        behind_doc: True to place image behind text

    Returns:
        Image ID (image_{hash}_{occurrence})
    """
    target = resolve_target(doc, target_id)

    # Get/add image to package and get relationship ID
    rId, image = doc.part.get_or_add_image(image_path)
    h = image.sha1[:8]

    # Calculate default dimensions in EMUs from pixels and DPI
    # EMU = (pixels * 914400) / DPI
    default_cx = int(image.px_width * _EMU_PER_INCH / image.horz_dpi)
    default_cy = int(image.px_height * _EMU_PER_INCH / image.vert_dpi)

    # Calculate dimensions in EMUs
    if width_inches and height_inches:
        cx = int(width_inches * _EMU_PER_INCH)
        cy = int(height_inches * _EMU_PER_INCH)
    elif width_inches:
        cx = int(width_inches * _EMU_PER_INCH)
        cy = int(cx * default_cy / default_cx)
    elif height_inches:
        cy = int(height_inches * _EMU_PER_INCH)
        cx = int(cy * default_cx / default_cy)
    else:
        cx, cy = default_cx, default_cy

    # Convert position to EMUs
    offset_h = int(position_h * _EMU_PER_INCH)
    offset_v = int(position_v * _EMU_PER_INCH)

    # Generate unique IDs for drawing elements
    import random

    doc_pr_id = random.randint(100000, 999999)

    # Map wrap_type API value to XML element name
    wrap_xml_name = _WRAP_API_TO_XML.get(wrap_type, "wrapSquare")
    wrap_element = (
        f"<wp:{wrap_xml_name}/>"
        if wrap_type == "none"
        else f'<wp:{wrap_xml_name} wrapText="bothSides"/>'
    )

    # Build anchor XML using OOXML structure
    # Note: namespace URIs use lowercase 'ml' to match python-docx's oxml_nsmap
    anchor_xml = f"""
    <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
               distT="0" distB="0" distL="114300" distR="114300"
               simplePos="0" relativeHeight="251658240" behindDoc="{1 if behind_doc else 0}"
               locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="{relative_h}">
            <wp:posOffset>{offset_h}</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="{relative_v}">
            <wp:posOffset>{offset_v}</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        {wrap_element}
        <wp:docPr id="{doc_pr_id}" name="Picture {doc_pr_id}"/>
        <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks noChangeAspect="1"/>
        </wp:cNvGraphicFramePr>
        <a:graphic>
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                <pic:pic>
                    <pic:nvPicPr>
                        <pic:cNvPr id="{doc_pr_id}" name="Picture {doc_pr_id}"/>
                        <pic:cNvPicPr/>
                    </pic:nvPicPr>
                    <pic:blipFill>
                        <a:blip r:embed="{rId}"/>
                        <a:stretch>
                            <a:fillRect/>
                        </a:stretch>
                    </pic:blipFill>
                    <pic:spPr>
                        <a:xfrm>
                            <a:off x="0" y="0"/>
                            <a:ext cx="{cx}" cy="{cy}"/>
                        </a:xfrm>
                        <a:prstGeom prst="rect">
                            <a:avLst/>
                        </a:prstGeom>
                    </pic:spPr>
                </pic:pic>
            </a:graphicData>
        </a:graphic>
    </wp:anchor>
    """

    # Parse anchor XML
    anchor_el = etree.fromstring(anchor_xml.encode())

    # Create drawing wrapper
    drawing_el = etree.Element(qn("w:drawing"))
    drawing_el.append(anchor_el)

    # Find target paragraph and add drawing
    if target.leaf_kind == "paragraph":
        para = target.leaf_obj
    elif target.leaf_kind == "cell":
        cell = target.leaf_obj
        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    else:
        # Create new paragraph for floating image
        para = doc.add_paragraph()
        if hasattr(target, "leaf_el"):
            target.leaf_el.addnext(para._element)

    # Add run with drawing
    run = para.add_run()
    run._element.append(drawing_el)

    occurrence = count_image_occurrence(doc, h, para._element)
    return f"image_{h}_{occurrence}"


# =============================================================================
# Track Changes / Revision Support
# =============================================================================

from lxml.etree import ElementBase as _LxmlElementBase

# Namespace for revision XPath queries
_REV_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _rev_xpath(element, expr: str) -> list:
    """Execute XPath with revision namespace on python-docx element.

    Uses lxml's ElementBase.xpath() directly rather than compiled XPath,
    since python-docx elements inherit from lxml but their .xpath() wrapper
    doesn't support the namespaces argument.
    """
    return _LxmlElementBase.xpath(element, expr, namespaces=_REV_NS)


# Content revision elements (accept/reject supported)
_CONTENT_REVISIONS = ("w:ins", "w:del")

# Move markers (for range-based pairing)
_MOVE_RANGE_MARKERS = (
    "w:moveFromRangeStart",
    "w:moveFromRangeEnd",
    "w:moveToRangeStart",
    "w:moveToRangeEnd",
)
_MOVE_WRAPPERS = ("w:moveFrom", "w:moveTo")

# Formatting/table revisions (report only, no accept/reject)
_FORMATTING_REVISIONS = (
    "w:pPrChange",
    "w:rPrChange",
    "w:sectPrChange",
    "w:numberingChange",
    "w:tcPrChange",
    "w:trPrChange",
    "w:tblPrChange",
    "w:tblPrExChange",
    "w:cellDel",
    "w:cellIns",
    "w:cellMerge",
    "w:tblGridChange",
)

# All revision tags for detection
_ALL_REVISION_TAGS = (
    _CONTENT_REVISIONS + _MOVE_RANGE_MARKERS + _MOVE_WRAPPERS + _FORMATTING_REVISIONS
)


def has_tracked_changes(doc: Document) -> bool:
    """Check if document body has any tracked changes.

    Searches only within w:body (not headers/footers/footnotes).
    """
    xpath_expr = " | ".join(f"//w:body//{tag}" for tag in _ALL_REVISION_TAGS)
    return bool(_rev_xpath(doc.element, xpath_expr))


def _get_revision_text(element, tag: str) -> str:
    """Extract text from a revision element.

    For deletions (w:del), uses w:delText.
    For insertions (w:ins) and other elements, uses w:t.
    """
    parts = []
    if tag == "del":
        # Deletions use w:delText
        for dt in _rev_xpath(element, ".//w:delText"):
            if dt.text:
                parts.append(dt.text)
    else:
        # Insertions and other elements use w:t
        for t in _rev_xpath(element, ".//w:t"):
            if t.text:
                parts.append(t.text)
    return "".join(parts)


def _classify_revision(tag: str) -> tuple[str, bool]:
    """Classify a revision tag into type and whether accept/reject is supported.

    Returns: (type_name, is_supported)
    """
    if tag == "ins":
        return "insertion", True
    if tag == "del":
        return "deletion", True
    if tag in ("moveFrom", "moveTo"):
        return "move", True  # Move wrappers are supported
    if tag in (
        "moveFromRangeStart",
        "moveFromRangeEnd",
        "moveToRangeStart",
        "moveToRangeEnd",
    ):
        return "move", False  # Range markers are metadata, not actionable
    if tag in ("pPrChange", "rPrChange", "sectPrChange", "numberingChange"):
        return "formatting", False
    if tag in (
        "cellDel",
        "cellIns",
        "cellMerge",
        "tcPrChange",
        "trPrChange",
        "tblPrChange",
        "tblPrExChange",
        "tblGridChange",
    ):
        return "table", False
    return "unknown", False


def _has_field_deletion(element) -> bool:
    """Check if element contains w:delInstrText (field instruction deletion)."""
    return bool(_rev_xpath(element, ".//w:delInstrText"))


def read_tracked_changes(doc: Document) -> list[dict]:
    """List all tracked changes in document body.

    Searches only within w:body (consistent with has_tracked_changes).
    Returns list in document order (no deduplication - same w:id may appear
    multiple times for multi-element revisions like cross-paragraph changes).

    Each entry contains:
    - id: str (w:id attribute, treated as string; may repeat for grouped changes)
    - type: str (insertion, deletion, move, formatting, table)
    - author: str
    - date: str
    - text: str (affected text, empty for formatting/markers)
    - supported: bool (whether accept/reject is implemented)
    - tag: str (original tag name for disambiguation)
    """
    results = []

    # Include all revision elements (content, moves, formatting)
    # Move range markers included for completeness but typically have no text
    all_elements = (
        _CONTENT_REVISIONS
        + _MOVE_WRAPPERS
        + _MOVE_RANGE_MARKERS
        + _FORMATTING_REVISIONS
    )
    xpath_expr = " | ".join(f"//w:body//{tag}" for tag in all_elements)

    for el in _rev_xpath(doc.element, xpath_expr):
        change_id = el.get(qn("w:id"))
        if not change_id:
            continue

        # Get tag name without namespace
        tag = el.tag.split("}")[-1]
        rev_type, supported = _classify_revision(tag)

        # Check for field deletions (unsupported for reject)
        if tag == "del" and _has_field_deletion(el):
            supported = False

        results.append(
            {
                "id": change_id,
                "type": rev_type,
                "author": el.get(qn("w:author"), ""),
                "date": el.get(qn("w:date"), ""),
                "text": _get_revision_text(el, tag),
                "supported": supported,
                "tag": tag,
            }
        )

    return results


# --- Phase 3: Safe XML Helpers for Accept/Reject ---


def _drop_tag_keep_content(element) -> None:
    """Remove element tag, promote children to parent at correct position.

    Handles lxml text/tail correctly:
    1. Find element's index in parent
    2. Preserve element.text by prepending to first child or attaching appropriately
    3. Remove children from element, insert at parent[index], incrementing
    4. Propagate element.tail to last child or previous sibling
    5. Remove the now-empty element

    Used for accepting insertions (unwrap w:ins) and rejecting deletions
    (unwrap w:del after converting delText to t).
    """
    parent = element.getparent()
    if parent is None:
        return

    index = parent.index(element)
    children = list(element)  # Snapshot before mutation

    # Handle element.text (text before first child)
    if element.text:
        if children:
            # Prepend to first child's text
            first_child = children[0]
            first_child.text = element.text + (first_child.text or "")
        else:
            # No children - attach to previous sibling's tail or parent's text
            prev = element.getprevious()
            if prev is not None:
                prev.tail = (prev.tail or "") + element.text
            else:
                parent.text = (parent.text or "") + element.text

    # Move children to parent at correct position
    for child in children:
        element.remove(child)
        parent.insert(index, child)
        index += 1

    # Handle tail text (attach to last promoted child or previous sibling)
    if element.tail:
        if children:
            last_child = children[-1]
            last_child.tail = (last_child.tail or "") + element.tail
        else:
            # No children - attach to previous sibling or parent text
            prev = element.getprevious()
            if prev is not None:
                prev.tail = (prev.tail or "") + element.tail
            else:
                parent.text = (parent.text or "") + element.tail

    parent.remove(element)


def _convert_deltext_to_text(element) -> None:
    """Convert w:delText elements to w:t elements within element.

    Creates new w:t elements rather than renaming tags in-place
    to avoid python-docx element class issues.

    Used when rejecting a deletion - the deleted text must be restored
    as normal text (w:t) rather than deleted text (w:delText).
    """
    for dt in _rev_xpath(element, ".//w:delText"):
        # Create new w:t element
        new_t = OxmlElement("w:t")
        new_t.text = dt.text
        # Preserve xml:space if present
        space_attr = dt.get("{http://www.w3.org/XML/1998/namespace}space")
        if space_attr:
            new_t.set("{http://www.w3.org/XML/1998/namespace}space", space_attr)
        # Replace in parent
        dt_parent = dt.getparent()
        dt_parent.replace(dt, new_t)


def _process_revisions_deepest_first(elements: list) -> list:
    """Sort revision elements by depth (ancestor count), deepest first.

    This ensures inner revisions are processed before outer ones,
    preventing parent removal from orphaning child operations.
    """
    return sorted(elements, key=lambda el: len(list(el.iterancestors())), reverse=True)


def _remove_element_preserve_tail(element) -> None:
    """Remove element from parent while preserving tail text.

    In lxml, each element's .tail contains text that follows the element.
    When removing an element, we must preserve its tail by attaching it
    to the previous sibling's tail or the parent's text.
    """
    parent = element.getparent()
    if parent is None:
        return

    if element.tail:
        prev = element.getprevious()
        if prev is not None:
            prev.tail = (prev.tail or "") + element.tail
        else:
            parent.text = (parent.text or "") + element.tail

    parent.remove(element)


# --- Phase 4: Accept/Reject Functions ---


def _find_elements_by_id(doc: Document, change_id: str, tags: tuple) -> list:
    """Find all elements with given w:id in document body.

    Uses Python filtering instead of XPath interpolation to avoid
    potential issues with special characters in change_id.
    """
    # Select all elements by tag, then filter by w:id in Python
    xpath_expr = " | ".join(f"//w:body//{tag}" for tag in tags)
    all_elements = _rev_xpath(doc.element, xpath_expr)
    w_id_qn = qn("w:id")
    return [el for el in all_elements if el.get(w_id_qn) == change_id]


# --- Move Handling Functions ---


def _find_move_range_markers(doc: Document, move_id: str) -> dict:
    """Find all range markers for a move operation.

    Returns dict with keys: from_start, from_end, to_start, to_end
    Each value is a list of matching elements (usually 0 or 1).
    """
    return {
        "from_start": _find_elements_by_id(doc, move_id, ("w:moveFromRangeStart",)),
        "from_end": _find_elements_by_id(doc, move_id, ("w:moveFromRangeEnd",)),
        "to_start": _find_elements_by_id(doc, move_id, ("w:moveToRangeStart",)),
        "to_end": _find_elements_by_id(doc, move_id, ("w:moveToRangeEnd",)),
    }


def _validate_move_completeness(move_from: list, move_to: list, markers: dict) -> None:
    """Validate that a move has all required components.

    Raises ValueError if move is incomplete (missing wrappers or range markers).
    All four range markers must be present for atomic processing.
    """
    if not move_from and not move_to:
        raise ValueError("Move has no wrappers (moveFrom/moveTo)")

    # Both sides should have wrappers for a complete move
    if not move_from:
        raise ValueError("Move is incomplete: missing moveFrom wrapper")
    if not move_to:
        raise ValueError("Move is incomplete: missing moveTo wrapper")

    # All four range markers must be present for consistent state
    if not markers["from_start"]:
        raise ValueError("Move is incomplete: missing moveFromRangeStart")
    if not markers["from_end"]:
        raise ValueError("Move is incomplete: missing moveFromRangeEnd")
    if not markers["to_start"]:
        raise ValueError("Move is incomplete: missing moveToRangeStart")
    if not markers["to_end"]:
        raise ValueError("Move is incomplete: missing moveToRangeEnd")


def _accept_move(doc: Document, move_id: str, move_from: list, move_to: list) -> None:
    """Accept a move: keep destination content, remove source.

    Processing:
    1. Validate move completeness
    2. Remove all w:moveFrom wrappers entirely (source content discarded)
    3. Unwrap all w:moveTo wrappers (keep destination content)
    4. Remove range markers
    """
    # Validate completeness before mutating
    markers = _find_move_range_markers(doc, move_id)
    _validate_move_completeness(move_from, move_to, markers)

    # Process deepest first
    all_from = _process_revisions_deepest_first(move_from)
    all_to = _process_revisions_deepest_first(move_to)

    # Remove source (moveFrom) - content is discarded, preserve tail
    for el in all_from:
        _remove_element_preserve_tail(el)

    # Keep destination (moveTo) - unwrap to keep content
    for el in all_to:
        _drop_tag_keep_content(el)

    # Clean up range markers (only after successful processing)
    for marker_list in markers.values():
        for marker in marker_list:
            _remove_element_preserve_tail(marker)


def _reject_move(doc: Document, move_id: str, move_from: list, move_to: list) -> None:
    """Reject a move: keep source content, remove destination.

    Processing:
    1. Validate move completeness
    2. Remove all w:moveTo wrappers entirely (destination discarded)
    3. Unwrap all w:moveFrom wrappers (keep source content)
    4. Remove range markers
    """
    # Validate completeness before mutating
    markers = _find_move_range_markers(doc, move_id)
    _validate_move_completeness(move_from, move_to, markers)

    # Process deepest first
    all_from = _process_revisions_deepest_first(move_from)
    all_to = _process_revisions_deepest_first(move_to)

    # Remove destination (moveTo) - content is discarded, preserve tail
    for el in all_to:
        _remove_element_preserve_tail(el)

    # Keep source (moveFrom) - unwrap to keep content
    for el in all_from:
        _drop_tag_keep_content(el)

    # Clean up range markers (only after successful processing)
    for marker_list in markers.values():
        for marker in marker_list:
            _remove_element_preserve_tail(marker)


def accept_change(doc: Document, change_id: str) -> None:
    """Accept a specific tracked change by ID.

    - Insertions (w:ins): Unwrap content, remove tag
    - Deletions (w:del): Remove entirely
    - Moves: Remove source, unwrap destination
    - Formatting: Raises ValueError (not supported)
    """
    # Find all elements with this ID
    ins_elements = _find_elements_by_id(doc, change_id, ("w:ins",))
    del_elements = _find_elements_by_id(doc, change_id, ("w:del",))
    move_from = _find_elements_by_id(doc, change_id, ("w:moveFrom",))
    move_to = _find_elements_by_id(doc, change_id, ("w:moveTo",))
    formatting = _find_elements_by_id(doc, change_id, _FORMATTING_REVISIONS)

    if formatting:
        raise ValueError(f"Cannot accept formatting change {change_id} (not supported)")

    # Handle moves
    if move_from or move_to:
        _accept_move(doc, change_id, move_from, move_to)
        return

    if not ins_elements and not del_elements:
        raise ValueError(f"Change not found: {change_id}")

    # Process deepest first to handle nested revisions
    all_elements = _process_revisions_deepest_first(ins_elements + del_elements)

    for el in all_elements:
        tag = el.tag.split("}")[-1]
        if tag == "ins":
            # Accept insertion: unwrap (keep content)
            _drop_tag_keep_content(el)
        elif tag == "del":
            # Accept deletion: remove entirely (content is deleted)
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)


def reject_change(doc: Document, change_id: str) -> None:
    """Reject a specific tracked change by ID.

    - Insertions (w:ins): Remove entirely
    - Deletions (w:del): Convert delText to t, unwrap
    - Moves: Remove destination, unwrap source
    - Formatting: Raises ValueError (not supported)
    """
    # Find all elements with this ID
    ins_elements = _find_elements_by_id(doc, change_id, ("w:ins",))
    del_elements = _find_elements_by_id(doc, change_id, ("w:del",))
    move_from = _find_elements_by_id(doc, change_id, ("w:moveFrom",))
    move_to = _find_elements_by_id(doc, change_id, ("w:moveTo",))
    formatting = _find_elements_by_id(doc, change_id, _FORMATTING_REVISIONS)

    if formatting:
        raise ValueError(f"Cannot reject formatting change {change_id} (not supported)")

    # Handle moves
    if move_from or move_to:
        _reject_move(doc, change_id, move_from, move_to)
        return

    if not ins_elements and not del_elements:
        raise ValueError(f"Change not found: {change_id}")

    # Check for field deletions (unsupported for reject)
    for el in del_elements:
        if _has_field_deletion(el):
            raise ValueError(
                f"Cannot reject field deletion {change_id} (not supported)"
            )

    # Process deepest first to handle nested revisions
    all_elements = _process_revisions_deepest_first(ins_elements + del_elements)

    for el in all_elements:
        tag = el.tag.split("}")[-1]
        if tag == "ins":
            # Reject insertion: remove entirely (content was never there)
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        elif tag == "del":
            # Reject deletion: restore content (convert delText to t, unwrap)
            _convert_deltext_to_text(el)
            _drop_tag_keep_content(el)


def accept_all_changes(doc: Document) -> int:
    """Accept all supported tracked changes. Returns count.

    Gathers all change IDs first, then processes to avoid
    iteration invalidation during tree mutation.
    """
    changes = read_tracked_changes(doc)
    # Filter to supported changes and get unique IDs (preserving order)
    seen = set()
    supported_ids = []
    for c in changes:
        if c["supported"] and c["id"] not in seen:
            seen.add(c["id"])
            supported_ids.append(c["id"])

    count = 0
    for change_id in supported_ids:
        try:
            accept_change(doc, change_id)
            count += 1
        except ValueError:
            # Skip unsupported changes (shouldn't happen but be safe)
            pass

    return count


def reject_all_changes(doc: Document) -> int:
    """Reject all supported tracked changes. Returns count."""
    changes = read_tracked_changes(doc)
    # Filter to supported changes and get unique IDs (preserving order)
    seen = set()
    supported_ids = []
    for c in changes:
        if c["supported"] and c["id"] not in seen:
            seen.add(c["id"])
            supported_ids.append(c["id"])

    count = 0
    for change_id in supported_ids:
        try:
            reject_change(doc, change_id)
            count += 1
        except ValueError:
            # Skip unsupported changes (e.g., field deletions)
            pass

    return count


# =============================================================================
# List Management (Phase 8)
# =============================================================================


def _get_numbering_part(doc: Document):
    """Get the numbering part from document. Returns None if not present."""
    try:
        return doc.part.numbering_part
    except Exception:
        return None


def _resolve_abstract_num_id(doc: Document, num_id: int) -> int | None:
    """Resolve num_id to abstractNumId via numbering.xml.

    Returns None if num_id not found.
    """
    numbering_part = _get_numbering_part(doc)
    if numbering_part is None:
        return None

    numbering_el = numbering_part._element
    # Find w:num with matching w:numId
    for num_el in _rev_xpath(numbering_el, ".//w:num"):
        if num_el.get(qn("w:numId")) == str(num_id):
            abstract_ref = num_el.find(qn("w:abstractNumId"))
            if abstract_ref is not None:
                return int(abstract_ref.get(qn("w:val")))
    return None


def _resolve_level_format(
    doc: Document, abstract_num_id: int, ilvl: int
) -> dict[str, str | int | None]:
    """Get level format info (numFmt, lvlText, start) from abstractNum.

    Returns dict with keys: format_type, level_text, start_value.
    """
    numbering_part = _get_numbering_part(doc)
    if numbering_part is None:
        return {"format_type": None, "level_text": None, "start_value": None}

    numbering_el = numbering_part._element
    # Find w:abstractNum with matching w:abstractNumId
    for abs_num in _rev_xpath(numbering_el, ".//w:abstractNum"):
        if abs_num.get(qn("w:abstractNumId")) == str(abstract_num_id):
            # Find w:lvl with matching w:ilvl
            for lvl in _rev_xpath(abs_num, ".//w:lvl"):
                if lvl.get(qn("w:ilvl")) == str(ilvl):
                    num_fmt = lvl.find(qn("w:numFmt"))
                    lvl_text = lvl.find(qn("w:lvlText"))
                    start = lvl.find(qn("w:start"))
                    return {
                        "format_type": num_fmt.get(qn("w:val"))
                        if num_fmt is not None
                        else None,
                        "level_text": lvl_text.get(qn("w:val"))
                        if lvl_text is not None
                        else None,
                        "start_value": int(start.get(qn("w:val")))
                        if start is not None
                        else None,
                    }
    return {"format_type": None, "level_text": None, "start_value": None}


def get_list_info(doc: Document, paragraph: Paragraph) -> dict | None:
    """Get list properties for a paragraph.

    Returns None if paragraph is not in a list.
    Returns dict with: num_id, abstract_num_id, level, format_type, start_value, level_text.
    """
    p_el = paragraph._element
    num_pr = p_el.find(qn("w:pPr"))
    if num_pr is None:
        return None

    num_pr_elem = num_pr.find(qn("w:numPr"))
    if num_pr_elem is None:
        return None

    ilvl_el = num_pr_elem.find(qn("w:ilvl"))
    num_id_el = num_pr_elem.find(qn("w:numId"))

    if num_id_el is None:
        return None

    num_id = int(num_id_el.get(qn("w:val")))
    # numId of 0 means "no list"
    if num_id == 0:
        return None

    ilvl = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0

    abstract_num_id = _resolve_abstract_num_id(doc, num_id)
    level_info = (
        _resolve_level_format(doc, abstract_num_id, ilvl)
        if abstract_num_id is not None
        else {}
    )

    return {
        "num_id": num_id,
        "abstract_num_id": abstract_num_id,
        "level": ilvl,
        "format_type": level_info.get("format_type"),
        "start_value": level_info.get("start_value"),
        "level_text": level_info.get("level_text"),
    }


def _ensure_pPr(p_el):
    """Ensure paragraph has pPr element, create if needed."""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    return pPr


def _ensure_numPr(pPr):
    """Ensure pPr has numPr element, create if needed."""
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.insert(0, numPr)
    return numPr


def set_list_level(paragraph: Paragraph, level: int) -> None:
    """Set list indentation level (0-8).

    Only works on paragraphs already in a list.
    Raises ValueError if paragraph is not in a list.
    """
    if not 0 <= level <= 8:
        raise ValueError(f"List level must be 0-8, got {level}")

    p_el = paragraph._element
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        raise ValueError("Paragraph is not in a list")

    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        raise ValueError("Paragraph is not in a list")

    num_id_el = numPr.find(qn("w:numId"))
    if num_id_el is None or num_id_el.get(qn("w:val")) == "0":
        raise ValueError("Paragraph is not in a list")

    # Set or update ilvl
    ilvl_el = numPr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.insert(0, ilvl_el)
    ilvl_el.set(qn("w:val"), str(level))


def promote_list_item(paragraph: Paragraph) -> int:
    """Decrease level (move left). Min level is 0. Returns new level."""
    p_el = paragraph._element
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        raise ValueError("Paragraph is not in a list")

    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        raise ValueError("Paragraph is not in a list")

    num_id_el = numPr.find(qn("w:numId"))
    if num_id_el is None or num_id_el.get(qn("w:val")) == "0":
        raise ValueError("Paragraph is not in a list")

    ilvl_el = numPr.find(qn("w:ilvl"))
    current = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0
    new_level = max(0, current - 1)

    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.insert(0, ilvl_el)
    ilvl_el.set(qn("w:val"), str(new_level))
    return new_level


def demote_list_item(paragraph: Paragraph) -> int:
    """Increase level (move right). Max level is 8. Returns new level."""
    p_el = paragraph._element
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        raise ValueError("Paragraph is not in a list")

    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        raise ValueError("Paragraph is not in a list")

    num_id_el = numPr.find(qn("w:numId"))
    if num_id_el is None or num_id_el.get(qn("w:val")) == "0":
        raise ValueError("Paragraph is not in a list")

    ilvl_el = numPr.find(qn("w:ilvl"))
    current = int(ilvl_el.get(qn("w:val"))) if ilvl_el is not None else 0
    new_level = min(8, current + 1)

    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        numPr.insert(0, ilvl_el)
    ilvl_el.set(qn("w:val"), str(new_level))
    return new_level


def _get_max_num_id(doc: Document) -> int:
    """Get the maximum numId currently in use."""
    numbering_part = _get_numbering_part(doc)
    if numbering_part is None:
        return 0

    max_id = 0
    for num_el in _rev_xpath(numbering_part._element, ".//w:num"):
        num_id = int(num_el.get(qn("w:numId"), "0"))
        max_id = max(max_id, num_id)
    return max_id


def restart_numbering(doc: Document, paragraph: Paragraph, start_value: int = 1) -> int:
    """Restart numbering from given value.

    Creates a new w:num in numbering.xml referencing the same abstractNum,
    with lvlOverride/startOverride, then updates the paragraph's numId.

    Returns the new numId.
    Raises ValueError if paragraph is not in a list.
    """
    list_info = get_list_info(doc, paragraph)
    if list_info is None:
        raise ValueError("Paragraph is not in a list")

    abstract_num_id = list_info["abstract_num_id"]
    ilvl = list_info["level"]

    if abstract_num_id is None:
        raise ValueError("Cannot determine abstract numbering definition")

    numbering_part = _get_numbering_part(doc)
    if numbering_part is None:
        raise ValueError("Document has no numbering part")

    # Create new w:num with unique numId
    new_num_id = _get_max_num_id(doc) + 1

    # Build the new w:num element
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))

    abstract_num_id_el = OxmlElement("w:abstractNumId")
    abstract_num_id_el.set(qn("w:val"), str(abstract_num_id))
    new_num.append(abstract_num_id_el)

    # Add lvlOverride with startOverride
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), str(ilvl))

    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start_value))
    lvl_override.append(start_override)

    new_num.append(lvl_override)

    # Add to numbering.xml
    numbering_part._element.append(new_num)

    # Update paragraph's numId
    p_el = paragraph._element
    pPr = p_el.find(qn("w:pPr"))
    numPr = pPr.find(qn("w:numPr"))
    num_id_el = numPr.find(qn("w:numId"))
    num_id_el.set(qn("w:val"), str(new_num_id))

    return new_num_id


def remove_list_formatting(paragraph: Paragraph) -> None:
    """Remove list formatting from paragraph (removes w:numPr)."""
    p_el = paragraph._element
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        return

    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


# =============================================================================
# TEXT BOX / FLOATING CONTENT SUPPORT
# =============================================================================

# Namespaces for text box discovery (DrawingML + VML)
_TEXTBOX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "v": "urn:schemas-microsoft-com:vml",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
}


def _textbox_xpath(element, expr: str) -> list:
    """Execute XPath with text box namespaces."""
    return _LxmlElementBase.xpath(element, expr, namespaces=_TEXTBOX_NS)


def _extract_text_from_txbxcontent(txbx_content) -> tuple[str, int]:
    """Extract text and paragraph count from w:txbxContent element."""
    paragraphs = _textbox_xpath(txbx_content, ".//w:p")
    para_texts = []
    for p in paragraphs:
        # Get all text nodes within the paragraph
        texts = _textbox_xpath(p, ".//w:t/text()")
        para_texts.append("".join(texts))
    return "\n".join(para_texts), len(paragraphs)


def _get_textbox_dimensions(ancestor) -> tuple[float, float]:
    """Extract width/height in inches from DrawingML extent or VML style."""
    # DrawingML: look for wp:extent
    extents = _textbox_xpath(ancestor, ".//wp:extent")
    if extents:
        ext = extents[0]
        cx = ext.get("cx")  # EMUs
        cy = ext.get("cy")
        width = int(cx) / _EMU_PER_INCH if cx else 0.0
        height = int(cy) / _EMU_PER_INCH if cy else 0.0
        return width, height

    # VML: parse style attribute
    shapes = _textbox_xpath(ancestor, ".//*[starts-with(name(), 'v:')][@style]")
    if shapes:
        style = shapes[0].get("style", "")
        width = height = 0.0
        for part in style.split(";"):
            if ":" in part:
                key, val = part.split(":", 1)
                key = key.strip().lower()
                val = val.strip().lower()
                if key == "width" and val.endswith("pt"):
                    width = float(val[:-2]) / 72.0
                elif key == "height" and val.endswith("pt"):
                    height = float(val[:-2]) / 72.0
        return width, height

    return 0.0, 0.0


def _get_textbox_id_and_name(
    ancestor, source_type: str, index: int
) -> tuple[str, str | None]:
    """Extract stable ID and name from DrawingML docPr or VML shape."""
    if source_type == "drawingml":
        # Look for wp:docPr
        doc_prs = _textbox_xpath(ancestor, ".//wp:docPr")
        if doc_prs:
            doc_pr = doc_prs[0]
            tb_id = doc_pr.get("id", "")
            name = doc_pr.get("name")
            if tb_id:
                return f"textbox_{tb_id}", name
    else:
        # VML: use v:shape @id
        shapes = _textbox_xpath(ancestor, ".//*[starts-with(name(), 'v:')][@id]")
        if shapes:
            shape_id = shapes[0].get("id", "")
            if shape_id:
                return f"textbox_vml_{shape_id}", None

    # Fallback to index-based ID
    return f"textbox_{source_type}_{index}", None


def _get_position_type(ancestor) -> str:
    """Determine if text box is anchored (floating) or inline."""
    # DrawingML: wp:anchor = floating, wp:inline = inline
    # Check both descendants (from w:drawing) and self (from wp:anchor/wp:inline)
    if _textbox_xpath(ancestor, ".//wp:anchor") or ancestor.tag == qn("wp:anchor"):
        return "anchor"
    if _textbox_xpath(ancestor, ".//wp:inline") or ancestor.tag == qn("wp:inline"):
        return "inline"
    # VML in w:pict is typically anchored
    if ancestor.tag == qn("w:pict") or _textbox_xpath(ancestor, ".//v:shape"):
        return "anchor"
    return "unknown"


def _get_wrap_type(ancestor) -> str | None:
    """Get wrap type from DrawingML wrap elements."""
    wrap_map = {
        "wrapSquare": "square",
        "wrapTight": "tight",
        "wrapThrough": "through",
        "wrapTopAndBottom": "topAndBottom",
        "wrapNone": "none",
    }
    for wrap_name, wrap_type in wrap_map.items():
        if _textbox_xpath(ancestor, f".//wp:{wrap_name}"):
            return wrap_type
    return None


def build_text_boxes(doc: Document) -> list[dict]:
    """Build list of all text boxes from both DrawingML and VML.

    Discovery strategy (per OpenAI review):
    1. Search for ALL w:txbxContent elements in document
    2. Classify source by ancestor chain (w:drawing = DrawingML, w:pict = VML)
    3. Handle mc:AlternateContent (both mc:Choice and mc:Fallback)
    """
    text_boxes = []
    seen_ids = set()

    # Search for all w:txbxContent elements
    txbx_contents = _textbox_xpath(doc.element, "//w:txbxContent")

    for idx, txbx in enumerate(txbx_contents):
        # Determine source type by ancestor chain
        parent_chain = list(txbx.iterancestors())
        has_drawing = any(
            p.tag == qn("w:drawing") or "drawing" in str(p.tag).lower()
            for p in parent_chain
        )
        has_pict = any(p.tag == qn("w:pict") for p in parent_chain)

        if has_drawing:
            source_type = "drawingml"
            # Find the w:drawing ancestor for dimension extraction
            ancestor = next(
                (p for p in parent_chain if p.tag == qn("w:drawing")),
                txbx.getparent(),
            )
        elif has_pict:
            source_type = "vml"
            ancestor = next(
                (p for p in parent_chain if p.tag == qn("w:pict")),
                txbx.getparent(),
            )
        else:
            # Unknown container type
            source_type = "unknown"
            ancestor = txbx.getparent()

        # Extract text and paragraph count
        text, para_count = _extract_text_from_txbxcontent(txbx)

        # Get ID and name
        tb_id, name = _get_textbox_id_and_name(ancestor, source_type, idx)

        # Skip duplicates (mc:AlternateContent can have same content twice)
        if tb_id in seen_ids:
            continue
        seen_ids.add(tb_id)

        # Get dimensions
        width, height = _get_textbox_dimensions(ancestor)

        # Get position type
        position_type = _get_position_type(ancestor)

        # Get wrap type
        wrap_type = _get_wrap_type(ancestor)

        text_boxes.append(
            {
                "id": tb_id,
                "name": name,
                "text": text,
                "paragraph_count": para_count,
                "width_inches": round(width, 2),
                "height_inches": round(height, 2),
                "position_type": position_type,
                "source_type": source_type,
                "wrap_type": wrap_type,
            }
        )

    return text_boxes


def _find_textbox_content_by_id(doc: Document, textbox_id: str):
    """Find w:txbxContent element by text box ID."""
    text_boxes = build_text_boxes(doc)
    for idx, tb in enumerate(text_boxes):
        if tb["id"] == textbox_id:
            # Re-find the actual element
            txbx_contents = _textbox_xpath(doc.element, "//w:txbxContent")
            # Need to find the matching one by index
            # Since build_text_boxes filters duplicates, we track which we've seen
            seen_ids = set()
            for tidx, txbx in enumerate(txbx_contents):
                parent_chain = list(txbx.iterancestors())
                has_drawing = any(
                    p.tag == qn("w:drawing") or "drawing" in str(p.tag).lower()
                    for p in parent_chain
                )
                has_pict = any(p.tag == qn("w:pict") for p in parent_chain)
                source_type = (
                    "drawingml" if has_drawing else "vml" if has_pict else "unknown"
                )
                # Find appropriate ancestor element
                drawing_ancestor = next(
                    (p for p in parent_chain if p.tag == qn("w:drawing")), None
                )
                pict_ancestor = next(
                    (p for p in parent_chain if p.tag == qn("w:pict")), None
                )
                if drawing_ancestor is not None:
                    ancestor = drawing_ancestor
                elif pict_ancestor is not None:
                    ancestor = pict_ancestor
                else:
                    ancestor = txbx.getparent()
                tb_id, _ = _get_textbox_id_and_name(ancestor, source_type, tidx)
                if tb_id in seen_ids:
                    continue
                seen_ids.add(tb_id)
                if tb_id == textbox_id:
                    return txbx
    return None


def read_text_box_content(doc: Document, textbox_id: str) -> list[dict]:
    """Read paragraphs inside a text box.

    Returns list of dicts with 'index', 'text', and basic formatting info.
    """
    txbx = _find_textbox_content_by_id(doc, textbox_id)
    if txbx is None:
        raise ValueError(f"Text box not found: {textbox_id}")

    paragraphs = _textbox_xpath(txbx, "./w:p")
    result = []
    for i, p in enumerate(paragraphs):
        texts = _textbox_xpath(p, ".//w:t/text()")
        text = "".join(texts)
        result.append(
            {
                "index": i,
                "text": text,
                "id": f"{textbox_id}/p{i}",
            }
        )
    return result


def edit_text_box_text(
    doc: Document, textbox_id: str, para_index: int, new_text: str
) -> None:
    """Edit text in a text box paragraph.

    Replaces all text in the specified paragraph with new_text.
    """
    txbx = _find_textbox_content_by_id(doc, textbox_id)
    if txbx is None:
        raise ValueError(f"Text box not found: {textbox_id}")

    paragraphs = _textbox_xpath(txbx, "./w:p")
    if para_index < 0 or para_index >= len(paragraphs):
        raise ValueError(
            f"Paragraph index {para_index} out of range (0-{len(paragraphs) - 1})"
        )

    p = paragraphs[para_index]

    # Find all runs and clear their text
    runs = _textbox_xpath(p, "./w:r")
    if runs:
        # Clear all runs and set text in first one
        for run in runs:
            for t in _textbox_xpath(run, "./w:t"):
                t.text = ""
        # Set new text in first run's first w:t
        first_run = runs[0]
        t_elements = _textbox_xpath(first_run, "./w:t")
        if t_elements:
            t_elements[0].text = new_text
        else:
            # Create w:t element
            t = OxmlElement("w:t")
            t.text = new_text
            first_run.append(t)
    else:
        # No runs - create one
        run = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_text
        run.append(t)
        p.append(run)


# =============================================================================
# Cross-References and Captions (Phase 10)
# =============================================================================

# Reserved bookmark prefixes (used internally by Word)
_RESERVED_BOOKMARK_PREFIXES = ("_Toc", "_Ref", "_Hlt", "_GoBack")


def _validate_bookmark_name(name: str) -> None:
    """Validate bookmark name per Word restrictions.

    Word requires:
    - Starts with a letter
    - No spaces (underscores allowed)
    - Max 40 characters
    """
    if not name:
        raise ValueError("Bookmark name cannot be empty")
    if not name[0].isalpha():
        raise ValueError("Bookmark name must start with a letter")
    if " " in name:
        raise ValueError("Bookmark name cannot contain spaces")
    if len(name) > 40:
        raise ValueError("Bookmark name cannot exceed 40 characters")


def _get_next_bookmark_id(doc: Document) -> int:
    """Get next available bookmark ID."""
    max_id = 0
    for start in _rev_xpath(doc.element, "//w:bookmarkStart"):
        bm_id = start.get(qn("w:id"))
        if bm_id:
            max_id = max(max_id, int(bm_id))
    return max_id + 1


def _is_reserved_bookmark(name: str) -> bool:
    """Check if bookmark name is a reserved Word internal bookmark."""
    return any(name.startswith(prefix) for prefix in _RESERVED_BOOKMARK_PREFIXES)


def build_bookmarks(doc: Document) -> list[dict]:
    """List all bookmarks in document.

    Skips reserved bookmarks (_Toc*, _Ref*, etc.) used internally by Word.
    Returns list of dicts with: id, name, block_id.
    """
    results = []
    seen_names = set()

    for start in _rev_xpath(doc.element, "//w:body//w:bookmarkStart"):
        bm_id = start.get(qn("w:id"))
        bm_name = start.get(qn("w:name"))

        if not bm_id or not bm_name:
            continue

        # Skip reserved bookmarks
        if _is_reserved_bookmark(bm_name):
            continue

        # Skip duplicates
        if bm_name in seen_names:
            continue
        seen_names.add(bm_name)

        # Find containing paragraph for block_id
        block_id = ""
        parent = start.getparent()
        while parent is not None:
            if parent.tag == qn("w:p"):
                p = Paragraph(parent, doc)
                kind, level = paragraph_kind_and_level(p)
                text = p.text or ""
                occurrence = count_occurrence(doc, kind, text, parent)
                block_id = make_block_id(kind, text, occurrence)
                break
            parent = parent.getparent()

        results.append(
            {
                "id": int(bm_id),
                "name": bm_name,
                "block_id": block_id,
            }
        )

    return results


def add_bookmark(doc: Document, name: str, paragraph: Paragraph) -> int:
    """Add bookmark at paragraph. Returns bookmark ID.

    Validates:
    - Name starts with letter
    - No spaces (Word restriction)
    - Unique within document
    """
    _validate_bookmark_name(name)

    # Check uniqueness
    for bm in build_bookmarks(doc):
        if bm["name"] == name:
            raise ValueError(f"Bookmark name already exists: {name}")

    bm_id = _get_next_bookmark_id(doc)
    p_el = paragraph._element

    # Insert bookmarkStart at beginning of paragraph
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bm_id))
    bookmark_start.set(qn("w:name"), name)
    p_el.insert(0, bookmark_start)

    # Insert bookmarkEnd at end of paragraph
    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bm_id))
    p_el.append(bookmark_end)

    return bm_id


def insert_cross_reference(
    paragraph: Paragraph, bookmark_name: str, ref_type: str = "text"
) -> None:
    """Insert cross-reference field.

    ref_type options:
    - 'text': REF <bookmark> field (bookmark text)
    - 'number': REF <bookmark> \\r field (paragraph number)
    - 'page': PAGEREF <bookmark> field (page number)

    Bookmark names are case-sensitive in field instructions.
    """
    if ref_type == "text":
        instr = f"REF {bookmark_name} \\h"
    elif ref_type == "number":
        instr = f"REF {bookmark_name} \\r \\h"
    elif ref_type == "page":
        instr = f"PAGEREF {bookmark_name} \\h"
    else:
        raise ValueError(
            f"Unknown ref_type: {ref_type}. Use 'text', 'number', or 'page'"
        )

    insert_field(paragraph, instr, uppercase=False, placeholder="[ref]")


def insert_caption(
    doc: Document,
    target_id: str,
    label: str,
    caption_text: str,
    position: str = "below",
) -> str:
    """Insert caption for table/image. Returns caption block ID.

    Creates paragraph with:
    1. Style "Caption"
    2. Label text (e.g., "Figure ")
    3. SEQ <label> field
    4. Separator (": ")
    5. Caption text

    Args:
        doc: Document to modify
        target_id: Block ID of element to caption
        label: Caption label (e.g., "Figure", "Table")
        caption_text: Text after the number
        position: "below" or "above" the target element
    """
    target = resolve_target(doc, target_id)

    # Create caption paragraph
    caption_p = doc.add_paragraph(style="Caption")

    # Add label text
    caption_p.add_run(f"{label} ")

    # Add SEQ field for auto-numbering
    insert_field(caption_p, f"SEQ {label}", uppercase=False)

    # Add separator and caption text
    caption_p.add_run(f": {caption_text}")

    # Move paragraph to correct position
    # For tables (even with hierarchical targets like table_X#r0c0), caption should be
    # relative to the table itself, not a cell paragraph inside it
    caption_el = caption_p._element
    if target.base_kind == "table":
        target_el = target.base_obj._tbl
    else:
        target_el = target.leaf_el

    if position == "above":
        target_el.addprevious(caption_el)
    else:  # below
        target_el.addnext(caption_el)

    # Generate block ID for the caption
    text = caption_p.text or ""
    occurrence = count_occurrence(doc, "paragraph", text, caption_el)
    return make_block_id("paragraph", text, occurrence)


def build_captions(doc: Document) -> list[dict]:
    """List all captions (paragraphs with style "Caption" containing SEQ fields).

    Returns list of dicts with: id, label, number, text, block_id, style.
    """
    results = []

    for kind, block_obj, element in iter_body_blocks(doc):
        if kind != "paragraph":
            continue

        p = block_obj
        style_name = p.style.name if p.style else ""

        if style_name != "Caption":
            continue

        # Extract text and look for SEQ field pattern
        full_text = p.text or ""

        # Parse the caption: "Label N: text" or "Label N – text"
        # The SEQ field result is embedded in the text
        match = re.match(r"^(\w+)\s+(\d+)[:–\-]\s*(.*)$", full_text)
        if match:
            label, number, caption_text = match.groups()
            number = int(number)
        else:
            # Can't parse - use full text
            label = "Unknown"
            number = 0
            caption_text = full_text

        # Generate block ID
        occurrence = count_occurrence(doc, kind, full_text, element)
        block_id = make_block_id(kind, full_text, occurrence)

        results.append(
            {
                "id": block_id,
                "label": label,
                "number": number,
                "text": caption_text,
                "block_id": block_id,
                "style": style_name,
            }
        )

    return results


# =============================================================================
# Phase 11: Comment Threading
# =============================================================================

# commentsExtended.xml namespace
_W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
_COMMENTS_EXT_NS = {"w15": _W15_NS}

# Content type for commentsExtended.xml
_COMMENTS_EXTENDED_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.commentsExtended+xml"


def _get_comments_extended_part(doc: Document):
    """Get commentsExtended.xml part if exists, else None.

    The part may not exist in older Word documents or those without threaded comments.
    """
    package = doc.part.package
    # Look for the part by content type
    for rel in package.rels.values():
        if hasattr(rel, "_target") and hasattr(rel._target, "content_type"):
            if rel._target.content_type == _COMMENTS_EXTENDED_CT:
                return rel._target
    # Also try by part name pattern
    try:
        for part in package.iter_parts():
            if "/commentsExtended.xml" in part.partname:
                return part
    except Exception:
        pass
    return None


def _parse_comment_threading(doc: Document) -> dict:
    """Parse threading/resolution from commentsExtended.xml.

    Returns: {para_id: {'parent_para_id': str|None, 'done': bool}}
    """
    result = {}
    ext_part = _get_comments_extended_part(doc)
    if ext_part is None:
        return result

    try:
        root = etree.fromstring(ext_part.blob)

        # commentsExtended contains w15:commentEx elements
        # Use Clark notation for w15 namespace since qn() doesn't know w15
        para_id_attr = f"{{{_W15_NS}}}paraId"
        done_attr = f"{{{_W15_NS}}}done"
        parent_attr = f"{{{_W15_NS}}}paraIdParent"

        for comment_ex in root.findall(".//w15:commentEx", namespaces=_COMMENTS_EXT_NS):
            # w15:paraId is the link to the comment (matches w:comment's w15:paraId)
            para_id = comment_ex.get(para_id_attr)
            if not para_id:
                continue

            # w15:done indicates resolution
            done_str = comment_ex.get(done_attr)
            done = done_str == "1" if done_str else False

            # w15:paraIdParent indicates parent comment for threading
            parent_para_id = comment_ex.get(parent_attr)

            result[para_id] = {"parent_para_id": parent_para_id, "done": done}
    except (etree.XMLSyntaxError, AttributeError):
        # Narrow exception handling - only catch XML parsing errors
        pass

    return result


def _get_comment_para_id_map(doc: Document) -> dict:
    """Build mapping from w15:paraId to comment_id.

    python-docx comments don't expose paraId, so we parse comments.xml directly.
    """
    para_id_to_comment_id = {}
    try:
        comments_part = doc.part._comments_part
        if comments_part is None:
            return para_id_to_comment_id

        root = etree.fromstring(comments_part.blob)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # Use Clark notation for attributes - w:id uses standard WML namespace
        w_id_attr = qn("w:id")
        # w15:paraId uses Word 2012 namespace (not in python-docx's qn)
        para_id_attr = f"{{{_W15_NS}}}paraId"

        for comment in root.findall(".//w:comment", namespaces=ns):
            comment_id = comment.get(w_id_attr)
            para_id = comment.get(para_id_attr)
            # Guard against None para_id
            if comment_id is not None and para_id is not None:
                para_id_to_comment_id[para_id] = int(comment_id)
    except (etree.XMLSyntaxError, AttributeError, ValueError):
        # Narrow exception handling for XML/parsing errors
        pass

    return para_id_to_comment_id


def build_comments_with_threading(doc: Document) -> list[dict]:
    """Build comments list with threading info from extended part.

    Falls back to flat list if commentsExtended.xml not present.
    Returns list of dicts compatible with CommentInfo model.
    """
    # Get base comment info from python-docx
    comments = {}
    for c in doc.comments:
        comments[c.comment_id] = {
            "id": c.comment_id,
            "author": c.author,
            "initials": c.initials,
            "timestamp": c.timestamp.isoformat() if c.timestamp else None,
            "text": c.text,
            "parent_id": None,
            "resolved": False,
            "replies": [],
        }

    # Parse extended info if available
    threading_info = _parse_comment_threading(doc)
    para_id_map = _get_comment_para_id_map(doc)

    # Apply threading info
    for para_id, info in threading_info.items():
        comment_id = para_id_map.get(para_id)
        if comment_id is None or comment_id not in comments:
            continue

        comments[comment_id]["resolved"] = info["done"]

        # Find parent comment by para_id
        parent_para_id = info.get("parent_para_id")
        if parent_para_id:
            parent_comment_id = para_id_map.get(parent_para_id)
            if parent_comment_id is not None:
                comments[comment_id]["parent_id"] = parent_comment_id

    # Build replies lists
    for comment_id, comment in comments.items():
        parent_id = comment["parent_id"]
        if parent_id is not None and parent_id in comments:
            comments[parent_id]["replies"].append(comment_id)

    # Return sorted by comment ID for deterministic ordering
    return sorted(comments.values(), key=lambda c: c["id"])


def reply_to_comment(
    doc: Document, parent_id: int, text: str, author: str = "", initials: str = ""
) -> int:
    """Add reply to existing comment. Returns new comment ID.

    Note: Full threading support requires commentsExtended.xml manipulation
    which involves OPC packaging. This creates a basic reply comment
    anchored to the same location as the parent.
    """
    # Validate parent comment exists
    parent_comment = None
    for c in doc.comments:
        if c.comment_id == parent_id:
            parent_comment = c
            break

    if parent_comment is None:
        raise ValueError(f"Parent comment not found: {parent_id}")

    # Find runs anchored to parent comment by searching for commentRangeStart
    # with matching ID in the document body
    anchored_runs = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Find the commentRangeStart for this comment ID
    range_start = doc.element.find(
        f".//w:commentRangeStart[@w:id='{parent_id}']", namespaces=ns
    )

    if range_start is not None:
        # Find the containing paragraph
        parent_el = range_start.getparent()
        while parent_el is not None:
            if parent_el.tag == qn("w:p"):
                # Found paragraph - get its runs
                para = Paragraph(parent_el, doc)
                anchored_runs = para.runs
                break
            parent_el = parent_el.getparent()

    # Fall back to first paragraph with runs if no anchored runs found
    if not anchored_runs:
        for para in doc.paragraphs:
            if para.runs:
                anchored_runs = para.runs
                break

    if not anchored_runs:
        raise ValueError("No runs available to anchor reply comment")

    # Create reply comment anchored to same location as parent
    new_comment = doc.add_comment(
        runs=anchored_runs, text=text, author=author, initials=initials
    )

    return new_comment.comment_id


def resolve_comment(doc: Document, comment_id: int) -> None:
    """Mark comment as resolved.

    Note: Full resolution support requires commentsExtended.xml manipulation.
    This is a placeholder that validates the comment exists.
    Word 2013+ uses commentsExtended.xml with w15:done="1" attribute.
    """
    # Validate comment exists
    found = False
    for c in doc.comments:
        if c.comment_id == comment_id:
            found = True
            break

    if not found:
        raise ValueError(f"Comment not found: {comment_id}")

    # Note: Actual resolution requires modifying commentsExtended.xml
    # which involves complex OPC packaging. This is a placeholder.
    # For now, we validate the comment exists - full implementation
    # would create/modify commentsExtended.xml part.


def unresolve_comment(doc: Document, comment_id: int) -> None:
    """Mark comment as unresolved (clears 'done' state).

    Note: Full support requires commentsExtended.xml manipulation.
    """
    # Validate comment exists
    found = False
    for c in doc.comments:
        if c.comment_id == comment_id:
            found = True
            break

    if not found:
        raise ValueError(f"Comment not found: {comment_id}")


# =============================================================================
# Phase 12: Table of Contents
# =============================================================================


def has_toc(doc: Document) -> bool:
    """Check if document has a Table of Contents.

    Searches for:
    1. w:instrText containing "TOC" (complex field)
    2. w:fldSimple[@w:instr] starting with "TOC"
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Check for complex field with TOC
    for instr in doc.element.findall(".//w:instrText", namespaces=ns):
        if instr.text and "TOC" in instr.text.upper():
            return True

    # Check for simple field with TOC
    for fld in doc.element.findall(".//w:fldSimple", namespaces=ns):
        instr = fld.get(qn("w:instr")) or ""
        if instr.strip().upper().startswith("TOC"):
            return True

    return False


def get_toc_info(doc: Document) -> dict:
    """Get TOC metadata if exists.

    Parses heading levels from field switches (e.g., \\o "1-3").
    Returns dict compatible with TOCInfo model.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    result = {
        "exists": False,
        "heading_levels": "1-3",
        "entry_count": 0,
        "block_id": None,
        "has_sdt_wrapper": False,
        "is_dirty": False,
    }

    # Find TOC field instruction
    toc_instr = None
    toc_para = None
    is_dirty = False

    # Check for complex field with TOC
    for instr_el in doc.element.findall(".//w:instrText", namespaces=ns):
        if instr_el.text and "TOC" in instr_el.text.upper():
            toc_instr = instr_el.text
            # Find containing paragraph
            parent = instr_el.getparent()
            while parent is not None:
                if parent.tag == qn("w:p"):
                    toc_para = parent
                    break
                parent = parent.getparent()

            # Check for dirty flag on fldChar begin
            if toc_para is not None:
                for fld_char in toc_para.findall(".//w:fldChar", namespaces=ns):
                    if fld_char.get(qn("w:fldCharType")) == "begin":
                        dirty = fld_char.get(qn("w:dirty"))
                        is_dirty = dirty == "true" or dirty == "1"
                        break
            break

    # Check for simple field with TOC
    if not toc_instr:
        for fld in doc.element.findall(".//w:fldSimple", namespaces=ns):
            instr = fld.get(qn("w:instr")) or ""
            if instr.strip().upper().startswith("TOC"):
                toc_instr = instr
                # Find containing paragraph
                parent = fld.getparent()
                while parent is not None:
                    if parent.tag == qn("w:p"):
                        toc_para = parent
                        break
                    parent = parent.getparent()
                # Check dirty on simple field
                dirty = fld.get(qn("w:dirty"))
                is_dirty = dirty == "true" or dirty == "1"
                break

    if not toc_instr:
        return result

    result["exists"] = True
    result["is_dirty"] = is_dirty

    # Parse heading levels from \o switch
    levels_match = re.search(r'\\o\s*"(\d+-\d+)"', toc_instr)
    if levels_match:
        result["heading_levels"] = levels_match.group(1)

    # Get block ID for the TOC paragraph
    if toc_para is not None:
        p = Paragraph(toc_para, doc)
        kind, level = paragraph_kind_and_level(p)
        text = p.text or ""
        occurrence = count_occurrence(doc, kind, text, toc_para)
        result["block_id"] = make_block_id(kind, text, occurrence)

        # Check for SDT wrapper
        parent = toc_para.getparent()
        if parent is not None and parent.tag == qn("w:sdt"):
            result["has_sdt_wrapper"] = True

    return result


def insert_toc(
    doc: Document,
    target_id: str,
    position: str = "before",
    heading_levels: str = "1-3",
) -> str:
    """Insert TOC field at position. Returns block ID.

    Field code: TOC \\o "1-3" \\h \\z \\u
    - \\o: heading levels
    - \\h: hyperlinks
    - \\z: hide tab leaders and page numbers in Web view
    - \\u: use applied paragraph outline level

    Sets w:dirty="true" so Word updates on open.
    """
    target = resolve_target(doc, target_id)

    # Create paragraph for TOC
    toc_para = doc.add_paragraph()

    # Build field instruction
    instr = f' TOC \\\\o "{heading_levels}" \\\\h \\\\z \\\\u '

    # Insert field with 5-run structure
    # Run 1: fldChar begin (with dirty flag)
    run1 = toc_para.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    run1._r.append(fld_char_begin)

    # Run 2: instrText
    run2 = toc_para.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = instr
    run2._r.append(instr_text)

    # Run 3: fldChar separate
    run3 = toc_para.add_run()
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_char_sep)

    # Run 4: result text (placeholder - Word will replace)
    toc_para.add_run("Update this field to generate Table of Contents")

    # Run 5: fldChar end
    run5 = toc_para.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_char_end)

    # Move paragraph to correct position
    toc_el = toc_para._element
    if target.base_kind == "table":
        target_el = target.base_obj._tbl
    else:
        target_el = target.leaf_el

    if position == "before":
        target_el.addprevious(toc_el)
    else:
        target_el.addnext(toc_el)

    # Generate block ID
    text = toc_para.text or ""
    occurrence = count_occurrence(doc, "paragraph", text, toc_el)
    return make_block_id("paragraph", text, occurrence)


def update_toc_field(doc: Document) -> bool:
    """Set dirty flag on TOC field begin marker.

    Sets w:dirty="true" on w:fldChar[@w:fldCharType="begin"] for complex fields,
    or w:dirty="true" on w:fldSimple for simple fields.
    Word recalculates field values when document opens.

    Returns True if TOC was found and updated, False otherwise.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    # Find complex field with TOC
    for instr_el in doc.element.findall(".//w:instrText", namespaces=ns):
        if instr_el.text and "TOC" in instr_el.text.upper():
            # Find containing paragraph and fldChar begin
            parent = instr_el.getparent()
            while parent is not None:
                if parent.tag == qn("w:p"):
                    # Find fldChar begin in this paragraph
                    for fld_char in parent.findall(".//w:fldChar", namespaces=ns):
                        if fld_char.get(qn("w:fldCharType")) == "begin":
                            fld_char.set(qn("w:dirty"), "true")
                            return True
                    break
                parent = parent.getparent()
            break

    # Find simple field with TOC
    for fld in doc.element.findall(".//w:fldSimple", namespaces=ns):
        instr = fld.get(qn("w:instr")) or ""
        if instr.strip().upper().startswith("TOC"):
            fld.set(qn("w:dirty"), "true")
            return True

    return False


# =============================================================================
# Phase 14: Footnotes & Endnotes
# =============================================================================

# Namespaces for footnote operations
_FN_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_FN_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_FN_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_FN_XML_NS = "http://www.w3.org/XML/1998/namespace"

# Reserved footnote/endnote IDs (separators only, user notes use 1+)
_RESERVED_NOTE_IDS = {-1, 0}


def build_footnotes(doc: Document) -> list[dict]:
    """List all footnotes and endnotes with their content.

    Reads from word/footnotes.xml and word/endnotes.xml via package parts,
    matching references in the document body. Returns list of dicts with
    id, type, text, block_id.
    """
    result = []
    ns = {"w": _FN_W_NS}

    # Build map of reference locations in document (id -> block_id)
    ref_locations: dict[tuple[str, int], str] = {}  # (type, id) -> block_id

    # Find footnote references in document body
    for fn_ref in doc.element.findall(".//w:footnoteReference", namespaces=ns):
        fn_id = fn_ref.get(qn("w:id"))
        if fn_id:
            # Find containing block
            parent = fn_ref.getparent()
            while parent is not None:
                if parent.tag == qn("w:p"):
                    p = Paragraph(parent, doc)
                    kind, level = paragraph_kind_and_level(p)
                    text = p.text or ""
                    occurrence = count_occurrence(doc, kind, text, parent)
                    block_id = make_block_id(kind, text, occurrence)
                    ref_locations[("footnote", int(fn_id))] = block_id
                    break
                parent = parent.getparent()

    # Find endnote references
    for en_ref in doc.element.findall(".//w:endnoteReference", namespaces=ns):
        en_id = en_ref.get(qn("w:id"))
        if en_id:
            parent = en_ref.getparent()
            while parent is not None:
                if parent.tag == qn("w:p"):
                    p = Paragraph(parent, doc)
                    kind, level = paragraph_kind_and_level(p)
                    text = p.text or ""
                    occurrence = count_occurrence(doc, kind, text, parent)
                    block_id = make_block_id(kind, text, occurrence)
                    ref_locations[("endnote", int(en_id))] = block_id
                    break
                parent = parent.getparent()

    # Read footnotes/endnotes from package parts
    for part in doc.part.package.iter_parts():
        partname = part.partname
        if partname == "/word/footnotes.xml":
            fn_root = etree.fromstring(part.blob)
            for fn in fn_root.xpath("//w:footnote", namespaces=ns):
                fn_id_str = fn.get(f"{{{_FN_W_NS}}}id")
                if fn_id_str:
                    fn_id = int(fn_id_str)
                    if fn_id in _RESERVED_NOTE_IDS:
                        continue  # Skip separators
                    # Extract text content
                    text_parts = fn.xpath(".//w:t/text()", namespaces=ns)
                    text = "".join(text_parts).strip()
                    block_id = ref_locations.get(("footnote", fn_id), "")
                    result.append(
                        {
                            "id": fn_id,
                            "type": "footnote",
                            "text": text,
                            "block_id": block_id,
                        }
                    )
        elif partname == "/word/endnotes.xml":
            en_root = etree.fromstring(part.blob)
            for en in en_root.xpath("//w:endnote", namespaces=ns):
                en_id_str = en.get(f"{{{_FN_W_NS}}}id")
                if en_id_str:
                    en_id = int(en_id_str)
                    if en_id in _RESERVED_NOTE_IDS:
                        continue  # Skip separators
                    text_parts = en.xpath(".//w:t/text()", namespaces=ns)
                    text = "".join(text_parts).strip()
                    block_id = ref_locations.get(("endnote", en_id), "")
                    result.append(
                        {
                            "id": en_id,
                            "type": "endnote",
                            "text": text,
                            "block_id": block_id,
                        }
                    )

    return result


def _get_safe_note_id(notes_root, ns: dict) -> int:
    """Get a safe footnote/endnote ID avoiding reserved values."""
    used_ids: set[int] = set()
    for fn in notes_root.xpath("//w:footnote | //w:endnote", namespaces=ns):
        fn_id = fn.get(f"{{{_FN_W_NS}}}id")
        if fn_id:
            try:
                used_ids.add(int(fn_id))
            except ValueError:
                pass

    candidate_id = 1  # User notes start at 1, reserved are -1 and 0
    while candidate_id in used_ids or candidate_id in _RESERVED_NOTE_IDS:
        candidate_id += 1
    return candidate_id


def _ensure_note_content_types(ct_xml: bytes, note_type: str) -> bytes:
    """Ensure content types include footnotes/endnotes."""
    ct_root = etree.fromstring(ct_xml)
    ns = {"ct": _FN_CT_NS}

    part_name = f"/word/{note_type}s.xml"
    if note_type == "footnote":
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
    else:
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"

    existing = ct_root.xpath(f"//ct:Override[@PartName='{part_name}']", namespaces=ns)
    if existing:
        return ct_xml

    override = etree.Element(
        f"{{{_FN_CT_NS}}}Override",
        PartName=part_name,
        ContentType=content_type,
    )
    ct_root.append(override)
    return etree.tostring(
        ct_root, encoding="UTF-8", xml_declaration=True, standalone="yes"
    )


def _ensure_note_relationship(rels_xml: bytes, note_type: str) -> bytes:
    """Ensure document relationships include footnotes/endnotes."""
    rels_root = etree.fromstring(rels_xml)
    ns = {"r": _FN_REL_NS}

    rel_type = f"http://schemas.openxmlformats.org/officeDocument/2006/relationships/{note_type}s"
    existing = rels_root.xpath(
        f"//r:Relationship[contains(@Type, '{note_type}s')]", namespaces=ns
    )
    if existing:
        return rels_xml

    # Generate unique rId
    all_rels = rels_root.xpath("//r:Relationship", namespaces=ns)
    existing_ids = {rel.get("Id") for rel in all_rels if rel.get("Id")}
    rid_num = 1
    while f"rId{rid_num}" in existing_ids:
        rid_num += 1

    rel = etree.Element(
        f"{{{_FN_REL_NS}}}Relationship",
        Id=f"rId{rid_num}",
        Type=rel_type,
        Target=f"{note_type}s.xml",
    )
    rels_root.append(rel)
    return etree.tostring(
        rels_root, encoding="UTF-8", xml_declaration=True, standalone="yes"
    )


def _create_minimal_notes_xml(note_type: str) -> bytes:
    """Create minimal footnotes.xml or endnotes.xml with separators."""
    tag = "footnotes" if note_type == "footnote" else "endnotes"
    item = "footnote" if note_type == "footnote" else "endnote"
    xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:{tag} xmlns:w="{_FN_W_NS}">
    <w:{item} w:type="separator" w:id="-1">
        <w:p>
            <w:pPr>
                <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
            </w:pPr>
            <w:r>
                <w:separator/>
            </w:r>
        </w:p>
    </w:{item}>
    <w:{item} w:type="continuationSeparator" w:id="0">
        <w:p>
            <w:pPr>
                <w:spacing w:after="0" w:line="240" w:lineRule="auto"/>
            </w:pPr>
            <w:r>
                <w:continuationSeparator/>
            </w:r>
        </w:p>
    </w:{item}>
</w:{tag}>'''
    return xml.encode("utf-8")


def _ensure_note_styles(styles_root, note_type: str) -> None:
    """Ensure footnote/endnote styles exist."""
    ns = {"w": _FN_W_NS}
    ref_style_id = f"{note_type.title()}Reference"
    text_style_id = f"{note_type.title()}Text"

    # Check for reference style
    ref_style = styles_root.xpath(
        f'//w:style[@w:styleId="{ref_style_id}"]', namespaces=ns
    )
    if not ref_style:
        style = etree.Element(
            f"{{{_FN_W_NS}}}style",
            attrib={
                f"{{{_FN_W_NS}}}type": "character",
                f"{{{_FN_W_NS}}}styleId": ref_style_id,
            },
        )
        name = etree.SubElement(style, f"{{{_FN_W_NS}}}name")
        name.set(f"{{{_FN_W_NS}}}val", f"{note_type} reference")
        base = etree.SubElement(style, f"{{{_FN_W_NS}}}basedOn")
        base.set(f"{{{_FN_W_NS}}}val", "DefaultParagraphFont")
        rPr = etree.SubElement(style, f"{{{_FN_W_NS}}}rPr")
        vert_align = etree.SubElement(rPr, f"{{{_FN_W_NS}}}vertAlign")
        vert_align.set(f"{{{_FN_W_NS}}}val", "superscript")
        styles_root.append(style)

    # Check for text style
    text_style = styles_root.xpath(
        f'//w:style[@w:styleId="{text_style_id}"]', namespaces=ns
    )
    if not text_style:
        style = etree.Element(
            f"{{{_FN_W_NS}}}style",
            attrib={
                f"{{{_FN_W_NS}}}type": "paragraph",
                f"{{{_FN_W_NS}}}styleId": text_style_id,
            },
        )
        name = etree.SubElement(style, f"{{{_FN_W_NS}}}name")
        name.set(f"{{{_FN_W_NS}}}val", f"{note_type} text")
        base = etree.SubElement(style, f"{{{_FN_W_NS}}}basedOn")
        base.set(f"{{{_FN_W_NS}}}val", "Normal")
        pPr = etree.SubElement(style, f"{{{_FN_W_NS}}}pPr")
        sz = etree.SubElement(pPr, f"{{{_FN_W_NS}}}sz")
        sz.set(f"{{{_FN_W_NS}}}val", "20")  # 10pt
        styles_root.append(style)


def add_footnote(
    doc_path: str,
    target_id: str,
    text: str,
    note_type: str = "footnote",
    position: str = "after",
) -> int:
    """Add a footnote or endnote to a document.

    Args:
        doc_path: Path to the Word document
        target_id: Block ID where the reference should be placed
        text: Content of the footnote/endnote
        note_type: "footnote" or "endnote"
        position: "after" (end of paragraph) or "before" (start)

    Returns:
        The ID of the new footnote/endnote

    Raises:
        ValueError: If target not found or invalid location
        FileNotFoundError: If document not found
    """
    import os
    import zipfile

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Document not found: {doc_path}")

    ns = {"w": _FN_W_NS}
    notes_file = f"word/{note_type}s.xml"
    note_tag = note_type

    # Read document parts
    doc_parts: dict[str, bytes] = {}
    with zipfile.ZipFile(doc_path, "r") as zf:
        doc_parts["document"] = zf.read("word/document.xml")
        doc_parts["content_types"] = zf.read("[Content_Types].xml")
        doc_parts["document_rels"] = zf.read("word/_rels/document.xml.rels")

        if notes_file in zf.namelist():
            doc_parts["notes"] = zf.read(notes_file)
        else:
            doc_parts["notes"] = _create_minimal_notes_xml(note_type)

        if "word/styles.xml" in zf.namelist():
            doc_parts["styles"] = zf.read("word/styles.xml")
        else:
            doc_parts["styles"] = (
                f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:styles xmlns:w="{_FN_W_NS}"/>'.encode()
            )

    # Parse XML
    doc_root = etree.fromstring(doc_parts["document"])
    notes_root = etree.fromstring(doc_parts["notes"])
    styles_root = etree.fromstring(doc_parts["styles"])

    # Load document with python-docx to find target paragraph
    temp_doc = Document(doc_path)

    # Use find_paragraph_by_id which handles all paragraph types
    target_para_obj = find_paragraph_by_id(temp_doc, target_id)
    if target_para_obj is None:
        raise ValueError(f"Target block not found: {target_id}")

    # Get all paragraphs from both python-docx body and raw XML
    # Use index-based matching (more robust than text matching for duplicates)
    all_paras_docx = list(temp_doc.element.body.iter(qn("w:p")))
    all_paras_xml = doc_root.xpath("//w:body//w:p", namespaces=ns)

    # Find index of target in python-docx element tree
    try:
        target_idx = all_paras_docx.index(target_para_obj._element)
    except ValueError:
        raise ValueError(f"Target element not found in document structure: {target_id}")

    if target_idx >= len(all_paras_xml):
        raise ValueError(f"Paragraph index mismatch: {target_idx}")

    target_para = all_paras_xml[target_idx]

    # Validate location (not in header/footer)
    parent = target_para.getparent()
    while parent is not None:
        if parent.tag in [f"{{{_FN_W_NS}}}hdr", f"{{{_FN_W_NS}}}ftr"]:
            raise ValueError("Cannot add footnote/endnote in header/footer")
        parent = parent.getparent()

    # Get safe note ID
    note_id = _get_safe_note_id(notes_root, ns)

    # Find insertion position in paragraph
    runs = target_para.xpath(".//w:r", namespaces=ns)
    if position == "after" and runs:
        last_run = runs[-1]
        insert_pos = list(target_para).index(last_run) + 1
    elif position == "before" and runs:
        first_run = runs[0]
        insert_pos = list(target_para).index(first_run)
    else:
        insert_pos = len(target_para)

    # Create footnote reference run in document
    ref_run = etree.Element(f"{{{_FN_W_NS}}}r")
    rPr = etree.SubElement(ref_run, f"{{{_FN_W_NS}}}rPr")
    rStyle = etree.SubElement(rPr, f"{{{_FN_W_NS}}}rStyle")
    rStyle.set(f"{{{_FN_W_NS}}}val", f"{note_type.title()}Reference")
    fn_ref = etree.SubElement(ref_run, f"{{{_FN_W_NS}}}{note_type}Reference")
    fn_ref.set(f"{{{_FN_W_NS}}}id", str(note_id))
    target_para.insert(insert_pos, ref_run)

    # Create footnote content in notes file
    new_note = etree.Element(
        f"{{{_FN_W_NS}}}{note_tag}",
        attrib={f"{{{_FN_W_NS}}}id": str(note_id)},
    )

    # Add paragraph with content
    fn_para = etree.SubElement(new_note, f"{{{_FN_W_NS}}}p")
    pPr = etree.SubElement(fn_para, f"{{{_FN_W_NS}}}pPr")
    pStyle = etree.SubElement(pPr, f"{{{_FN_W_NS}}}pStyle")
    pStyle.set(f"{{{_FN_W_NS}}}val", f"{note_type.title()}Text")

    # Add footnote reference marker
    marker_run = etree.SubElement(fn_para, f"{{{_FN_W_NS}}}r")
    marker_rPr = etree.SubElement(marker_run, f"{{{_FN_W_NS}}}rPr")
    marker_rStyle = etree.SubElement(marker_rPr, f"{{{_FN_W_NS}}}rStyle")
    marker_rStyle.set(f"{{{_FN_W_NS}}}val", f"{note_type.title()}Reference")
    etree.SubElement(marker_run, f"{{{_FN_W_NS}}}{note_type}Ref")

    # Add space
    space_run = etree.SubElement(fn_para, f"{{{_FN_W_NS}}}r")
    space_text = etree.SubElement(space_run, f"{{{_FN_W_NS}}}t")
    space_text.set(f"{{{_FN_XML_NS}}}space", "preserve")
    space_text.text = " "

    # Add content text
    text_run = etree.SubElement(fn_para, f"{{{_FN_W_NS}}}r")
    text_elem = etree.SubElement(text_run, f"{{{_FN_W_NS}}}t")
    text_elem.text = text

    notes_root.append(new_note)

    # Ensure styles exist
    _ensure_note_styles(styles_root, note_type)

    # Ensure content types and relationships
    ct_xml = _ensure_note_content_types(doc_parts["content_types"], note_type)
    rels_xml = _ensure_note_relationship(doc_parts["document_rels"], note_type)

    # Write modified document
    temp_path = doc_path + ".tmp"
    with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        with zipfile.ZipFile(doc_path, "r") as zin:
            for item in zin.infolist():
                if item.filename not in [
                    "word/document.xml",
                    notes_file,
                    "word/styles.xml",
                    "[Content_Types].xml",
                    "word/_rels/document.xml.rels",
                ]:
                    zout.writestr(item, zin.read(item.filename))

        zout.writestr(
            "word/document.xml",
            etree.tostring(
                doc_root, encoding="UTF-8", xml_declaration=True, standalone="yes"
            ),
        )
        zout.writestr(
            notes_file,
            etree.tostring(
                notes_root, encoding="UTF-8", xml_declaration=True, standalone="yes"
            ),
        )
        zout.writestr(
            "word/styles.xml",
            etree.tostring(
                styles_root, encoding="UTF-8", xml_declaration=True, standalone="yes"
            ),
        )
        zout.writestr("[Content_Types].xml", ct_xml)
        zout.writestr("word/_rels/document.xml.rels", rels_xml)

    os.replace(temp_path, doc_path)
    return note_id


def delete_footnote(doc_path: str, note_id: int, note_type: str = "footnote") -> None:
    """Delete a footnote or endnote from a document.

    Args:
        doc_path: Path to the Word document
        note_id: ID of the footnote/endnote to delete
        note_type: "footnote" or "endnote"

    Raises:
        ValueError: If note not found
        FileNotFoundError: If document not found
    """
    import os
    import zipfile

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Document not found: {doc_path}")

    ns = {"w": _FN_W_NS}
    notes_file = f"word/{note_type}s.xml"

    # Read document parts
    with zipfile.ZipFile(doc_path, "r") as zf:
        doc_xml = zf.read("word/document.xml")
        if notes_file not in zf.namelist():
            raise ValueError(f"No {note_type}s in document")
        notes_xml = zf.read(notes_file)

    # Parse XML
    doc_root = etree.fromstring(doc_xml)
    notes_root = etree.fromstring(notes_xml)

    # Remove reference from document
    refs_removed = 0
    for fn_ref in doc_root.xpath(
        f"//w:{note_type}Reference[@w:id='{note_id}']", namespaces=ns
    ):
        run = fn_ref.getparent()
        if run is not None and run.tag == f"{{{_FN_W_NS}}}r":
            para = run.getparent()
            if para is not None:
                para.remove(run)
                refs_removed += 1

    if refs_removed == 0:
        raise ValueError(f"{note_type.title()} {note_id} not found")

    # Remove note content
    for fn in notes_root.xpath(f"//w:{note_type}[@w:id='{note_id}']", namespaces=ns):
        notes_root.remove(fn)

    # Write modified document
    temp_path = doc_path + ".tmp"
    with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        with zipfile.ZipFile(doc_path, "r") as zin:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(
                        item,
                        etree.tostring(
                            doc_root,
                            encoding="UTF-8",
                            xml_declaration=True,
                            standalone="yes",
                        ),
                    )
                elif item.filename == notes_file:
                    zout.writestr(
                        item,
                        etree.tostring(
                            notes_root,
                            encoding="UTF-8",
                            xml_declaration=True,
                            standalone="yes",
                        ),
                    )
                else:
                    zout.writestr(item, zin.read(item.filename))

    os.replace(temp_path, doc_path)


# =============================================================================
# Content Control (SDT) Functions
# =============================================================================


def _get_sdt_type(sdt_pr) -> str:
    """Determine the type of content control from its properties."""
    # Check for specific type elements - w14:checkbox for newer Word versions
    if sdt_pr.find("w14:checkbox", namespaces=_SDT_NSMAP) is not None:
        return "checkbox"
    if sdt_pr.find("w:checkbox", namespaces=_SDT_NSMAP) is not None:
        return "checkbox"
    if sdt_pr.find("w:dropDownList", namespaces=_SDT_NSMAP) is not None:
        return "dropdown"
    if sdt_pr.find("w:comboBox", namespaces=_SDT_NSMAP) is not None:
        return "dropdown"
    if sdt_pr.find("w:date", namespaces=_SDT_NSMAP) is not None:
        return "date"
    if sdt_pr.find("w15:color", namespaces=_SDT_NSMAP) is not None:
        return "color"
    if sdt_pr.find("w:richText", namespaces=_SDT_NSMAP) is not None:
        return "richText"
    if sdt_pr.find("w:text", namespaces=_SDT_NSMAP) is not None:
        return "text"
    # Default to text for unrecognized types
    return "text"


def _get_sdt_value(sdt) -> str:
    """Extract the current value from a content control."""
    sdt_content = sdt.find("w:sdtContent", namespaces=oxml_nsmap)
    if sdt_content is None:
        return ""

    # Collect all text from paragraphs and runs
    texts = []
    for p in sdt_content.iter(qn("w:p")):
        para_texts = []
        for t in p.iter(qn("w:t")):
            if t.text:
                para_texts.append(t.text)
        texts.append("".join(para_texts))

    return "\n".join(texts)


def _get_sdt_checked_state(sdt_pr) -> bool | None:
    """Get checkbox checked state from SDT properties."""
    checkbox = sdt_pr.find("w14:checkbox", namespaces=_SDT_NSMAP)
    if checkbox is None:
        # Try w:checkbox for older format
        checkbox = sdt_pr.find("w:checkbox", namespaces=_SDT_NSMAP)

    if checkbox is not None:
        checked = checkbox.find("w14:checked", namespaces=_SDT_NSMAP)
        if checked is None:
            checked = checkbox.find("w:checked", namespaces=_SDT_NSMAP)
        if checked is not None:
            # w14:val attribute
            ns_w14 = "http://schemas.microsoft.com/office/word/2010/wordml"
            val = checked.get(f"{{{ns_w14}}}val") or checked.get(qn("w:val"))
            return val == "1" or val == "true"
        return False

    return None


def _get_sdt_dropdown_options(sdt_pr) -> list[str]:
    """Get dropdown/combobox options from SDT properties."""
    options = []

    # Check dropDownList
    dropdown = sdt_pr.find("w:dropDownList", namespaces=oxml_nsmap)
    if dropdown is None:
        dropdown = sdt_pr.find("w:comboBox", namespaces=oxml_nsmap)

    if dropdown is not None:
        for list_item in dropdown.findall("w:listItem", namespaces=oxml_nsmap):
            display_text = list_item.get(qn("w:displayText"))
            value = list_item.get(qn("w:value"))
            options.append(display_text or value or "")

    return options


def _get_sdt_date_format(sdt_pr) -> str | None:
    """Get date format from SDT properties."""
    date = sdt_pr.find("w:date", namespaces=oxml_nsmap)
    if date is not None:
        date_format = date.find("w:dateFormat", namespaces=oxml_nsmap)
        if date_format is not None:
            return date_format.get(qn("w:val"))
    return None


def build_content_controls(doc: Document) -> list[dict]:
    """Build list of all content controls (SDTs) in the document."""
    content_controls: list[dict] = []
    block_hash_counts: dict[str, int] = {}

    # Find all SDTs in document body
    body = doc._element.find("w:body", namespaces=oxml_nsmap)
    if body is None:
        return content_controls

    # Track parent paragraph for block_id
    for sdt in body.iter(qn("w:sdt")):
        sdt_pr = sdt.find("w:sdtPr", namespaces=oxml_nsmap)
        if sdt_pr is None:
            continue

        # Get ID
        id_el = sdt_pr.find("w:id", namespaces=oxml_nsmap)
        sdt_id = int(id_el.get(qn("w:val"))) if id_el is not None else 0

        # Get tag
        tag_el = sdt_pr.find("w:tag", namespaces=oxml_nsmap)
        tag = tag_el.get(qn("w:val")) if tag_el is not None else None

        # Get alias
        alias_el = sdt_pr.find("w:alias", namespaces=oxml_nsmap)
        alias = alias_el.get(qn("w:val")) if alias_el is not None else None

        # Determine type
        sdt_type = _get_sdt_type(sdt_pr)

        # Get value
        value = _get_sdt_value(sdt)

        # Get type-specific info
        checked = _get_sdt_checked_state(sdt_pr) if sdt_type == "checkbox" else None
        options = _get_sdt_dropdown_options(sdt_pr) if sdt_type == "dropdown" else []
        date_format = _get_sdt_date_format(sdt_pr) if sdt_type == "date" else None

        # Build block_id (find nearest parent paragraph or use document body)
        parent = sdt.getparent()
        block_id = "document"
        while parent is not None:
            if parent.tag == qn("w:p"):
                # Found parent paragraph
                block_id = build_block_id_from_element(parent, block_hash_counts, {})
                break
            parent = parent.getparent()

        content_controls.append(
            {
                "id": sdt_id,
                "tag": tag,
                "alias": alias,
                "type": sdt_type,
                "value": value,
                "options": options,
                "checked": checked,
                "date_format": date_format,
                "block_id": block_id,
            }
        )

    return content_controls


def build_block_id_from_element(
    element, block_hash_counts: dict[str, int], para_cache: dict
) -> str:
    """Build block ID from an element using consistent ID system with build_blocks().

    Uses content_hash() for normalization and tracks occurrence by block_type + hash.
    """
    # Determine block type - check if it's a heading by looking at pPr/pStyle
    block_type = "paragraph"
    pPr = element.find("w:pPr", namespaces=oxml_nsmap)
    if pPr is not None:
        pStyle = pPr.find("w:pStyle", namespaces=oxml_nsmap)
        if pStyle is not None:
            style_val = pStyle.get(qn("w:val"), "")
            # Check for heading styles (Heading1, Heading 1, heading1, etc.)
            if style_val.lower().replace(" ", "").startswith("heading"):
                try:
                    level_str = (
                        style_val.lower().replace("heading", "").replace(" ", "")
                    )
                    if level_str.isdigit():
                        level = int(level_str)
                        if 1 <= level <= 9:
                            block_type = f"heading{level}"
                except (ValueError, IndexError):
                    pass

    # Extract text content from element
    text_content = ""
    for t in element.iter(qn("w:t")):
        if t.text:
            text_content += t.text

    # Use content_hash for consistent normalization
    text_hash = content_hash(text_content)

    # Track occurrence by block_type + hash (same as build_blocks)
    key = f"{block_type}_{text_hash}"
    occurrence = block_hash_counts.get(key, 0)
    block_hash_counts[key] = occurrence + 1

    return make_block_id(block_type, text_content, occurrence)


def set_content_control_value(doc: Document, sdt_id: int, value: str) -> None:
    """Set the value of a content control.

    For dropdown: value must match one of the options
    For checkbox: value should be "true" or "false"
    For date: value should be ISO date string
    For text: value is the text content
    """
    body = doc._element.find("w:body", namespaces=oxml_nsmap)
    if body is None:
        raise ValueError("Document has no body")

    # Find SDT with matching ID
    for sdt in body.iter(qn("w:sdt")):
        sdt_pr = sdt.find("w:sdtPr", namespaces=oxml_nsmap)
        if sdt_pr is None:
            continue

        id_el = sdt_pr.find("w:id", namespaces=oxml_nsmap)
        if id_el is None:
            continue

        if int(id_el.get(qn("w:val"))) != sdt_id:
            continue

        # Found matching SDT
        sdt_type = _get_sdt_type(sdt_pr)

        if sdt_type == "checkbox":
            _set_checkbox_value(sdt, sdt_pr, value)
        elif sdt_type == "dropdown":
            _set_dropdown_value(sdt, sdt_pr, value)
        else:
            # Text, richText, date, etc.
            _set_text_value(sdt, value)

        return

    raise ValueError(f"Content control with ID {sdt_id} not found")


def _set_checkbox_value(sdt, sdt_pr, value: str) -> None:
    """Set checkbox checked state and update displayed content."""
    is_checked = value.lower() in ("true", "1", "yes")
    checked_val = "1" if is_checked else "0"
    ns_w14 = "http://schemas.microsoft.com/office/word/2010/wordml"

    # Find or create checkbox element
    checkbox = sdt_pr.find("w14:checkbox", namespaces=_SDT_NSMAP)
    if checkbox is None:
        checkbox = sdt_pr.find("w:checkbox", namespaces=_SDT_NSMAP)

    if checkbox is not None:
        # Find or create checked element
        checked = checkbox.find("w14:checked", namespaces=_SDT_NSMAP)
        if checked is None:
            checked = checkbox.find("w:checked", namespaces=_SDT_NSMAP)

        if checked is not None:
            # Update existing - use explicit namespace URI
            val_attr_w14 = f"{{{ns_w14}}}val"
            if val_attr_w14 in checked.attrib:
                checked.set(val_attr_w14, checked_val)
            elif qn("w:val") in checked.attrib:
                checked.set(qn("w:val"), checked_val)
            else:
                checked.set(val_attr_w14, checked_val)
        else:
            # Create new checked element
            checked = OxmlElement("w14:checked")
            checked.set(f"{{{ns_w14}}}val", checked_val)
            checkbox.append(checked)

    # Also update the displayed content (checkbox glyph in w:sdtContent)
    # Unicode checkbox characters: checked = ☒ (U+2612), unchecked = ☐ (U+2610)
    display_char = "\u2612" if is_checked else "\u2610"
    _set_text_value(sdt, display_char)


def _set_dropdown_value(sdt, sdt_pr, value: str) -> None:
    """Set dropdown selected value."""
    # Verify value is in options
    options = _get_sdt_dropdown_options(sdt_pr)
    if options and value not in options:
        raise ValueError(f"Value '{value}' not in dropdown options: {options}")

    # Set the text content
    _set_text_value(sdt, value)


def _set_text_value(sdt, value: str) -> None:
    """Set text content of an SDT."""
    sdt_content = sdt.find("w:sdtContent", namespaces=oxml_nsmap)
    if sdt_content is None:
        return

    # Find first paragraph and set its text
    for p in sdt_content.findall("w:p", namespaces=oxml_nsmap):
        # Clear existing runs
        for r in list(p.findall("w:r", namespaces=oxml_nsmap)):
            p.remove(r)

        # Add new run with text
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = value
        run.append(text)
        p.append(run)

        return  # Only update first paragraph

    # No paragraph found, create one
    p = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = value
    run.append(text)
    p.append(run)
    sdt_content.append(p)


# =============================================================================
# Math Equations (OMML) Functions
# =============================================================================

# Math namespace
_MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_MATH_NSMAP = {"m": _MATH_NS}


def _get_equation_text(omath) -> str:
    """Extract simplified text representation from OMML structure.

    Converts OMML elements to readable text:
    - m:f (fraction) -> (numerator)/(denominator)
    - m:sSup (superscript) -> base^exponent
    - m:sSub (subscript) -> base_subscript
    - m:rad (radical) -> sqrt(radicand) or nrt(degree, radicand)
    - m:nary (n-ary operator) -> sum, prod, int with limits
    - m:d (delimiter/parentheses) -> (content)
    - m:m (matrix) -> [a, b; c, d]
    - m:r (run) -> text content

    Container elements (e, num, den, sub, sup, deg, lim, oMath, oMathPara, fName)
    are handled by recursively processing their children.
    """
    # Container tags that just hold other content - recurse into children
    CONTAINER_TAGS = {
        "e",
        "num",
        "den",
        "sub",
        "sup",
        "deg",
        "lim",
        "oMath",
        "oMathPara",
        "fName",
    }

    parts = []

    for child in omath:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "r":
            # Math run - extract text from m:t
            for t in child.iter(f"{{{_MATH_NS}}}t"):
                if t.text:
                    parts.append(t.text)

        elif tag == "f":
            # Fraction: (num)/(den)
            num = child.find("m:num", namespaces=_MATH_NSMAP)
            den = child.find("m:den", namespaces=_MATH_NSMAP)
            num_text = _get_equation_text(num) if num is not None else "?"
            den_text = _get_equation_text(den) if den is not None else "?"
            parts.append(f"({num_text})/({den_text})")

        elif tag == "sSup":
            # Superscript: base^exp
            base = child.find("m:e", namespaces=_MATH_NSMAP)
            sup = child.find("m:sup", namespaces=_MATH_NSMAP)
            base_text = _get_equation_text(base) if base is not None else "?"
            sup_text = _get_equation_text(sup) if sup is not None else "?"
            parts.append(f"{base_text}^{sup_text}")

        elif tag == "sSub":
            # Subscript: base_sub
            base = child.find("m:e", namespaces=_MATH_NSMAP)
            sub = child.find("m:sub", namespaces=_MATH_NSMAP)
            base_text = _get_equation_text(base) if base is not None else "?"
            sub_text = _get_equation_text(sub) if sub is not None else "?"
            parts.append(f"{base_text}_{sub_text}")

        elif tag == "sSubSup":
            # Sub-superscript: base_sub^sup
            base = child.find("m:e", namespaces=_MATH_NSMAP)
            sub = child.find("m:sub", namespaces=_MATH_NSMAP)
            sup = child.find("m:sup", namespaces=_MATH_NSMAP)
            base_text = _get_equation_text(base) if base is not None else "?"
            sub_text = _get_equation_text(sub) if sub is not None else "?"
            sup_text = _get_equation_text(sup) if sup is not None else "?"
            parts.append(f"{base_text}_{sub_text}^{sup_text}")

        elif tag == "rad":
            # Radical: sqrt(x) or nrt(n, x)
            deg = child.find("m:deg", namespaces=_MATH_NSMAP)
            e = child.find("m:e", namespaces=_MATH_NSMAP)
            e_text = _get_equation_text(e) if e is not None else "?"
            if deg is not None:
                deg_text = _get_equation_text(deg)
                if deg_text and deg_text.strip():
                    parts.append(f"nrt({deg_text}, {e_text})")
                else:
                    parts.append(f"sqrt({e_text})")
            else:
                parts.append(f"sqrt({e_text})")

        elif tag == "nary":
            # N-ary operator (sum, product, integral)
            nary_pr = child.find("m:naryPr", namespaces=_MATH_NSMAP)
            chr_el = (
                nary_pr.find("m:chr", namespaces=_MATH_NSMAP)
                if nary_pr is not None
                else None
            )
            chr_val = chr_el.get(f"{{{_MATH_NS}}}val") if chr_el is not None else None

            # Map Unicode symbols to names
            op_map = {
                "\u2211": "sum",
                "\u220f": "prod",
                "\u222b": "int",
                "\u222c": "iint",
                "\u222d": "iiint",
            }
            op = op_map.get(chr_val, chr_val or "nary")

            sub = child.find("m:sub", namespaces=_MATH_NSMAP)
            sup = child.find("m:sup", namespaces=_MATH_NSMAP)
            e = child.find("m:e", namespaces=_MATH_NSMAP)

            sub_text = _get_equation_text(sub) if sub is not None else ""
            sup_text = _get_equation_text(sup) if sup is not None else ""
            e_text = _get_equation_text(e) if e is not None else "?"

            if sub_text or sup_text:
                parts.append(f"{op}({sub_text}..{sup_text})[{e_text}]")
            else:
                parts.append(f"{op}[{e_text}]")

        elif tag == "d":
            # Delimiter (parentheses, brackets, braces)
            d_pr = child.find("m:dPr", namespaces=_MATH_NSMAP)
            beg_chr = "("
            end_chr = ")"
            if d_pr is not None:
                beg_el = d_pr.find("m:begChr", namespaces=_MATH_NSMAP)
                end_el = d_pr.find("m:endChr", namespaces=_MATH_NSMAP)
                if beg_el is not None:
                    beg_chr = beg_el.get(f"{{{_MATH_NS}}}val") or "("
                if end_el is not None:
                    end_chr = end_el.get(f"{{{_MATH_NS}}}val") or ")"

            e = child.find("m:e", namespaces=_MATH_NSMAP)
            e_text = _get_equation_text(e) if e is not None else ""
            parts.append(f"{beg_chr}{e_text}{end_chr}")

        elif tag == "m":
            # Matrix
            rows = []
            for mr in child.findall("m:mr", namespaces=_MATH_NSMAP):
                row_parts = []
                for me in mr.findall("m:e", namespaces=_MATH_NSMAP):
                    row_parts.append(_get_equation_text(me))
                rows.append(", ".join(row_parts))
            parts.append("[" + "; ".join(rows) + "]")

        elif tag == "func":
            # Function: sin, cos, etc.
            fname = child.find("m:fName", namespaces=_MATH_NSMAP)
            e = child.find("m:e", namespaces=_MATH_NSMAP)
            fname_text = _get_equation_text(fname) if fname is not None else "f"
            e_text = _get_equation_text(e) if e is not None else ""
            parts.append(f"{fname_text}({e_text})")

        elif tag == "limLow":
            # Lower limit: lim_{x->0}
            e = child.find("m:e", namespaces=_MATH_NSMAP)
            lim = child.find("m:lim", namespaces=_MATH_NSMAP)
            e_text = _get_equation_text(e) if e is not None else ""
            lim_text = _get_equation_text(lim) if lim is not None else ""
            parts.append(f"{e_text}_{{{lim_text}}}")

        elif tag == "limUpp":
            # Upper limit
            e = child.find("m:e", namespaces=_MATH_NSMAP)
            lim = child.find("m:lim", namespaces=_MATH_NSMAP)
            e_text = _get_equation_text(e) if e is not None else ""
            lim_text = _get_equation_text(lim) if lim is not None else ""
            parts.append(f"{e_text}^{{{lim_text}}}")

        elif tag in CONTAINER_TAGS:
            # Container elements - recurse into their children
            parts.append(_get_equation_text(child))

        else:
            # Unknown structural element - recurse into children to find text
            child_text = _get_equation_text(child)
            if child_text:
                parts.append(child_text)

    return "".join(parts)


def _get_equation_complexity(omath) -> str:
    """Determine the complexity of a math equation."""
    # Check for complex elements
    has_matrix = omath.find(".//m:m", namespaces=_MATH_NSMAP) is not None
    has_nary = omath.find(".//m:nary", namespaces=_MATH_NSMAP) is not None
    has_fraction = omath.find(".//m:f", namespaces=_MATH_NSMAP) is not None
    has_radical = omath.find(".//m:rad", namespaces=_MATH_NSMAP) is not None
    has_script = (
        omath.find(".//m:sSup", namespaces=_MATH_NSMAP) is not None
        or omath.find(".//m:sSub", namespaces=_MATH_NSMAP) is not None
        or omath.find(".//m:sSubSup", namespaces=_MATH_NSMAP) is not None
    )

    if has_matrix:
        return "matrix"
    if has_nary or (has_fraction and has_radical):
        return "complex"
    if has_fraction:
        return "fraction"
    if has_script or has_radical:
        return "simple"
    return "simple"


def _extract_equations_from_paragraph(
    para, block_id: str, equation_hash_counts: dict[str, int]
) -> list[dict]:
    """Extract all equations from a paragraph element."""
    equations = []
    for omath in para._element.iter(f"{{{_MATH_NS}}}oMath"):
        text = _get_equation_text(omath)
        eq_hash = content_hash(text)
        occurrence = equation_hash_counts.get(eq_hash, 0)
        equation_hash_counts[eq_hash] = occurrence + 1
        eq_id = f"equation_{eq_hash}_{occurrence}"
        complexity = _get_equation_complexity(omath)
        equations.append(
            {
                "id": eq_id,
                "text": text,
                "block_id": block_id,
                "complexity": complexity,
            }
        )
    return equations


def build_equations(doc: Document) -> list[dict]:
    """Build list of all math equations (OMML) in the document."""
    equations: list[dict] = []
    block_hash_counts: dict[str, int] = {}
    equation_hash_counts: dict[str, int] = {}

    # Use iter_body_blocks to match build_blocks() block ID computation
    for kind, obj, _el in iter_body_blocks(doc):
        if kind == "paragraph":
            # Use SAME logic as build_blocks for block_id
            block_type, _ = paragraph_kind_and_level(obj)
            text = obj.text or ""
            block_hash_key = f"{block_type}_{content_hash(text)}"
            block_occurrence = block_hash_counts.get(block_hash_key, 0)
            block_id = make_block_id(block_type, text, block_occurrence)
            block_hash_counts[block_hash_key] = block_occurrence + 1

            equations.extend(
                _extract_equations_from_paragraph(obj, block_id, equation_hash_counts)
            )

        elif kind == "table":
            # Compute table's block_id
            table_content = table_content_for_hash(obj)
            block_hash_key = f"table_{content_hash(table_content)}"
            block_occurrence = block_hash_counts.get(block_hash_key, 0)
            table_block_id = make_block_id("table", table_content, block_occurrence)
            block_hash_counts[block_hash_key] = block_occurrence + 1

            # Search all cells for equations with hierarchical block_id
            rows, cols = len(obj.rows), len(obj.columns)
            for r in range(rows):
                for c in range(cols):
                    cell = obj.cell(r, c)
                    for p_idx, para in enumerate(cell.paragraphs):
                        hier_block_id = f"{table_block_id}#r{r}c{c}/p{p_idx}"
                        equations.extend(
                            _extract_equations_from_paragraph(
                                para, hier_block_id, equation_hash_counts
                            )
                        )

    return equations
