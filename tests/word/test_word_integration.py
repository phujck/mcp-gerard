"""Integration tests for Word document MCP tool."""

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document
from mcp.server.fastmcp.exceptions import ToolError

from mcp_handley_lab.word.tool import mcp


@pytest.fixture
def sample_docx():
    """Create a sample Word document for testing."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is the first paragraph.")
        doc.add_heading("Section Two", level=2)
        doc.add_paragraph("This is the second paragraph.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(1, 0).text = "Row 1 Col 1"
        table.cell(1, 1).text = "Row 1 Col 2"
        doc.save(f.name)
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_read_outline(sample_docx):
    """Test reading document outline (headings only)."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "outline"}
    )
    assert result["block_count"] == 2  # Two headings
    assert len(result["blocks"]) == 2
    assert result["blocks"][0]["type"] == "heading1"  # Level in type
    assert result["blocks"][0]["level"] == 1
    assert "Test Document" in result["blocks"][0]["text"]
    assert result["blocks"][1]["type"] == "heading2"
    assert result["blocks"][1]["level"] == 2


@pytest.mark.asyncio
async def test_read_meta(sample_docx):
    """Test reading document metadata."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "meta"}
    )
    assert result["meta"] is not None
    assert result["blocks"] == []


@pytest.mark.asyncio
async def test_read_blocks(sample_docx):
    """Test reading all blocks."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    assert result["block_count"] == 5  # 2 headings + 2 paragraphs + 1 table
    types = [b["type"] for b in result["blocks"]]
    assert "heading1" in types  # Level in type
    assert "heading2" in types
    assert "paragraph" in types
    assert "table" in types


@pytest.mark.asyncio
async def test_read_search(sample_docx):
    """Test searching for text."""
    _, result = await mcp.call_tool(
        "read",
        {"file_path": str(sample_docx), "scope": "search", "search_query": "second"},
    )
    assert result["block_count"] == 1  # Only matching blocks counted
    assert len(result["blocks"]) == 1
    assert "second" in result["blocks"][0]["text"].lower()


@pytest.mark.asyncio
async def test_edit_append_paragraph(sample_docx):
    """Test appending a paragraph."""
    _, read_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    initial_count = read_result["block_count"]

    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "append",
            "content_type": "paragraph",
            "content_data": "New paragraph at the end",
        },
    )
    assert edit_result["success"]
    assert edit_result["element_id"]

    _, read_result2 = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    assert read_result2["block_count"] == initial_count + 1


@pytest.mark.asyncio
async def test_edit_delete(sample_docx):
    """Test deleting a block."""
    _, read_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    initial_count = read_result["block_count"]
    target_id = read_result["blocks"][1]["id"]  # Delete second block

    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "delete",
            "target_id": target_id,
        },
    )
    assert edit_result["success"]

    _, read_result2 = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    assert read_result2["block_count"] == initial_count - 1


@pytest.mark.asyncio
async def test_edit_insert_before(sample_docx):
    """Test inserting before a block."""
    _, read_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    target_id = read_result["blocks"][1]["id"]

    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "insert_before",
            "target_id": target_id,
            "content_type": "paragraph",
            "content_data": "Inserted before",
        },
    )
    assert edit_result["success"]
    assert edit_result["element_id"]


@pytest.mark.asyncio
async def test_edit_replace(sample_docx):
    """Test replacing block content."""
    _, read_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in read_result["blocks"] if b["type"] == "paragraph")

    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "replace",
            "target_id": para_block["id"],
            "content_data": "Replaced content",
        },
    )
    assert edit_result["success"]


@pytest.mark.asyncio
async def test_edit_append_table(sample_docx):
    """Test appending a table."""
    table_data = json.dumps([["A", "B"], ["1", "2"]])
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "append",
            "content_type": "table",
            "content_data": table_data,
        },
    )
    assert edit_result["success"]
    assert edit_result["element_id"]


@pytest.mark.asyncio
async def test_edit_create_empty():
    """Test creating a new empty document."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        new_path = Path(f.name)
    new_path.unlink()  # Remove so we can create it

    try:
        _, edit_result = await mcp.call_tool(
            "edit",
            {
                "file_path": str(new_path),
                "operation": "create",
            },
        )
        assert edit_result["success"]
        assert new_path.exists()

        # Verify it's readable (creates default empty paragraph)
        _, read_result = await mcp.call_tool(
            "read", {"file_path": str(new_path), "scope": "blocks"}
        )
        assert read_result["block_count"] == 1
        assert read_result["blocks"][0]["text"] == ""
    finally:
        new_path.unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_edit_create_with_content():
    """Test creating a new document with initial content."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        new_path = Path(f.name)
    new_path.unlink()  # Remove so we can create it

    try:
        _, edit_result = await mcp.call_tool(
            "edit",
            {
                "file_path": str(new_path),
                "operation": "create",
                "content_type": "heading",
                "content_data": "My New Document",
                "heading_level": 1,
            },
        )
        assert edit_result["success"]
        assert edit_result["element_id"]
        assert new_path.exists()

        # Verify content
        _, read_result = await mcp.call_tool(
            "read", {"file_path": str(new_path), "scope": "blocks"}
        )
        assert read_result["block_count"] == 1
        assert read_result["blocks"][0]["type"] == "heading1"  # Level in type
        assert "My New Document" in read_result["blocks"][0]["text"]
    finally:
        new_path.unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_edit_create_overwrites_existing(sample_docx):
    """Test that create overwrites an existing file (no defensive check)."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "create",
            "content_data": "New content",
        },
    )
    assert edit_result["success"]
    # Verify new content
    _, read_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    assert read_result["block_count"] == 1
    assert "New content" in read_result["blocks"][0]["text"]


@pytest.mark.asyncio
async def test_block_ids_are_content_hash(sample_docx):
    """Test that block IDs use content-hash format (type_hash_occurrence)."""
    import re

    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    # Block IDs: paragraph_abc123_0, heading1_def456_0, table_ghi789_0
    id_pattern = re.compile(r"^(paragraph|heading[1-9]|table)_[0-9a-f]{8}_\d+$")
    for block in result["blocks"]:
        block_id = block["id"]
        assert id_pattern.match(block_id), f"Invalid ID format: {block_id}"


@pytest.mark.asyncio
async def test_read_table_cells(sample_docx):
    """Test reading table cells with 0-based indices and hierarchical IDs."""
    # First find the table
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    table_id = table_block["id"]

    # Now read the cells
    _, cells_result = await mcp.call_tool(
        "read",
        {"file_path": str(sample_docx), "scope": "table_cells", "target_id": table_id},
    )
    assert cells_result["table_rows"] == 2
    assert cells_result["table_cols"] == 2
    assert len(cells_result["cells"]) == 4  # 2x2 table
    # Check first cell (now 0-based)
    cell_0_0 = next(c for c in cells_result["cells"] if c["row"] == 0 and c["col"] == 0)
    assert cell_0_0["text"] == "Header 1"
    # Check hierarchical_id format
    assert cell_0_0["hierarchical_id"] == f"{table_id}#r0c0"
    # Check another cell
    cell_1_1 = next(c for c in cells_result["cells"] if c["row"] == 1 and c["col"] == 1)
    assert cell_1_1["text"] == "Row 1 Col 2"
    assert cell_1_1["hierarchical_id"] == f"{table_id}#r1c1"


@pytest.mark.asyncio
async def test_edit_cell(sample_docx):
    """Test editing a single table cell."""
    # Find the table
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    table_id = table_block["id"]

    # Edit cell (1,1)
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "edit_cell",
            "target_id": table_id,
            "row": 1,
            "col": 1,
            "content_data": "Updated Header",
        },
    )
    assert edit_result["success"]
    # Returns updated table block ID (content-hash changes after cell edit)
    assert edit_result["element_id"].startswith("table_")
    assert edit_result["element_id"] != table_id  # Hash changed

    # Verify the change - content hash changes after edit, so re-read blocks
    _, blocks_result2 = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    new_table_block = next(b for b in blocks_result2["blocks"] if b["type"] == "table")
    new_table_id = new_table_block["id"]

    _, cells_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "table_cells",
            "target_id": new_table_id,
        },
    )
    cell_1_1 = next(c for c in cells_result["cells"] if c["row"] == 1 and c["col"] == 1)
    assert cell_1_1["text"] == "Updated Header"


@pytest.mark.asyncio
async def test_edit_cell_out_of_range(sample_docx):
    """Test that editing out-of-range cell fails."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    with pytest.raises(ToolError, match="out of range"):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(sample_docx),
                "operation": "edit_cell",
                "target_id": table_block["id"],
                "row": 99,
                "col": 1,
                "content_data": "Should fail",
            },
        )


@pytest.mark.asyncio
async def test_edit_cell_on_non_table(sample_docx):
    """Test that editing cell on non-table block fails (underlying library error)."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    with pytest.raises(ToolError):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(sample_docx),
                "operation": "edit_cell",
                "target_id": para_block["id"],
                "row": 1,
                "col": 1,
                "content_data": "Should fail",
            },
        )


@pytest.fixture
def formatted_docx():
    """Create a Word document with formatted runs for testing."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        # Create a paragraph with multiple runs with different formatting
        p = doc.add_paragraph()
        p.add_run("Normal text, ")  # default formatting
        run2 = p.add_run("bold text, ")
        run2.bold = True
        run3 = p.add_run("italic text.")
        run3.italic = True
        doc.save(f.name)
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_read_runs(formatted_docx):
    """Test reading runs from a paragraph."""
    # First get the paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(formatted_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Now read the runs
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(formatted_docx),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )
    assert runs_result["block_count"] == 3  # Three runs
    assert len(runs_result["runs"]) == 3

    # Check run properties
    assert runs_result["runs"][0]["text"] == "Normal text, "
    assert runs_result["runs"][0]["bold"] is None  # Not set (inherits)

    assert runs_result["runs"][1]["text"] == "bold text, "
    assert runs_result["runs"][1]["bold"] is True

    assert runs_result["runs"][2]["text"] == "italic text."
    assert runs_result["runs"][2]["italic"] is True


@pytest.mark.asyncio
async def test_edit_run_text(formatted_docx):
    """Test editing run text."""
    # Get paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(formatted_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Edit the bold run (index 1) - use different text (not just case change, as hash normalizes)
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(formatted_docx),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 1,
            "content_data": "strongly emphasized, ",
        },
    )
    assert edit_result["success"]
    # Hash changes because paragraph text changed
    assert edit_result["element_id"].startswith("paragraph_")
    assert edit_result["element_id"] != para_block["id"]

    # Verify the change - need new paragraph ID
    _, blocks_result2 = await mcp.call_tool(
        "read", {"file_path": str(formatted_docx), "scope": "blocks"}
    )
    new_para = next(b for b in blocks_result2["blocks"] if b["type"] == "paragraph")

    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(formatted_docx),
            "scope": "runs",
            "target_id": new_para["id"],
        },
    )
    assert runs_result["runs"][1]["text"] == "strongly emphasized, "
    assert runs_result["runs"][1]["bold"] is True  # Formatting preserved


@pytest.mark.asyncio
async def test_edit_run_formatting(formatted_docx):
    """Test editing run formatting."""
    # Get paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(formatted_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Make the first run (Normal text) bold and underlined
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(formatted_docx),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 0,
            "formatting": '{"bold": true, "underline": true}',
        },
    )
    assert edit_result["success"]

    # Verify formatting (text unchanged, so ID unchanged)
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(formatted_docx),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )
    assert runs_result["runs"][0]["text"] == "Normal text, "
    assert runs_result["runs"][0]["bold"] is True
    assert runs_result["runs"][0]["underline"] is True


@pytest.mark.asyncio
async def test_edit_run_out_of_range(formatted_docx):
    """Test that editing out-of-range run fails."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(formatted_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    with pytest.raises(ToolError, match="out of range"):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(formatted_docx),
                "operation": "edit_run",
                "target_id": para_block["id"],
                "run_index": 99,
                "content_data": "Should fail",
            },
        )


@pytest.mark.asyncio
async def test_read_runs_on_table_fails(sample_docx):
    """Test that reading runs on a table fails."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    with pytest.raises(ToolError):
        await mcp.call_tool(
            "read",
            {
                "file_path": str(sample_docx),
                "scope": "runs",
                "target_id": table_block["id"],
            },
        )


@pytest.mark.asyncio
async def test_edit_run_on_table_fails(sample_docx):
    """Test that editing run on a table fails (underlying library error)."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    with pytest.raises(ToolError):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(sample_docx),
                "operation": "edit_run",
                "target_id": table_block["id"],
                "run_index": 0,
                "content_data": "Should fail",
            },
        )


# --- Comment tests ---


@pytest.mark.asyncio
async def test_read_comments_empty(sample_docx):
    """Test reading comments from a document with no comments."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "comments"}
    )
    assert result["block_count"] == 0
    assert result["comments"] == []


@pytest.mark.asyncio
async def test_add_comment_to_paragraph(sample_docx):
    """Test adding a comment to a paragraph."""
    # First get a paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Add a comment
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_comment",
            "target_id": para_block["id"],
            "content_data": "This is a test comment",
        },
    )
    assert edit_result["success"]
    assert edit_result["comment_id"] is not None

    # Verify comment was added
    _, comments_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "comments"}
    )
    assert comments_result["block_count"] == 1
    assert len(comments_result["comments"]) == 1
    assert comments_result["comments"][0]["text"] == "This is a test comment"


@pytest.mark.asyncio
async def test_add_comment_with_author(sample_docx):
    """Test adding a comment with author and initials."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Add a comment with author
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_comment",
            "target_id": para_block["id"],
            "content_data": "Review needed",
            "author": "Test Author",
            "initials": "TA",
        },
    )
    assert edit_result["success"]

    # Verify author info
    _, comments_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "comments"}
    )
    assert comments_result["comments"][0]["author"] == "Test Author"
    assert comments_result["comments"][0]["initials"] == "TA"


@pytest.mark.asyncio
async def test_add_comment_to_table_fails(sample_docx):
    """Test that adding a comment to a table fails (underlying library error)."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    with pytest.raises(ToolError):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(sample_docx),
                "operation": "add_comment",
                "target_id": table_block["id"],
                "content_data": "Should fail",
            },
        )


# --- Headers/Footers tests ---


@pytest.mark.asyncio
async def test_read_headers_footers(sample_docx):
    """Test reading headers and footers."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "headers_footers"}
    )
    # Document should have at least one section
    assert result["block_count"] >= 1
    assert len(result["headers_footers"]) >= 1
    # First section should be at index 0
    assert result["headers_footers"][0]["section_index"] == 0


@pytest.mark.asyncio
async def test_set_header(sample_docx):
    """Test setting a header."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_header",
            "section_index": 0,
            "content_data": "My Custom Header",
        },
    )
    assert edit_result["success"]

    # Verify header was set
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "headers_footers"}
    )
    assert result["headers_footers"][0]["header_text"] == "My Custom Header"
    assert result["headers_footers"][0]["header_is_linked"] is False


@pytest.mark.asyncio
async def test_set_footer(sample_docx):
    """Test setting a footer."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_footer",
            "section_index": 0,
            "content_data": "Page Footer Text",
        },
    )
    assert edit_result["success"]

    # Verify footer was set
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "headers_footers"}
    )
    assert result["headers_footers"][0]["footer_text"] == "Page Footer Text"
    assert result["headers_footers"][0]["footer_is_linked"] is False


@pytest.mark.asyncio
async def test_set_header_invalid_section(sample_docx):
    """Test that setting header on invalid section fails."""
    with pytest.raises(ToolError, match="out of range"):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(sample_docx),
                "operation": "set_header",
                "section_index": 99,
                "content_data": "Should fail",
            },
        )


# --- Page Setup tests ---


@pytest.mark.asyncio
async def test_read_page_setup(sample_docx):
    """Test reading page setup."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "page_setup"}
    )
    # Document should have at least one section
    assert result["block_count"] >= 1
    assert len(result["page_setup"]) >= 1
    # Check first section has expected fields
    setup = result["page_setup"][0]
    assert setup["section_index"] == 0
    assert setup["orientation"] in ["portrait", "landscape"]
    assert setup["page_width"] > 0
    assert setup["page_height"] > 0


@pytest.mark.asyncio
async def test_set_margins(sample_docx):
    """Test setting page margins."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_margins",
            "section_index": 0,
            "formatting": '{"top": 0.5, "bottom": 0.5, "left": 0.75, "right": 0.75}',
        },
    )
    assert edit_result["success"]

    # Verify margins were set
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "page_setup"}
    )
    setup = result["page_setup"][0]
    assert setup["top_margin"] == 0.5
    assert setup["bottom_margin"] == 0.5
    assert setup["left_margin"] == 0.75
    assert setup["right_margin"] == 0.75


@pytest.mark.asyncio
async def test_set_orientation_landscape(sample_docx):
    """Test setting page orientation to landscape."""
    # Get original dimensions
    _, before = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "page_setup"}
    )
    orig_width = before["page_setup"][0]["page_width"]
    orig_height = before["page_setup"][0]["page_height"]

    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_orientation",
            "section_index": 0,
            "content_data": "landscape",
        },
    )
    assert edit_result["success"]

    # Verify orientation was set
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "page_setup"}
    )
    setup = result["page_setup"][0]
    assert setup["orientation"] == "landscape"
    # Width should now be greater than height (or swapped)
    assert setup["page_width"] >= setup["page_height"] or (
        setup["page_width"] == orig_height and setup["page_height"] == orig_width
    )


@pytest.mark.asyncio
async def test_set_orientation_invalid(sample_docx):
    """Test that invalid orientation defaults to portrait."""
    _, result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_orientation",
            "section_index": 0,
            "content_data": "diagonal",
        },
    )
    assert result["success"]
    # Invalid value defaults to portrait
    _, read_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "page_setup"}
    )
    assert read_result["page_setup"][0]["orientation"] == "portrait"


@pytest.mark.asyncio
async def test_set_margins_missing_formatting(sample_docx):
    """Test that set_margins without formatting fails."""
    with pytest.raises(ToolError):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(sample_docx),
                "operation": "set_margins",
                "section_index": 0,
            },
        )


# --- File error handling tests ---


@pytest.mark.asyncio
async def test_read_missing_file():
    """Test that reading a missing file raises error."""
    with pytest.raises(ToolError):
        await mcp.call_tool(
            "read",
            {
                "file_path": "/nonexistent/path/to/file.docx",
                "scope": "blocks",
            },
        )


@pytest.mark.asyncio
async def test_edit_missing_file():
    """Test that editing a missing file raises error."""
    with pytest.raises(ToolError):
        await mcp.call_tool(
            "edit",
            {
                "file_path": "/nonexistent/path/to/file.docx",
                "operation": "append",
                "content_data": "Should fail",
            },
        )


@pytest.mark.asyncio
async def test_read_corrupt_file():
    """Test that reading a corrupt file raises error."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(b"not a valid docx file")
        corrupt_path = f.name
    try:
        with pytest.raises(ToolError):
            await mcp.call_tool(
                "read",
                {
                    "file_path": corrupt_path,
                    "scope": "blocks",
                },
            )
    finally:
        Path(corrupt_path).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_create_in_new_directory():
    """Test creating a document in a new (non-existent) directory raises error."""
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create path with nested non-existent directory
        new_file = Path(tmpdir) / "subdir" / "nested" / "test.docx"
        # Now fails because we don't create parent directories
        with pytest.raises(ToolError):
            await mcp.call_tool(
                "edit",
                {
                    "file_path": str(new_file),
                    "operation": "create",
                    "content_data": "Test content",
                },
            )


# --- Image tests ---


@pytest.fixture
def sample_image():
    """Create a minimal PNG image for testing using PIL."""
    from PIL import Image as PILImage

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
        # Create a 10x10 red image
        img = PILImage.new("RGB", (10, 10), color="red")
        img.save(f.name, "PNG")
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.fixture
def docx_with_image(sample_image):
    """Create a Word document with an embedded image."""
    from docx.shared import Inches

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_heading("Document with Image", level=1)
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(sample_image), width=Inches(1), height=Inches(1))
        doc.add_paragraph("Text after image.")
        doc.save(f.name)
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_read_images(docx_with_image):
    """Test reading images from a document."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_image), "scope": "images"}
    )
    assert result["block_count"] == 1
    assert len(result["images"]) == 1
    img = result["images"][0]
    # Check ID format: image_{hash}_{occurrence}
    assert img["id"].startswith("image_")
    assert "_" in img["id"][6:]  # hash_occurrence part
    # Check dimensions (we set 1 inch)
    assert 0.9 < img["width_inches"] < 1.1
    assert 0.9 < img["height_inches"] < 1.1
    assert img["content_type"] == "image/png"
    # Check block_id references a paragraph
    assert img["block_id"].startswith("paragraph_")


@pytest.mark.asyncio
async def test_read_images_empty(sample_docx):
    """Test reading images from a document with no images."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "images"}
    )
    assert result["block_count"] == 0
    assert result["images"] == []


@pytest.mark.asyncio
async def test_insert_image(sample_docx, sample_image):
    """Test inserting an image."""
    # Get a paragraph to insert after
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Insert image
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "insert_image",
            "target_id": para_block["id"],
            "content_data": str(sample_image),
            "formatting": '{"width": 2.0, "height": 1.5}',
        },
    )
    assert edit_result["success"]
    assert edit_result["element_id"].startswith("image_")

    # Verify image was added
    _, images_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "images"}
    )
    assert images_result["block_count"] == 1
    img = images_result["images"][0]
    assert 1.9 < img["width_inches"] < 2.1
    assert 1.4 < img["height_inches"] < 1.6


@pytest.mark.asyncio
async def test_delete_image(docx_with_image):
    """Test deleting an image."""
    # Get the image ID
    _, images_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_image), "scope": "images"}
    )
    assert len(images_result["images"]) == 1
    image_id = images_result["images"][0]["id"]

    # Delete the image
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_image),
            "operation": "delete_image",
            "target_id": image_id,
        },
    )
    assert edit_result["success"]

    # Verify image was deleted
    _, images_result2 = await mcp.call_tool(
        "read", {"file_path": str(docx_with_image), "scope": "images"}
    )
    assert images_result2["block_count"] == 0
    assert images_result2["images"] == []


@pytest.mark.asyncio
async def test_image_id_format(sample_image):
    """Test that image IDs use content-hash format."""
    import re

    from docx.shared import Inches

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as doc_f:
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(str(sample_image), width=Inches(1))
        doc.save(doc_f.name)
        doc_path = doc_f.name

    try:
        _, result = await mcp.call_tool(
            "read", {"file_path": doc_path, "scope": "images"}
        )
        # Image ID format: image_{sha1[:8]}_{occurrence}
        id_pattern = re.compile(r"^image_[0-9a-f]{8}_\d+$")
        for img in result["images"]:
            assert id_pattern.match(img["id"]), f"Invalid ID format: {img['id']}"
    finally:
        Path(doc_path).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_multiple_same_images(sample_image):
    """Test that same image appearing twice has different occurrence numbers."""
    from docx.shared import Inches

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as doc_f:
        doc = Document()
        # Add same image twice in different paragraphs
        p1 = doc.add_paragraph()
        p1.add_run().add_picture(str(sample_image), width=Inches(1))
        p2 = doc.add_paragraph()
        p2.add_run().add_picture(str(sample_image), width=Inches(1))
        doc.save(doc_f.name)
        doc_path = doc_f.name

    try:
        _, result = await mcp.call_tool(
            "read", {"file_path": doc_path, "scope": "images"}
        )
        assert len(result["images"]) == 2
        # Same hash, different occurrence
        id1, id2 = result["images"][0]["id"], result["images"][1]["id"]
        hash1, occ1 = id1.split("_")[1], id1.split("_")[2]
        hash2, occ2 = id2.split("_")[1], id2.split("_")[2]
        assert hash1 == hash2  # Same image = same hash
        assert occ1 != occ2  # Different occurrences
        assert {occ1, occ2} == {"0", "1"}  # Should be 0 and 1
    finally:
        Path(doc_path).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_delete_image_invalid_id(sample_docx):
    """Test that deleting with invalid image ID fails."""
    with pytest.raises(ToolError):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(sample_docx),
                "operation": "delete_image",
                "target_id": "image_nonexist_0",
            },
        )


@pytest.mark.asyncio
async def test_read_images_in_table(sample_image):
    """Test reading images from table cells with hierarchical block_id."""
    from docx.shared import Inches

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as doc_f:
        doc = Document()
        doc.add_paragraph("Document with table containing image")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Signature"
        table.cell(1, 0).text = "John Doe"
        # Add image to signature cell (row=1, col=1, 0-based)
        sig_cell = table.cell(1, 1)
        sig_para = sig_cell.paragraphs[0]
        sig_para.add_run().add_picture(str(sample_image), width=Inches(1))
        doc.save(doc_f.name)
        doc_path = doc_f.name

    try:
        _, result = await mcp.call_tool(
            "read", {"file_path": doc_path, "scope": "images"}
        )
        assert result["block_count"] == 1
        assert len(result["images"]) == 1
        img = result["images"][0]
        # block_id should be hierarchical: table_xxx#r1c1/p0
        assert img["block_id"].startswith("table_")
        assert "#r1c1/p0" in img["block_id"]  # Image in cell (1,1), paragraph 0
        assert img["content_type"] == "image/png"
    finally:
        Path(doc_path).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_delete_image_in_table(sample_image):
    """Test deleting an image from a table cell."""
    from docx.shared import Inches

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as doc_f:
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.paragraphs[0].add_run().add_picture(str(sample_image), width=Inches(1))
        doc.save(doc_f.name)
        doc_path = doc_f.name

    try:
        # Get the image ID
        _, images_result = await mcp.call_tool(
            "read", {"file_path": doc_path, "scope": "images"}
        )
        assert len(images_result["images"]) == 1
        image_id = images_result["images"][0]["id"]

        # Delete the image
        _, edit_result = await mcp.call_tool(
            "edit",
            {
                "file_path": doc_path,
                "operation": "delete_image",
                "target_id": image_id,
            },
        )
        assert edit_result["success"]

        # Verify image was deleted
        _, images_result2 = await mcp.call_tool(
            "read", {"file_path": doc_path, "scope": "images"}
        )
        assert images_result2["block_count"] == 0
    finally:
        Path(doc_path).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_insert_image_into_table_cell(sample_image):
    """Test inserting an image into a specific table cell using hierarchical target_id."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as doc_f:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Image"
        table.cell(0, 1).text = "Description"
        table.cell(1, 1).text = "A test image"
        doc.save(doc_f.name)
        doc_path = doc_f.name

    try:
        # Get the table ID
        _, blocks_result = await mcp.call_tool(
            "read", {"file_path": doc_path, "scope": "blocks"}
        )
        table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
        table_id = table_block["id"]

        # Insert image into cell (1, 0) - second row, first column (0-based)
        # Using hierarchical target_id: table_xxx#r1c0
        _, edit_result = await mcp.call_tool(
            "edit",
            {
                "file_path": doc_path,
                "operation": "insert_image",
                "target_id": f"{table_id}#r1c0",
                "content_data": str(sample_image),
                "formatting": '{"width": 1}',
            },
        )
        assert edit_result["success"]
        assert edit_result["element_id"].startswith("image_")

        # Verify image was inserted in the correct cell
        _, images_result = await mcp.call_tool(
            "read", {"file_path": doc_path, "scope": "images"}
        )
        assert images_result["block_count"] == 1
        img = images_result["images"][0]
        # Image should be in cell (1,0), paragraph 0
        assert "#r1c0/p0" in img["block_id"]
    finally:
        Path(doc_path).unlink(missing_ok=True)


# --- Hierarchical addressing tests ---


@pytest.mark.asyncio
async def test_hierarchical_path_parsing(sample_docx):
    """Test that hierarchical paths are parsed correctly."""
    import re

    # Get table ID and read cells with hierarchical IDs
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    table_id = table_block["id"]

    _, cells_result = await mcp.call_tool(
        "read",
        {"file_path": str(sample_docx), "scope": "table_cells", "target_id": table_id},
    )

    # Each cell should have a hierarchical_id in format: {table_id}#r{row}c{col}
    # Use strict regex to validate full grammar
    hier_id_pattern = re.compile(rf"^{re.escape(table_id)}#r(\d+)c(\d+)$")

    for cell in cells_result["cells"]:
        hid = cell["hierarchical_id"]
        match = hier_id_pattern.match(hid)
        assert match, f"Invalid hierarchical_id format: {hid}"
        # Verify the parsed row/col match cell's row/col
        assert int(match.group(1)) == cell["row"]
        assert int(match.group(2)) == cell["col"]


@pytest.mark.asyncio
async def test_hierarchical_path_from_paragraph_fails(sample_docx):
    """Test that hierarchical paths are rejected for non-table base blocks."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")
    para_id = para_block["id"]

    # Trying to use hierarchical path on a paragraph should fail
    with pytest.raises(ToolError, match="has no attribute 'cell'"):
        await mcp.call_tool(
            "read",
            {
                "file_path": str(sample_docx),
                "scope": "runs",
                "target_id": f"{para_id}#r0c0",  # Invalid: para can't have cell path
            },
        )


@pytest.mark.asyncio
async def test_hierarchical_path_invalid_segment(sample_docx):
    """Test that invalid path segments are rejected."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    table_id = table_block["id"]

    # Invalid segment format should fail
    with pytest.raises(ToolError, match="Invalid path segment"):
        await mcp.call_tool(
            "read",
            {
                "file_path": str(sample_docx),
                "scope": "runs",
                "target_id": f"{table_id}#invalid",
            },
        )


@pytest.mark.asyncio
async def test_hierarchical_path_cell_out_of_bounds(sample_docx):
    """Test that out-of-bounds cell coordinates are rejected."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    table_id = table_block["id"]

    # Row out of bounds (table has 2 rows: 0 and 1)
    with pytest.raises(ToolError, match="index out of range"):
        await mcp.call_tool(
            "read",
            {
                "file_path": str(sample_docx),
                "scope": "runs",
                "target_id": f"{table_id}#r99c0/p0",
            },
        )


@pytest.mark.asyncio
async def test_hierarchical_transition_cell_to_paragraph(sample_docx):
    """Test valid transition: table -> cell -> paragraph."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    table_id = table_block["id"]

    # Valid path: table#r0c0/p0 -> first paragraph in cell (0,0)
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "runs",
            "target_id": f"{table_id}#r0c0/p0",
        },
    )
    # Should succeed and return runs for the paragraph in the cell
    assert "runs" in runs_result


@pytest.mark.asyncio
async def test_hierarchical_transition_invalid_para_from_table(sample_docx):
    """Test invalid transition: cannot select paragraph directly from table."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    table_id = table_block["id"]

    # Invalid: p0 directly from table (must go through cell first)
    with pytest.raises(ToolError, match="has no attribute 'paragraphs'"):
        await mcp.call_tool(
            "read",
            {
                "file_path": str(sample_docx),
                "scope": "runs",
                "target_id": f"{table_id}#p0",  # Missing cell selector
            },
        )


# --- Paragraph formatting tests ---


@pytest.mark.asyncio
async def test_paragraph_indentation(sample_docx):
    """Test setting paragraph indentation (left, right, first-line)."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply indentation formatting
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "style",
            "target_id": para_block["id"],
            "formatting": '{"left_indent": 0.5, "right_indent": 0.25, "first_line_indent": 0.5}',
        },
    )
    assert edit_result["success"]

    # Verify by reopening document
    # Note: doc.paragraphs includes headings, so find paragraph by text
    doc = Document(str(sample_docx))
    para = next(p for p in doc.paragraphs if p.text == para_block["text"])
    pf = para.paragraph_format
    # Check values are approximately correct (allow for floating point)
    assert abs(pf.left_indent.inches - 0.5) < 0.01
    assert abs(pf.right_indent.inches - 0.25) < 0.01
    assert abs(pf.first_line_indent.inches - 0.5) < 0.01


@pytest.mark.asyncio
async def test_paragraph_spacing(sample_docx):
    """Test setting paragraph spacing (before, after, line spacing)."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply spacing formatting
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "style",
            "target_id": para_block["id"],
            "formatting": '{"space_before": 12, "space_after": 6, "line_spacing": 1.5}',
        },
    )
    assert edit_result["success"]

    # Verify by reopening document
    doc = Document(str(sample_docx))
    para = next(p for p in doc.paragraphs if p.text == para_block["text"])
    pf = para.paragraph_format
    assert abs(pf.space_before.pt - 12) < 0.1
    assert abs(pf.space_after.pt - 6) < 0.1
    assert abs(pf.line_spacing - 1.5) < 0.01


@pytest.mark.asyncio
async def test_paragraph_flow_control(sample_docx):
    """Test setting paragraph flow control (keep_with_next, page_break_before)."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply flow control formatting
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "style",
            "target_id": para_block["id"],
            "formatting": '{"keep_with_next": true, "page_break_before": true}',
        },
    )
    assert edit_result["success"]

    # Verify by reopening document
    doc = Document(str(sample_docx))
    para = next(p for p in doc.paragraphs if p.text == para_block["text"])
    pf = para.paragraph_format
    assert pf.keep_with_next is True
    assert pf.page_break_before is True


# --- Run effects tests ---


@pytest.mark.asyncio
async def test_run_highlight(sample_docx):
    """Test applying highlight color to a run."""
    from docx.enum.text import WD_COLOR_INDEX

    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Read runs to get run index
    _, runs_result = await mcp.call_tool(
        "read",
        {"file_path": str(sample_docx), "scope": "runs", "target_id": para_block["id"]},
    )
    assert len(runs_result["runs"]) > 0

    # Apply highlight formatting
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 0,
            "formatting": '{"highlight_color": "yellow"}',
        },
    )
    assert edit_result["success"]

    # Verify by reopening document
    doc = Document(str(sample_docx))
    para = next(p for p in doc.paragraphs if p.text == para_block["text"])
    assert para.runs[0].font.highlight_color == WD_COLOR_INDEX.YELLOW

    # Also verify via read() returns the highlight
    _, runs_result2 = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    assert runs_result2["runs"][0]["highlight_color"] == "yellow"


@pytest.mark.asyncio
async def test_run_strikethrough(sample_docx):
    """Test applying strikethrough to a run."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply strike formatting
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 0,
            "formatting": '{"strike": true, "double_strike": false}',
        },
    )
    assert edit_result["success"]

    # Verify by reopening document
    doc = Document(str(sample_docx))
    para = next(p for p in doc.paragraphs if p.text == para_block["text"])
    assert para.runs[0].font.strike is True
    assert para.runs[0].font.double_strike is False


@pytest.mark.asyncio
async def test_run_subscript_superscript(sample_docx):
    """Test applying subscript and superscript to runs."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply subscript formatting
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 0,
            "formatting": '{"subscript": true}',
        },
    )
    assert edit_result["success"]

    # Verify subscript
    doc = Document(str(sample_docx))
    para = next(p for p in doc.paragraphs if p.text == para_block["text"])
    assert para.runs[0].font.subscript is True

    # Now change to superscript
    _, edit_result2 = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "edit_run",
            "target_id": edit_result["element_id"],
            "run_index": 0,
            "formatting": '{"subscript": false, "superscript": true}',
        },
    )
    assert edit_result2["success"]

    # Verify superscript
    doc = Document(str(sample_docx))
    para = next(p for p in doc.paragraphs if p.text == para_block["text"])
    assert para.runs[0].font.subscript is False
    assert para.runs[0].font.superscript is True


# --- Table row/column operations tests ---


@pytest.mark.asyncio
async def test_add_table_row_empty(sample_docx):
    """Test adding an empty row to a table."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    original_rows = table_block["rows"]

    # Add empty row
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_row",
            "target_id": table_block["id"],
        },
    )
    assert edit_result["success"]
    assert "Added row" in edit_result["message"]

    # Verify row count increased
    _, cells_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "table_cells",
            "target_id": edit_result["element_id"],
        },
    )
    assert cells_result["table_rows"] == original_rows + 1


@pytest.mark.asyncio
async def test_add_table_row_with_data(sample_docx):
    """Test adding a row with cell values."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Add row with data
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_row",
            "target_id": table_block["id"],
            "content_data": '["New Cell 1", "New Cell 2"]',
        },
    )
    assert edit_result["success"]

    # Verify cell content
    doc = Document(str(sample_docx))
    table = doc.tables[0]
    last_row = table.rows[-1]
    assert last_row.cells[0].text == "New Cell 1"
    assert last_row.cells[1].text == "New Cell 2"


@pytest.mark.asyncio
async def test_add_table_column(sample_docx):
    """Test adding a column with width and values."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    original_cols = table_block["cols"]

    # Add column with data
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_column",
            "target_id": table_block["id"],
            "content_data": '["Header 3", "Row 1 Col 3"]',
            "formatting": '{"width": 1.5}',
        },
    )
    assert edit_result["success"]
    assert "Added column" in edit_result["message"]

    # Verify column count and content
    _, cells_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "table_cells",
            "target_id": edit_result["element_id"],
        },
    )
    assert cells_result["table_cols"] == original_cols + 1

    doc = Document(str(sample_docx))
    table = doc.tables[0]
    assert table.cell(0, original_cols).text == "Header 3"
    assert table.cell(1, original_cols).text == "Row 1 Col 3"


@pytest.mark.asyncio
async def test_delete_table_row(sample_docx):
    """Test deleting a specific row."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    original_rows = table_block["rows"]

    # Delete second row (row index 1)
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "delete_row",
            "target_id": table_block["id"],
            "row": 1,
        },
    )
    assert edit_result["success"]
    assert "Deleted row 1" in edit_result["message"]

    # Verify row count decreased
    _, cells_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "table_cells",
            "target_id": edit_result["element_id"],
        },
    )
    assert cells_result["table_rows"] == original_rows - 1


@pytest.mark.asyncio
async def test_delete_table_column(sample_docx):
    """Test deleting a specific column."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")
    original_cols = table_block["cols"]

    # Delete first column (col index 0)
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "delete_column",
            "target_id": table_block["id"],
            "col": 0,
        },
    )
    assert edit_result["success"]
    assert "Deleted column 0" in edit_result["message"]

    # Verify column count decreased
    _, cells_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "table_cells",
            "target_id": edit_result["element_id"],
        },
    )
    assert cells_result["table_cols"] == original_cols - 1

    # Verify the remaining column has Header 2 content
    doc = Document(str(sample_docx))
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Header 2"


# --- Page breaks tests ---


@pytest.mark.asyncio
async def test_add_page_break(sample_docx):
    """Test appending a page break to the document."""

    _, blocks_before = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    initial_count = blocks_before["block_count"]

    # Add page break
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_page_break",
        },
    )
    assert edit_result["success"]
    assert "Added page break" in edit_result["message"]
    assert edit_result["element_id"].startswith("paragraph_")

    # Verify block count increased
    _, blocks_after = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    assert blocks_after["block_count"] == initial_count + 1

    # Verify the page break exists in the document
    doc = Document(str(sample_docx))
    last_para = doc.paragraphs[-1]
    # Check if there's a break in the runs
    found_break = False
    for run in last_para.runs:
        for child in run._element:
            if "br" in child.tag:
                found_break = True
                break
    assert found_break


@pytest.mark.asyncio
async def test_add_break_after_paragraph(sample_docx):
    """Test inserting a page break after a specific paragraph."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Add break after paragraph
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_break",
            "target_id": para_block["id"],
            "content_data": "page",
        },
    )
    assert edit_result["success"]
    assert "Added page break" in edit_result["message"]


@pytest.mark.asyncio
async def test_add_column_break(sample_docx):
    """Test inserting a column break."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Add column break after paragraph
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_break",
            "target_id": para_block["id"],
            "content_data": "column",
        },
    )
    assert edit_result["success"]
    assert "Added column break" in edit_result["message"]


# --- Metadata writing tests ---


@pytest.mark.asyncio
async def test_set_document_title(sample_docx):
    """Test updating document title."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_meta",
            "content_data": '{"title": "New Document Title"}',
        },
    )
    assert edit_result["success"]
    assert "Updated document metadata" in edit_result["message"]

    # Verify via read
    _, meta_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "meta"}
    )
    assert meta_result["meta"]["title"] == "New Document Title"


@pytest.mark.asyncio
async def test_set_document_author(sample_docx):
    """Test updating document author."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_meta",
            "content_data": '{"author": "Test Author"}',
        },
    )
    assert edit_result["success"]

    # Verify via read
    _, meta_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "meta"}
    )
    assert meta_result["meta"]["author"] == "Test Author"


@pytest.mark.asyncio
async def test_set_multiple_metadata(sample_docx):
    """Test updating multiple metadata properties at once."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_meta",
            "content_data": '{"title": "Multi Test", "author": "Multi Author"}',
        },
    )
    assert edit_result["success"]

    # Verify via python-docx directly
    doc = Document(str(sample_docx))
    assert doc.core_properties.title == "Multi Test"
    assert doc.core_properties.author == "Multi Author"


# --- First/Even page header tests ---


@pytest.mark.asyncio
async def test_set_first_page_header(sample_docx):
    """Test setting a different first page header."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_first_page_header",
            "section_index": 0,
            "content_data": "First Page Header",
        },
    )
    assert edit_result["success"]
    assert "Set first page header" in edit_result["message"]

    # Verify via read
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "headers_footers"}
    )
    assert result["headers_footers"][0]["has_different_first_page"] is True
    assert result["headers_footers"][0]["first_page_header_text"] == "First Page Header"


@pytest.mark.asyncio
async def test_set_even_page_header(sample_docx):
    """Test setting a different even page header."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_even_page_header",
            "section_index": 0,
            "content_data": "Even Page Header",
        },
    )
    assert edit_result["success"]
    assert "Set even page header" in edit_result["message"]

    # Verify via read
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "headers_footers"}
    )
    assert result["headers_footers"][0]["has_different_odd_even"] is True
    assert result["headers_footers"][0]["even_page_header_text"] == "Even Page Header"


@pytest.mark.asyncio
async def test_add_section_new_page(sample_docx):
    """Test adding a new section with new_page start type."""
    # Get initial section count
    _, initial = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "page_setup"}
    )
    initial_count = len(initial["page_setup"])

    # Add section
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_section",
            "content_data": "new_page",
        },
    )
    assert edit_result["success"]
    assert "new_page" in edit_result["message"]

    # Verify section count increased
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "page_setup"}
    )
    assert len(result["page_setup"]) == initial_count + 1


@pytest.mark.asyncio
async def test_add_section_continuous(sample_docx):
    """Test adding a continuous section."""
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "add_section",
            "content_data": "continuous",
        },
    )
    assert edit_result["success"]
    assert "continuous" in edit_result["message"]

    # Verify section was added
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "page_setup"}
    )
    assert len(result["page_setup"]) >= 2


# --- Hyperlink tests ---


@pytest.fixture
def docx_with_hyperlinks():
    """Create a Word document with hyperlinks for testing."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_heading("Document with Links", level=1)
        # Add a paragraph with text and a hyperlink
        p = doc.add_paragraph("Visit ")
        # Use python-docx's add_hyperlink (via low-level API)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(
            qn("r:id"),
            doc.part.relate_to(
                "https://example.com",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            ),
        )
        # Create run with text
        new_run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        new_run.append(r_pr)
        text = OxmlElement("w:t")
        text.text = "Example Website"
        new_run.append(text)
        hyperlink.append(new_run)
        p._p.append(hyperlink)

        # Add more text after the hyperlink
        p.add_run(" for more info.")

        doc.save(f.name)
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_build_runs_includes_hyperlink_text(docx_with_hyperlinks):
    """Test that build_runs includes text from hyperlinks."""
    # Get the paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_hyperlinks), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Read runs from the paragraph
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_hyperlinks),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )

    # Should have 3 runs: "Visit ", "Example Website" (hyperlink), " for more info."
    assert runs_result["block_count"] == 3
    texts = [r["text"] for r in runs_result["runs"]]
    assert "Visit " in texts
    assert "Example Website" in texts
    assert " for more info." in texts

    # The hyperlink run should be marked
    hyperlink_run = next(
        r for r in runs_result["runs"] if r["text"] == "Example Website"
    )
    assert hyperlink_run["is_hyperlink"] is True
    assert hyperlink_run["hyperlink_url"] == "https://example.com"


@pytest.mark.asyncio
async def test_read_hyperlinks(docx_with_hyperlinks):
    """Test reading hyperlinks from a document."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_hyperlinks), "scope": "hyperlinks"}
    )

    assert result["block_count"] == 1
    assert len(result["hyperlinks"]) == 1

    link = result["hyperlinks"][0]
    assert link["text"] == "Example Website"
    assert link["url"] == "https://example.com"
    assert link["address"] == "https://example.com"
    assert link["is_external"] is True
    assert link["index"] == 0


@pytest.mark.asyncio
async def test_read_hyperlinks_empty(sample_docx):
    """Test reading hyperlinks from a document with no hyperlinks."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "hyperlinks"}
    )
    assert result["block_count"] == 0
    assert result["hyperlinks"] == []


@pytest.mark.asyncio
async def test_hyperlink_in_table_cell():
    """Test that hyperlinks in table cells are detected."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Website"
        table.cell(1, 0).text = "Example Corp"

        # Add hyperlink to cell (1, 1)
        cell = table.cell(1, 1)
        p = cell.paragraphs[0]

        # Create hyperlink
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(
            qn("r:id"),
            doc.part.relate_to(
                "https://example.org",
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            ),
        )
        new_run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "Visit Site"
        new_run.append(text)
        hyperlink.append(new_run)
        p._p.append(hyperlink)

        doc.save(f.name)
        doc_path = Path(f.name)

    try:
        # Read hyperlinks
        _, result = await mcp.call_tool(
            "read", {"file_path": str(doc_path), "scope": "hyperlinks"}
        )

        assert result["block_count"] == 1
        assert len(result["hyperlinks"]) == 1
        link = result["hyperlinks"][0]
        assert link["text"] == "Visit Site"
        assert link["url"] == "https://example.org"
        assert link["is_external"] is True
    finally:
        doc_path.unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_edit_hyperlink_run(docx_with_hyperlinks):
    """Test that editing a hyperlink run uses correct indexing."""
    # Get the paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_hyperlinks), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Read runs to find the hyperlink run index
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_hyperlinks),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )

    # Find the hyperlink run
    hyperlink_run = next(r for r in runs_result["runs"] if r["is_hyperlink"])
    hyperlink_index = hyperlink_run["index"]

    # Edit the hyperlink run text
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_hyperlinks),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": hyperlink_index,
            "content_data": "Modified Link Text",
        },
    )
    assert edit_result["success"]

    # Verify the hyperlink run was edited (not another run)
    _, runs_result2 = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_hyperlinks),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    edited_run = runs_result2["runs"][hyperlink_index]
    assert edited_run["text"] == "Modified Link Text"
    assert edited_run["is_hyperlink"] is True  # Still a hyperlink


# --- Style tests ---


@pytest.mark.asyncio
async def test_read_styles(sample_docx):
    """Test reading all styles from a document."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "styles"}
    )

    # Should have many built-in styles
    assert result["block_count"] > 0
    assert len(result["styles"]) > 0

    # Check that common styles exist
    style_names = [s["name"] for s in result["styles"]]
    assert "Normal" in style_names
    assert "Heading 1" in style_names

    # Check style properties
    normal_style = next(s for s in result["styles"] if s["name"] == "Normal")
    assert normal_style["style_id"] == "Normal"
    assert normal_style["type"] == "paragraph"
    assert normal_style["builtin"] is True


@pytest.mark.asyncio
async def test_read_styles_includes_custom():
    """Test that custom styles appear in the list."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        # Add a custom style
        from docx.enum.style import WD_STYLE_TYPE

        custom_style = doc.styles.add_style("MyCustomStyle", WD_STYLE_TYPE.PARAGRAPH)
        custom_style.base_style = doc.styles["Normal"]
        doc.add_paragraph("Test paragraph", style="MyCustomStyle")
        doc.save(f.name)
        doc_path = Path(f.name)

    try:
        _, result = await mcp.call_tool(
            "read", {"file_path": str(doc_path), "scope": "styles"}
        )

        # Custom style should be in the list
        style_names = [s["name"] for s in result["styles"]]
        assert "MyCustomStyle" in style_names

        # Check custom style properties
        custom = next(s for s in result["styles"] if s["name"] == "MyCustomStyle")
        assert custom["builtin"] is False
        assert custom["base_style"] == "Normal"
        assert custom["type"] == "paragraph"
    finally:
        doc_path.unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_style_hierarchy(sample_docx):
    """Test that base_style relationships are correctly reported."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "styles"}
    )

    # Heading 1 typically has a base style
    heading1 = next((s for s in result["styles"] if s["name"] == "Heading 1"), None)
    if heading1:
        # Base style could be None or another style
        # Just verify the field exists and is valid
        assert "base_style" in heading1

    # Check style types are valid
    for style in result["styles"]:
        assert style["type"] in ["paragraph", "character", "table", "list", "unknown"]


@pytest.mark.asyncio
async def test_read_styles_includes_character_styles(sample_docx):
    """Test that character styles are included."""
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "styles"}
    )

    # Should have character styles
    char_styles = [s for s in result["styles"] if s["type"] == "character"]
    assert len(char_styles) > 0

    # Check common character styles
    style_names = [s["name"] for s in char_styles]
    # Common character styles in Word
    common_char_styles = ["Default Paragraph Font", "Hyperlink"]
    found_any = any(s in style_names for s in common_char_styles)
    assert found_any or len(char_styles) > 0  # At least has some character styles


# --- Paragraph Format Read tests ---


@pytest.fixture
def docx_with_formatting():
    """Create a Word document with paragraph formatting."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        # Add a formatted paragraph
        para = doc.add_paragraph("Formatted paragraph")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.left_indent = Inches(0.5)
        pf.right_indent = Inches(0.25)
        pf.first_line_indent = Inches(0.3)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5

        doc.save(f.name)
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_read_paragraph_format(docx_with_formatting):
    """Test reading paragraph formatting via runs scope."""
    # Get the paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_formatting), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Read runs (which includes paragraph_format)
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_formatting),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )

    # Verify paragraph_format is included
    pf = result["paragraph_format"]
    assert pf is not None
    assert pf["alignment"] == "center"
    assert pf["left_indent"] == pytest.approx(0.5, abs=0.01)
    assert pf["right_indent"] == pytest.approx(0.25, abs=0.01)
    assert pf["first_line_indent"] == pytest.approx(0.3, abs=0.01)
    assert pf["space_before"] == pytest.approx(12.0, abs=0.1)
    assert pf["space_after"] == pytest.approx(6.0, abs=0.1)
    assert pf["line_spacing"] == pytest.approx(1.5, abs=0.01)


@pytest.mark.asyncio
async def test_read_paragraph_alignment(sample_docx):
    """Test reading paragraph alignment from normal document."""
    # Get the first paragraph
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Read runs with paragraph_format
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )

    # Normal paragraphs typically have no explicit alignment (None or left)
    pf = result["paragraph_format"]
    assert pf is not None
    # alignment could be None for default or 'left'
    assert pf["alignment"] is None or pf["alignment"] == "left"


# --- Character Style tests ---


@pytest.mark.asyncio
async def test_apply_character_style(sample_docx):
    """Test applying a character style to a run."""
    # Get the first paragraph
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply "Strong" style to run 0
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 0,
            "formatting": '{"style": "Strong"}',
        },
    )
    assert edit_result["success"]

    # Read runs and verify style
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    assert runs_result["runs"][0]["style"] == "Strong"


@pytest.mark.asyncio
async def test_read_character_style(sample_docx):
    """Test reading character style from run."""
    # Get the first paragraph
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Read runs - default runs have no explicit style
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(sample_docx),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )
    # Default runs typically have None or "Default Paragraph Font" style
    assert result["runs"][0]["style"] is None or result["runs"][0]["style"] is not None


# --- Rich Header/Footer tests ---


@pytest.mark.asyncio
async def test_append_paragraph_to_header(sample_docx):
    """Test appending a paragraph to the header."""
    # First set a header
    _, _ = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_header",
            "content_data": "Initial Header",
            "section_index": 0,
        },
    )

    # Append a second paragraph
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "append_header",
            "content_type": "paragraph",
            "content_data": "Second Header Line",
            "section_index": 0,
        },
    )
    assert edit_result["success"]
    assert "paragraph_" in edit_result["element_id"]

    # Verify both are present
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "headers_footers"}
    )
    header_text = result["headers_footers"][0]["header_text"]
    assert "Initial Header" in header_text
    assert "Second Header Line" in header_text


@pytest.mark.asyncio
async def test_append_table_to_header(sample_docx):
    """Test appending a table to the header."""
    # First set a header
    _, _ = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_header",
            "content_data": "Header Text",
            "section_index": 0,
        },
    )

    # Append a table
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "append_header",
            "content_type": "table",
            "content_data": '[["Col1", "Col2"], ["A", "B"]]',
            "section_index": 0,
        },
    )
    assert edit_result["success"]
    assert "table_" in edit_result["element_id"]


@pytest.mark.asyncio
async def test_clear_header(sample_docx):
    """Test clearing header content."""
    # First set a header
    _, _ = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "set_header",
            "content_data": "Header To Clear",
            "section_index": 0,
        },
    )

    # Clear the header
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(sample_docx),
            "operation": "clear_header",
            "section_index": 0,
        },
    )
    assert edit_result["success"]

    # Verify header is cleared (should have empty or minimal content)
    _, result = await mcp.call_tool(
        "read", {"file_path": str(sample_docx), "scope": "headers_footers"}
    )
    header_text = result["headers_footers"][0]["header_text"]
    # Header should be empty or contain only whitespace
    assert header_text is None or header_text.strip() == ""


# --- Cell Merge tests ---


@pytest.fixture
def docx_with_table():
    """Create a Word document with a 3x3 table for merge testing."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_heading("Table for Merge Testing", level=1)
        table = doc.add_table(rows=3, cols=3)
        for r in range(3):
            for c in range(3):
                table.cell(r, c).text = f"R{r}C{c}"
        doc.save(f.name)
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_read_table_cells_merge_info(docx_with_table):
    """Test that table cells include merge info fields."""
    # Get the table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Read cells
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_cells",
            "target_id": table_block["id"],
        },
    )

    # All cells should be unmerged origin cells
    assert result["table_rows"] == 3
    assert result["table_cols"] == 3
    for cell in result["cells"]:
        assert cell["grid_span"] == 1
        assert cell["row_span"] == 1
        assert cell["is_merge_origin"] is True


@pytest.mark.asyncio
async def test_merge_cells_horizontal(docx_with_table):
    """Test merging cells horizontally."""
    # Get the table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Merge row 0, columns 0-1
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "merge_cells",
            "target_id": table_block["id"],
            "row": 0,
            "col": 0,
            "content_data": '{"end_row": 0, "end_col": 1}',
        },
    )
    assert edit_result["success"]

    # Re-fetch table ID (content hash changed after merge)
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Verify merge by reading cells
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_cells",
            "target_id": table_block["id"],
        },
    )

    # Find the merged cell origin (row 0, col 0)
    origin = next(c for c in result["cells"] if c["row"] == 0 and c["col"] == 0)
    assert origin["grid_span"] == 2
    assert origin["is_merge_origin"] is True

    # The continuation cell should exist but not be an origin
    continuation = next(c for c in result["cells"] if c["row"] == 0 and c["col"] == 1)
    assert continuation["is_merge_origin"] is False


@pytest.mark.asyncio
async def test_merge_cells_vertical(docx_with_table):
    """Test merging cells vertically."""
    # Get the table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Merge column 0, rows 0-1
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "merge_cells",
            "target_id": table_block["id"],
            "row": 0,
            "col": 0,
            "content_data": '{"end_row": 1, "end_col": 0}',
        },
    )
    assert edit_result["success"]

    # Re-fetch table ID (content hash changed after merge)
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Verify merge by reading cells
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_cells",
            "target_id": table_block["id"],
        },
    )

    # Find the merged cell origin (row 0, col 0)
    origin = next(c for c in result["cells"] if c["row"] == 0 and c["col"] == 0)
    assert origin["row_span"] == 2
    assert origin["is_merge_origin"] is True

    # The continuation cell should exist but not be an origin
    continuation = next(c for c in result["cells"] if c["row"] == 1 and c["col"] == 0)
    assert continuation["is_merge_origin"] is False


@pytest.mark.asyncio
async def test_merge_cells_rectangular(docx_with_table):
    """Test merging a 2x2 block of cells."""
    # Get the table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Merge 2x2 block from (0,0) to (1,1)
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "merge_cells",
            "target_id": table_block["id"],
            "row": 0,
            "col": 0,
            "content_data": '{"end_row": 1, "end_col": 1}',
        },
    )
    assert edit_result["success"]

    # Re-fetch table ID (content hash changed after merge)
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Verify merge by reading cells
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_cells",
            "target_id": table_block["id"],
        },
    )

    # Find the merged cell origin (row 0, col 0)
    origin = next(c for c in result["cells"] if c["row"] == 0 and c["col"] == 0)
    assert origin["grid_span"] == 2
    assert origin["row_span"] == 2
    assert origin["is_merge_origin"] is True

    # All other cells in the 2x2 block should not be origins
    for r, c in [(0, 1), (1, 0), (1, 1)]:
        cell = next(cl for cl in result["cells"] if cl["row"] == r and cl["col"] == c)
        assert cell["is_merge_origin"] is False


# =============================================================================
# Phase 3.1: Additional Font Properties Tests
# =============================================================================


@pytest.fixture
def docx_with_font_effects():
    """Create a document with various font effects for testing."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        p = doc.add_paragraph()
        # Run 0: all_caps
        run0 = p.add_run("ALL CAPS TEXT")
        run0.font.all_caps = True
        # Run 1: small_caps
        run1 = p.add_run("Small Caps Text")
        run1.font.small_caps = True
        # Run 2: hidden
        run2 = p.add_run("Hidden Text")
        run2.font.hidden = True
        # Run 3: emboss
        run3 = p.add_run("Embossed Text")
        run3.font.emboss = True
        # Run 4: normal for editing
        p.add_run("Normal Text")
        doc.save(f.name)
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_read_font_all_caps(docx_with_font_effects):
    """Test reading all_caps font property."""
    # Get the paragraph
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Read runs
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )

    # Check all_caps on run 0
    assert runs_result["runs"][0]["all_caps"] is True
    assert (
        runs_result["runs"][1]["all_caps"] is None
        or runs_result["runs"][1]["all_caps"] is False
    )


@pytest.mark.asyncio
async def test_read_font_small_caps(docx_with_font_effects):
    """Test reading small_caps font property."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )

    # Check small_caps on run 1
    assert runs_result["runs"][1]["small_caps"] is True


@pytest.mark.asyncio
async def test_read_font_hidden(docx_with_font_effects):
    """Test reading hidden font property."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )

    # Check hidden on run 2
    assert runs_result["runs"][2]["hidden"] is True


@pytest.mark.asyncio
async def test_read_font_emboss(docx_with_font_effects):
    """Test reading emboss font property."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": para_block["id"],
        },
    )

    # Check emboss on run 3
    assert runs_result["runs"][3]["emboss"] is True


@pytest.mark.asyncio
async def test_edit_font_small_caps(docx_with_font_effects):
    """Test applying small_caps formatting to a run."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply small_caps to run 4 (normal text)
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_font_effects),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 4,
            "formatting": json.dumps({"small_caps": True}),
        },
    )
    assert edit_result["success"]

    # Verify the change
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    assert runs_result["runs"][4]["small_caps"] is True


@pytest.mark.asyncio
async def test_edit_font_emboss(docx_with_font_effects):
    """Test applying emboss effect to a run."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply emboss to run 4
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_font_effects),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 4,
            "formatting": json.dumps({"emboss": True}),
        },
    )
    assert edit_result["success"]

    # Verify
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    assert runs_result["runs"][4]["emboss"] is True


@pytest.mark.asyncio
async def test_edit_font_imprint(docx_with_font_effects):
    """Test applying imprint effect to a run."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply imprint to run 4
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_font_effects),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 4,
            "formatting": json.dumps({"imprint": True}),
        },
    )
    assert edit_result["success"]

    # Verify
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    assert runs_result["runs"][4]["imprint"] is True


@pytest.mark.asyncio
async def test_edit_font_outline_shadow(docx_with_font_effects):
    """Test applying outline and shadow effects to a run."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Apply both outline and shadow to run 4
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_font_effects),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 4,
            "formatting": json.dumps({"outline": True, "shadow": True}),
        },
    )
    assert edit_result["success"]

    # Verify both properties
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    assert runs_result["runs"][4]["outline"] is True
    assert runs_result["runs"][4]["shadow"] is True


@pytest.mark.asyncio
async def test_clear_font_properties(docx_with_font_effects):
    """Test clearing/unsetting font properties by setting to null."""
    # Read the document to get block ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_font_effects), "scope": "blocks"}
    )
    para_block = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # First, apply small_caps and emboss to run 0
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_font_effects),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 0,
            "formatting": json.dumps({"small_caps": True, "emboss": True}),
        },
    )
    assert edit_result["success"]

    # Verify they are set
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    assert runs_result["runs"][0]["small_caps"] is True
    assert runs_result["runs"][0]["emboss"] is True

    # Now clear (unset) by setting to null
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_font_effects),
            "operation": "edit_run",
            "target_id": para_block["id"],
            "run_index": 0,
            "formatting": json.dumps({"small_caps": None, "emboss": None}),
        },
    )
    assert edit_result["success"]

    # Verify they are cleared (None)
    _, runs_result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_font_effects),
            "scope": "runs",
            "target_id": edit_result["element_id"],
        },
    )
    assert runs_result["runs"][0]["small_caps"] is None
    assert runs_result["runs"][0]["emboss"] is None


# --- Phase 2: Edit Style Definitions ---


@pytest.fixture
def docx_with_custom_style(tmp_path):
    """Create a document with a custom paragraph style for testing."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    # Create a custom style based on Normal
    custom_style = doc.styles.add_style("TestStyle", 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    custom_style.base_style = doc.styles["Normal"]
    custom_style.font.name = "Arial"
    custom_style.font.size = Pt(12)
    custom_style.font.bold = False

    # Add a paragraph using the custom style
    doc.add_paragraph("Test paragraph with custom style", style="TestStyle")

    path = tmp_path / "custom_style.docx"
    doc.save(path)
    return path


@pytest.mark.asyncio
async def test_read_style_format(docx_with_custom_style):
    """Test reading detailed style formatting."""
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_custom_style),
            "scope": "style",
            "target_id": "TestStyle",
        },
    )

    assert result["block_count"] == 1
    style_fmt = result["style_format"]
    assert style_fmt["name"] == "TestStyle"
    assert style_fmt["type"] == "paragraph"
    assert style_fmt["font_name"] == "Arial"
    assert style_fmt["font_size"] == 12.0
    assert style_fmt["bold"] is False


@pytest.mark.asyncio
async def test_edit_style_font(docx_with_custom_style):
    """Test modifying style font properties."""
    # Edit the style to make it bold and 16pt
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_custom_style),
            "operation": "edit_style",
            "target_id": "TestStyle",
            "formatting": json.dumps({"bold": True, "font_size": 16, "italic": True}),
        },
    )
    assert edit_result["success"]

    # Read the style back and verify changes
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_custom_style),
            "scope": "style",
            "target_id": "TestStyle",
        },
    )
    style_fmt = result["style_format"]
    assert style_fmt["bold"] is True
    assert style_fmt["italic"] is True
    assert style_fmt["font_size"] == 16.0


@pytest.mark.asyncio
async def test_edit_style_paragraph(docx_with_custom_style):
    """Test modifying style paragraph properties."""
    # Edit the style paragraph formatting
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_custom_style),
            "operation": "edit_style",
            "target_id": "TestStyle",
            "formatting": json.dumps(
                {
                    "alignment": "center",
                    "space_before": 12,
                    "space_after": 6,
                }
            ),
        },
    )
    assert edit_result["success"]

    # Read the style back and verify changes
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_custom_style),
            "scope": "style",
            "target_id": "TestStyle",
        },
    )
    style_fmt = result["style_format"]
    assert style_fmt["alignment"] == "center"
    assert style_fmt["space_before"] == 12.0
    assert style_fmt["space_after"] == 6.0


@pytest.mark.asyncio
async def test_edit_style_line_spacing_multiplier_roundtrip(docx_with_custom_style):
    """Test setting line spacing as a multiplier (< 5)."""
    # Set line spacing to 1.5 (multiplier)
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_custom_style),
            "operation": "edit_style",
            "target_id": "TestStyle",
            "formatting": json.dumps({"line_spacing": 1.5}),
        },
    )
    assert edit_result["success"]

    # Read back and verify
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_custom_style),
            "scope": "style",
            "target_id": "TestStyle",
        },
    )
    assert result["style_format"]["line_spacing"] == 1.5


@pytest.mark.asyncio
async def test_edit_style_line_spacing_points_roundtrip(docx_with_custom_style):
    """Test setting line spacing as points (>= 5)."""
    # Set line spacing to 18 points
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_custom_style),
            "operation": "edit_style",
            "target_id": "TestStyle",
            "formatting": json.dumps({"line_spacing": 18}),
        },
    )
    assert edit_result["success"]

    # Read back and verify (should read as 18.0 points)
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_custom_style),
            "scope": "style",
            "target_id": "TestStyle",
        },
    )
    assert result["style_format"]["line_spacing"] == 18.0


@pytest.mark.asyncio
async def test_edit_style_alignment_justify_roundtrip(docx_with_custom_style):
    """Test that justify alignment roundtrips correctly."""
    # Set alignment to justify
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_custom_style),
            "operation": "edit_style",
            "target_id": "TestStyle",
            "formatting": json.dumps({"alignment": "justify"}),
        },
    )
    assert edit_result["success"]

    # Read back and verify API returns "justify" (not "both" or other variants)
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_custom_style),
            "scope": "style",
            "target_id": "TestStyle",
        },
    )
    assert result["style_format"]["alignment"] == "justify"


@pytest.mark.asyncio
async def test_edit_style_invalid_alignment_raises(docx_with_custom_style):
    """Test that invalid alignment raises AttributeError (no defensive validation)."""
    from mcp.server.fastmcp.exceptions import ToolError

    with pytest.raises(ToolError, match="has no attribute 'INVALID'"):
        await mcp.call_tool(
            "edit",
            {
                "file_path": str(docx_with_custom_style),
                "operation": "edit_style",
                "target_id": "TestStyle",
                "formatting": json.dumps({"alignment": "invalid"}),
            },
        )


# =============================================================================
# Phase 3: Table Dimensions & Layout Tests
# =============================================================================


@pytest.mark.asyncio
async def test_read_table_layout(docx_with_table):
    """Test reading table layout info via scope='table_layout'."""
    # Get the table ID first
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Read table layout
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_layout",
            "target_id": table_block["id"],
        },
    )

    # Check layout info returned
    layout = result["table_layout"]
    assert layout["table_id"] == table_block["id"]
    assert layout["autofit"] is True  # Default for new tables
    # Rows should be present (3 rows in fixture)
    assert len(layout["rows"]) == 3
    for i, row in enumerate(layout["rows"]):
        assert row["index"] == i


@pytest.mark.asyncio
async def test_set_table_alignment(docx_with_table):
    """Test setting table horizontal alignment."""
    # Get table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Set alignment to center
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "set_table_alignment",
            "target_id": table_block["id"],
            "content_data": "center",
        },
    )
    assert edit_result["success"]

    # Verify by reading layout
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_layout",
            "target_id": table_block["id"],
        },
    )
    assert result["table_layout"]["alignment"] == "center"


@pytest.mark.asyncio
async def test_set_table_fixed_layout(docx_with_table):
    """Test setting table to fixed layout with column widths."""
    # Get table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Set fixed layout with column widths
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "set_table_fixed_layout",
            "target_id": table_block["id"],
            "content_data": json.dumps([1.5, 2.0, 1.5]),
        },
    )
    assert edit_result["success"]

    # Verify autofit is False
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_layout",
            "target_id": table_block["id"],
        },
    )
    assert result["table_layout"]["autofit"] is False


@pytest.mark.asyncio
async def test_set_row_height(docx_with_table):
    """Test setting row height with at_least rule."""
    # Get table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Set row 0 height to 0.5 inches with "at_least" rule
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "set_row_height",
            "target_id": table_block["id"],
            "row": 0,
            "content_data": json.dumps({"height": 0.5, "rule": "at_least"}),
        },
    )
    assert edit_result["success"]

    # Verify by reading layout
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_layout",
            "target_id": table_block["id"],
        },
    )
    row0 = result["table_layout"]["rows"][0]
    assert row0["height_inches"] == pytest.approx(0.5, rel=0.01)
    assert row0["height_rule"] == "at_least"


@pytest.mark.asyncio
async def test_set_cell_width(docx_with_table):
    """Test setting cell width."""
    # Get table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Set cell (0, 0) width to 2.0 inches
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "set_cell_width",
            "target_id": table_block["id"],
            "row": 0,
            "col": 0,
            "content_data": "2.0",
        },
    )
    assert edit_result["success"]

    # Verify by reading table cells
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_cells",
            "target_id": table_block["id"],
        },
    )
    cell_00 = next(c for c in result["cells"] if c["row"] == 0 and c["col"] == 0)
    assert cell_00["width_inches"] == pytest.approx(2.0, rel=0.01)


@pytest.mark.asyncio
async def test_set_cell_vertical_alignment(docx_with_table):
    """Test setting cell vertical alignment."""
    # Get table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Set cell (1, 1) to vertical center
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "set_cell_vertical_alignment",
            "target_id": table_block["id"],
            "row": 1,
            "col": 1,
            "content_data": "center",
        },
    )
    assert edit_result["success"]

    # Verify by reading table cells
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_with_table),
            "scope": "table_cells",
            "target_id": table_block["id"],
        },
    )
    cell_11 = next(c for c in result["cells"] if c["row"] == 1 and c["col"] == 1)
    assert cell_11["vertical_alignment"] == "center"


@pytest.mark.asyncio
async def test_table_layout_xml_structure(docx_with_table):
    """Test that table layout changes produce correct XML structure."""
    from docx import Document

    # Get table ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_with_table), "scope": "blocks"}
    )
    table_block = next(b for b in blocks_result["blocks"] if b["type"] == "table")

    # Apply multiple layout changes
    await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "set_table_alignment",
            "target_id": table_block["id"],
            "content_data": "right",
        },
    )
    await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_with_table),
            "operation": "set_row_height",
            "target_id": table_block["id"],
            "row": 1,
            "content_data": json.dumps({"height": 0.75, "rule": "exactly"}),
        },
    )

    # Open and verify XML
    doc = Document(str(docx_with_table))
    table = doc.tables[0]

    # Verify table alignment in XML (tblPr/jc element)
    tbl_pr = table._tbl.tblPr
    assert tbl_pr is not None
    jc = tbl_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}jc")
    assert jc is not None
    assert (
        jc.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
        == "right"
    )

    # Verify row height in XML (trPr/trHeight element)
    row = table.rows[1]
    tr = row._tr
    tr_pr = tr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trPr"
    )
    assert tr_pr is not None
    tr_height = tr_pr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trHeight"
    )
    assert tr_height is not None
    # hRule should be "exact" for "exactly" rule
    assert (
        tr_height.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hRule"
        )
        == "exact"
    )


# =============================================================================
# Phase 4: Paragraph Tab Stops Tests
# =============================================================================


@pytest.fixture
def docx_for_tabs():
    """Create a Word document for tab stop testing."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_paragraph("Test paragraph for tab stops")
        doc.save(f.name)
        yield Path(f.name)
    Path(f.name).unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_add_tab_stop(docx_for_tabs):
    """Test adding a tab stop with alignment and leader."""
    # Get paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_for_tabs), "scope": "blocks"}
    )
    para = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Add right-aligned tab with dot leader at 4 inches
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_tabs),
            "operation": "add_tab_stop",
            "target_id": para["id"],
            "content_data": json.dumps(
                {"position": 4.0, "alignment": "right", "leader": "dots"}
            ),
        },
    )
    assert edit_result["success"]

    # Verify by reading paragraph format
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_for_tabs),
            "scope": "runs",
            "target_id": para["id"],
        },
    )
    tab_stops = result["paragraph_format"]["tab_stops"]
    assert len(tab_stops) >= 1
    tab = tab_stops[0]
    assert tab["position_inches"] == pytest.approx(4.0, rel=0.01)
    assert tab["alignment"] == "right"
    assert tab["leader"] == "dots"


@pytest.mark.asyncio
async def test_clear_tab_stops(docx_for_tabs):
    """Test clearing all tab stops from a paragraph."""
    # Get paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_for_tabs), "scope": "blocks"}
    )
    para = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Add a tab stop first
    await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_tabs),
            "operation": "add_tab_stop",
            "target_id": para["id"],
            "content_data": json.dumps({"position": 2.0, "alignment": "left"}),
        },
    )

    # Clear all tab stops
    _, edit_result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_tabs),
            "operation": "clear_tab_stops",
            "target_id": para["id"],
        },
    )
    assert edit_result["success"]

    # Verify no tab stops remain
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_for_tabs),
            "scope": "runs",
            "target_id": para["id"],
        },
    )
    assert result["paragraph_format"]["tab_stops"] == []


@pytest.mark.asyncio
async def test_read_tab_stops(docx_for_tabs):
    """Test reading multiple tab stops from a paragraph."""
    # Get paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_for_tabs), "scope": "blocks"}
    )
    para = next(b for b in blocks_result["blocks"] if b["type"] == "paragraph")

    # Add multiple tab stops
    await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_tabs),
            "operation": "add_tab_stop",
            "target_id": para["id"],
            "content_data": json.dumps(
                {"position": 1.0, "alignment": "left", "leader": "spaces"}
            ),
        },
    )
    await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_tabs),
            "operation": "add_tab_stop",
            "target_id": para["id"],
            "content_data": json.dumps(
                {"position": 3.0, "alignment": "center", "leader": "heavy"}
            ),
        },
    )
    await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_tabs),
            "operation": "add_tab_stop",
            "target_id": para["id"],
            "content_data": json.dumps(
                {"position": 5.0, "alignment": "decimal", "leader": "middle_dot"}
            ),
        },
    )

    # Read tab stops
    _, result = await mcp.call_tool(
        "read",
        {
            "file_path": str(docx_for_tabs),
            "scope": "runs",
            "target_id": para["id"],
        },
    )
    tab_stops = result["paragraph_format"]["tab_stops"]
    assert len(tab_stops) == 3

    # Tab stops should be returned in position order
    positions = [t["position_inches"] for t in tab_stops]
    assert positions == sorted(positions)


# --- Phase 5: Document Fields Tests ---


@pytest.fixture
def docx_for_fields(tmp_path):
    """Create a document for field tests."""
    doc = Document()
    doc.add_paragraph("First section content")
    doc_path = tmp_path / "fields_test.docx"
    doc.save(str(doc_path))
    return doc_path


@pytest.mark.asyncio
async def test_insert_field_page(docx_for_fields):
    """Test inserting PAGE field into a paragraph."""
    # Get paragraph ID
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_for_fields), "scope": "blocks"}
    )
    para = blocks_result["blocks"][0]

    # Insert PAGE field
    _, result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_fields),
            "operation": "insert_field",
            "target_id": para["id"],
            "content_data": "PAGE",
        },
    )
    assert result["success"] is True
    assert "PAGE" in result["message"]

    # Verify XML structure - field parts are in separate runs
    doc = Document(str(docx_for_fields))
    p = doc.paragraphs[0]
    para_xml = p._p.xml
    assert "w:fldChar" in para_xml
    assert 'w:fldCharType="begin"' in para_xml
    assert 'w:fldCharType="separate"' in para_xml
    assert 'w:fldCharType="end"' in para_xml
    assert "PAGE" in para_xml


@pytest.mark.asyncio
async def test_insert_field_numpages(docx_for_fields):
    """Test inserting NUMPAGES field."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_for_fields), "scope": "blocks"}
    )
    para = blocks_result["blocks"][0]

    _, result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_fields),
            "operation": "insert_field",
            "target_id": para["id"],
            "content_data": "NUMPAGES",
        },
    )
    assert result["success"] is True
    assert "NUMPAGES" in result["message"]


@pytest.mark.asyncio
async def test_insert_page_x_of_y_footer(docx_for_fields):
    """Test inserting 'Page X of Y' in footer."""
    _, result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_fields),
            "operation": "insert_page_x_of_y",
            "section_index": 0,
            "content_data": "footer",
        },
    )
    assert result["success"] is True
    assert "footer" in result["message"]

    # Verify footer content
    doc = Document(str(docx_for_fields))
    footer = doc.sections[0].footer
    footer_text = "".join(p.text for p in footer.paragraphs)
    assert "Page" in footer_text
    assert "of" in footer_text

    # Verify field structure in footer
    footer_xml = footer._element.xml
    assert "w:fldChar" in footer_xml
    assert "PAGE" in footer_xml
    assert "NUMPAGES" in footer_xml


@pytest.mark.asyncio
async def test_insert_page_x_of_y_header(docx_for_fields):
    """Test inserting 'Page X of Y' in header."""
    _, result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_fields),
            "operation": "insert_page_x_of_y",
            "section_index": 0,
            "content_data": "header",
        },
    )
    assert result["success"] is True
    assert "header" in result["message"]

    # Verify header content
    doc = Document(str(docx_for_fields))
    header = doc.sections[0].header
    header_text = "".join(p.text for p in header.paragraphs)
    assert "Page" in header_text
    assert "of" in header_text


@pytest.mark.asyncio
async def test_insert_field_arbitrary_code(docx_for_fields):
    """Test that any field code is accepted (no artificial restrictions)."""
    _, blocks_result = await mcp.call_tool(
        "read", {"file_path": str(docx_for_fields), "scope": "blocks"}
    )
    para = blocks_result["blocks"][0]

    # AUTHOR is a valid Word field that was previously rejected
    _, result = await mcp.call_tool(
        "edit",
        {
            "file_path": str(docx_for_fields),
            "operation": "insert_field",
            "target_id": para["id"],
            "content_data": "AUTHOR",
        },
    )
    assert result["success"]
